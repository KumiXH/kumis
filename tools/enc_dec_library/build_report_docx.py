from __future__ import annotations

import csv
import json
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from tools.enc_dec_library.config import PAPERS, ROOT


REPORT = ROOT / "report" / "DiT编解码器发展架构训练与数据工程深度洞察.docx"
METADATA = ROOT / "metadata"
FIGURES = ROOT / "figures"

NAVY = RGBColor(22, 50, 79)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
RED = RGBColor(181, 27, 52)
INK = RGBColor(23, 33, 43)
MUTED = RGBColor(82, 99, 112)
GREEN = RGBColor(7, 132, 95)
GOLD = RGBColor(122, 90, 0)
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E5F1EA"
LIGHT_GOLD = "FFF2CC"
LIGHT_GRAY = "F4F6F9"


def read_csv(name: str):
    with (METADATA / name).open("r", encoding="utf-8-sig", newline="") as stream:
        return list(csv.DictReader(stream))


ARCH = read_csv("architecture_matrix.csv")
TRAINING = read_csv("training_matrix.csv")
DATASETS = read_csv("dataset_matrix.csv")
FLASH = read_csv("flashvsr_modules.csv")
TERMS = read_csv("terminology.csv")
SOURCES = json.loads((METADATA / "source_manifest.json").read_text(encoding="utf-8"))
SOURCE_BY_KEY = {row["key"]: row for row in SOURCES}
REF_NO = {paper["key"]: index for index, paper in enumerate(PAPERS, start=1)}


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths_in[index] * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=11, bold=False, italic=False, color=INK, font="Calibri"):
    run.font.name = font
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), font)
    r_pr.rFonts.set(qn("w:hAnsi"), font)
    r_pr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run(run, **kwargs)
    return run


def add_para(doc, text="", *, before=0, after=6, line=1.25, align=None, keep=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_together = keep
    if align is not None:
        paragraph.alignment = align
    if text:
        add_text(paragraph, text)
    return paragraph


def add_mixed(doc, parts, *, before=0, after=6, line=1.25, align=None):
    paragraph = add_para(doc, before=before, after=after, line=line, align=align)
    for text, kwargs in parts:
        add_text(paragraph, text, **kwargs)
    return paragraph


def heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    add_text(
        paragraph,
        text,
        size={1: 16, 2: 13, 3: 12}[level],
        bold=True,
        color=BLUE if level < 3 else DARK_BLUE,
    )
    return paragraph


def bullet(doc, text, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    add_text(paragraph, text)
    return paragraph


def numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    add_text(paragraph, text)
    return paragraph


def equation(doc, text, label=""):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_text(p, text, size=10.5, font="Consolas", color=NAVY)
    if label:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_after = Pt(0)
        add_text(p2, label, size=8.5, italic=True, color=MUTED)
    add_para(doc, after=3)


def callout(doc, label, text, fill=LIGHT_GRAY, label_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    add_text(p, f"{label} | ", size=10.5, bold=True, color=label_color)
    add_text(p, text, size=10.5)
    add_para(doc, after=4)


def add_table(doc, headers, rows, widths, font_size=9.2, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_text(p, header, size=font_size, bold=True, color=DARK_BLUE)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            add_text(p, str(value), size=font_size)
    add_para(doc, after=5)
    return table


def add_source_note(doc, text):
    p = add_para(doc, before=2, after=5, line=1.15)
    add_text(p, "证据锚点 / Evidence anchor: ", size=8.5, bold=True, color=MUTED)
    add_text(p, text, size=8.5, italic=True, color=MUTED)


def image_size(path: Path, max_width=6.2, max_height=6.65):
    with Image.open(path) as image:
        width, height = image.size
    scale = min(max_width / width, max_height / height)
    return Inches(width * scale), Inches(height * scale)


def add_figure(doc, path: Path, caption: str, source: str, *, max_width=6.2, max_height=6.65):
    if not path.exists() or path.suffix == ".part":
        return False
    try:
        width, height = image_size(path, max_width=max_width, max_height=max_height)
    except Exception:
        return False
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    inline_shape = paragraph.add_run().add_picture(str(path), width=width, height=height)
    try:
        doc_pr = inline_shape._inline.docPr
        doc_pr.set("descr", caption)
        doc_pr.set("title", path.stem)
    except Exception:
        pass
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(2)
    cp.paragraph_format.keep_with_next = True
    add_text(cp, caption, size=9, italic=True, color=MUTED)
    add_source_note(doc, source)
    return True


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, 18, 10, BLUE),
        "Heading 2": (13, 14, 7, BLUE),
        "Heading 3": (12, 10, 5, DARK_BLUE),
    }
    for name, (size, before, after, color) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(header, "DiT Encoder-Decoder Deep Insight | 2026-08-21", size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, "Research companion | Page ", size=8.5, color=MUTED)
    add_page_number(footer)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def cover(doc):
    add_para(doc, after=38)
    p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_text(p, "TECHNICAL STUDY COMPANION", size=10, bold=True, color=RED)
    p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_text(p, "DiT 编解码器发展、架构、训练与数据工程", size=28, bold=True, color=NAVY)
    p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(p, "A Deep Insight into Encoders, Decoders and Visual Tokenizers for Diffusion Transformers", size=14, color=DARK_BLUE)
    add_para(doc, after=22)
    add_figure(
        doc,
        FIGURES / "explanatory" / "encoder_decoder_timeline.png",
        "图 0-1  DiT 编解码器技术演进总览 / Encoder-Decoder Evolution Overview",
        "解释图，依据本资料库架构矩阵与原始论文证据绘制。",
        max_width=6.3,
        max_height=3.1,
    )
    p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=4)
    add_text(p, "中文主叙述 · English terminology · Paper-native figures · Code-level anchors", size=10.5, color=MUTED)
    p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_text(p, "版本 / Version: 2026-08-21", size=10.5, bold=True, color=NAVY)
    p = add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_text(p, "研究资料库：daily/20260821_ENC_DEC", size=9.5, color=MUTED)
    page_break(doc)


def abstract_and_method(doc):
    heading(doc, "摘要 / Abstract", 1)
    add_para(
        doc,
        "编解码器在 DiT 系统中不是被动的输入输出接口，而是同时决定信息上限、token 数量、训练成本、时序边界和最终纹理风格。本报告从 VAE、VQ-VAE、VQGAN 和 LDM 的历史出发，系统比较 SD3、FLUX.1、TiTok、DC-AE，以及 MAGVIT、CogVideoX、Open-Sora、WF-VAE、VidTok、LTX-Video、Cosmos 与 Wan 等主流图像/视频编解码器；进一步把 tokenizer 预训练、冻结潜空间 DiT 训练和图像复原适配拆成三个不同优化问题，并以 FlashVSR 的 LQ_proj_in、WanDecoder 与 TCDecoder 为代码级案例。报告强调一个核心结论：压缩倍率、潜变量通道数、因果缓存和解码器监督共同构成生成模型的系统级设计空间，不能只以重建 PSNR 或单次推理速度孤立评价。",
    )
    add_para(
        doc,
        "Encoders and decoders are not passive I/O components in a Diffusion Transformer system. They define the information ceiling, token budget, optimization difficulty, temporal boundary and final texture prior. This report traces the evolution from VAE, VQ-VAE, VQGAN and latent diffusion to modern image and video tokenizers, separates tokenizer pretraining from frozen-latent DiT training and restoration adaptation, and provides a code-level FlashVSR case study. The central systems insight is that compression ratio, latent dimensionality, causality, cache semantics and decoder supervision must be designed jointly.",
    )
    callout(
        doc,
        "关键词 / Keywords",
        "Diffusion Transformer (DiT); Autoencoder; Variational Autoencoder (VAE); Visual Tokenizer; Causal Video VAE; Flow Matching; Perceptual Loss; FlashVSR; TCDecoder; Data Engineering.",
        fill=LIGHT_BLUE,
    )

    heading(doc, "证据口径与阅读方法", 2)
    add_table(
        doc,
        ["标签", "含义", "在本文中的使用方式"],
        [
            ["paper-verified", "原论文正文、附录或原文图明确给出", "可写作论文事实，并标注页码/图号"],
            ["code-verified", "官方仓库源码或配置明确给出", "可写作实现事实，并标注文件与行号"],
            ["primary-source-indexed", "权威原始来源已定位，但本地完整 PDF 尚未落盘", "仅做基础定义或来源索引，不声称本地全文核验"],
            ["analysis", "由多条证据归纳形成的工程解释", "明确写为分析、建议或推论"],
            ["undisclosed", "公开论文与源码未披露", "保持空白边界，不用其他模型配方代填"],
        ],
        [1.35, 2.45, 2.7],
        font_size=9.2,
    )
    callout(
        doc,
        "本地资料状态",
        "来源索引共 40 条，当前有效完整论文 PDF 为 7 篇，官方源码快照 16/16 有效；另外 11 篇较新论文具有本地 ar5iv 全文快照与原文图。VAE、VQ-VAE 等历史基础论文当前保留权威来源索引，完整 PDF 待补。",
        fill=LIGHT_GOLD,
        label_color=GOLD,
    )
    heading(doc, "阅读导航", 2)
    for item in [
        "先读第 1-2 章：建立 encoder、decoder、tokenizer、latent channel、compression factor 与 causality 的统一语义。",
        "再读第 3-4 章：分别理解图像与视频路线，重点观察 token 形状和 decoder 职责如何变化。",
        "第 5-6 章回答如何训练、LOSS 如何组合、数据如何制作和张量如何进入模型。",
        "第 7 章是 FlashVSR 深入案例，直接对应 LQ_proj_in、WanDecoder、TCDecoder 三个易混模块。",
        "第 8-9 章用于架构选型、复现检查和后续研究。Excel 是全量审计底稿，Word 是结构化学习手册。",
    ]:
        numbered(doc, item)


def foundations(doc):
    heading(doc, "1. 概念边界与历史演进", 1)
    heading(doc, "1.1 五个容易混淆的对象", 2)
    add_table(
        doc,
        ["对象", "英文", "核心职责", "不应被误读为"],
        [
            ["编码器", "Encoder", "把像素/帧映射到潜变量；VAE 编码器还预测后验参数", "任何把 RGB 投影到特征的条件分支"],
            ["瓶颈", "Bottleneck", "定义连续分布、离散码本或高压缩 token 形状", "单纯的 DiT patchify"],
            ["解码器", "Decoder", "把潜变量恢复到像素/帧，可附带条件、去噪或缓存", "一定与编码器对称的固定逆运算"],
            ["视觉分词器", "Visual Tokenizer", "encoder-bottleneck-decoder 的完整 token 契约", "只产生条件特征的 projector"],
            ["潜空间接口", "Latent Interface", "规定 DiT 输入的尺度、通道、归一化和解码方式", "一个可随意替换而不影响主干的前处理"],
        ],
        [1.05, 1.25, 2.35, 1.85],
        font_size=9,
    )
    add_para(
        doc,
        "判断一个模块是不是 encoder，最实用的问题不是看名字，而是看它是否定义了完整的潜变量统计或 token 契约：输入是什么，输出能否被标准 decoder 接受，是否有 posterior/quantizer，是否与同一 tokenizer 的训练目标绑定。FlashVSR 的 LQ_proj_in 虽然把低清 RGB 压到与 Wan latent 对齐的尺度，但它不产生 Wan posterior，因此应称为条件投影器而不是 Wan encoder。",
    )

    heading(doc, "1.2 VAE：连续概率潜空间", 2)
    add_para(
        doc,
        "VAE 的 encoder 学习近似后验 q_phi(z|x)，decoder 学习 p_theta(x|z)。重参数化 z = mu + sigma * epsilon 让随机采样仍可反向传播。它建立了现代连续 latent autoencoder 的概率基础，但原始 VAE 并没有规定今天常见的 f8、4 通道或 16 通道接口。[1]",
    )
    equation(doc, "L_VAE = E_q(z|x)[-log p_theta(x|z)] + beta * D_KL(q_phi(z|x) || p(z))", "重建项 + KL 正则")
    add_para(
        doc,
        "直觉上，重建项要求 z 保留输入信息；KL 项要求潜空间接近规则先验，便于生成模型采样。beta 过大时细节被压缩，过小时潜空间不规则。现代感知 VAE 通常加入 LPIPS 或 GAN，因此不能把原始 ELBO 等同于 SD/FLUX 类 autoencoder 的完整训练配方。",
    )

    heading(doc, "1.3 VQ-VAE：离散码本与 straight-through", 2)
    add_para(
        doc,
        "VQ-VAE 把连续 encoder 特征映射到最近的 codebook entry，生成离散索引。量化不可导，因此使用 straight-through estimator 把 decoder 梯度近似传回 encoder；codebook loss 和 commitment loss 分别更新词典并约束 encoder 不要漂离码本。[2]",
    )
    equation(doc, "k* = argmin_k || z_e(x) - e_k ||_2 ;  z_q(x) = e_k*", "nearest-neighbor vector quantization")
    equation(doc, "L_VQ = L_rec + ||sg[z_e]-e||_2^2 + beta ||z_e-sg[e]||_2^2", "codebook + commitment")
    callout(doc, "典型失败", "codebook collapse、dead codes 和过低的有效码本利用率会让名义 token 容量远大于实际容量。", fill=LIGHT_GOLD, label_color=GOLD)

    heading(doc, "1.4 VQGAN 与 LDM：从压缩到感知重建", 2)
    add_para(
        doc,
        "VQGAN 用感知损失和 patch discriminator 把重建目标从逐像素平均推进到局部纹理真实感；LDM 则把 autoencoder 与生成模型训练彻底分离：先训练 perceptual compression，再冻结 encoder/decoder，在低维 latent 上训练扩散模型。[4][5] 这一步奠定了今天 DiT 的主流系统边界。",
    )
    add_figure(
        doc,
        FIGURES / "paper_figures" / "crops" / "vqgan_architecture.png",
        "图 1-1  VQGAN 的感知式离散 tokenizer 架构。",
        "VQGAN 原论文第 3 页 Fig.2；本地 PDF: papers/01_history_and_tokenizers/vqgan_2012.09841.pdf。",
    )
    add_figure(
        doc,
        FIGURES / "paper_figures" / "crops" / "ldm_perceptual_compression.png",
        "图 1-2  LDM 将感知压缩与 latent diffusion 解耦。",
        "LDM 原论文第 2 页；本地 PDF: papers/01_history_and_tokenizers/ldm_2112.10752.pdf。",
    )
    callout(
        doc,
        "系统结论",
        "一旦 tokenizer 冻结，DiT 训练无法修复 decoder 的色偏、纹理涂抹或高频缺失。decoder 成为生成质量的硬上限，同时 encoder 决定 DiT 看到的任务难度。",
        fill=LIGHT_GREEN,
        label_color=GREEN,
    )


def interface_and_image_models(doc):
    heading(doc, "2. DiT 的潜空间接口与张量流", 1)
    add_para(
        doc,
        "DiT 并不直接规定 encoder/decoder。经典 DiT 使用预训练 Stable Diffusion VAE，把 RGB 图像压成 f8、4 通道 latent，再按 p x p latent patch 转为 Transformer tokens。[6] 因此总 token 数由两次压缩共同决定：VAE 的空间压缩 f 和 DiT 的 patch size p。",
    )
    equation(doc, "N_tokens(image) = (H / f / p) * (W / f / p)", "图像 token 数")
    equation(doc, "N_tokens(video) = ceil(T / f_t) * (H / f_s / p_h) * (W / f_s / p_w)", "视频 token 数")
    add_figure(
        doc,
        FIGURES / "explanatory" / "image_video_tensor_flow.png",
        "图 2-1  图像/视频 tokenizer 与 DiT 的张量流。",
        "解释图；数值来自 DiT、SD3、CogVideoX、LTX-Video、Wan 与 FlashVSR 的公开结构。",
        max_height=4.4,
    )
    add_figure(
        doc,
        FIGURES / "paper_figures" / "crops" / "dit_architecture.png",
        "图 2-2  DiT 原论文：latent patchify、Transformer blocks 与输出头。",
        "DiT 原论文第 3 页 Fig.3；source_code/dit/models.py。",
    )
    callout(
        doc,
        "形状示例",
        "512x512 RGB 经 f8 VAE 得到 [B,4,64,64]。若 DiT patch size p=2，则 token grid 为 32x32，共 1024 tokens；若改成 16 通道 latent，token 数不变，但每个位置的信息维度和预测难度显著增加。",
        fill=LIGHT_BLUE,
    )

    heading(doc, "3. 主流图像编解码器", 1)
    heading(doc, "3.1 SD3 VAE 与 MM-DiT：16 通道 f8 接口", 2)
    add_para(
        doc,
        "SD3 将图像 latent 从经典 4 通道扩展到 16 通道，同时保持 f8 空间压缩。官方参考实现使用四级 ch_mult=(1,2,4,4)、两层 residual blocks 和 mid attention，encoder 输出 2*z_channels 以形成后验参数；latent 还需经过 scale=1.5305、shift=0.0609 的归一化接口。[8] 更高通道容量改善重建，但也扩大 flow model 的输入/输出维度。",
    )
    add_table(
        doc,
        ["结构锚点", "SD3 实现", "对 MM-DiT 的影响"],
        [
            ["Encoder backbone", "2D ResNet；ch_mult=(1,2,4,4)；每级 2 个 residual blocks", "保持 f8 网格，同时扩大单个 latent 位置的表达容量"],
            ["Bottleneck", "mid-block self-attention", "在进入 flow model 前补充低分辨率全局信息交互"],
            ["Posterior", "encoder 输出 2 x 16 channels，拆分为 mean / log-variance", "采样后以 16-channel continuous latent 作为图像 token 来源"],
            ["Latent contract", "scale=1.5305；shift=0.0609", "MM-DiT 权重依赖该统计接口，不可与其他 VAE latent 直接互换"],
        ],
        [1.15, 2.75, 2.6],
        font_size=8.5,
    )
    add_figure(
        doc,
        FIGURES / "paper_figures" / "crops" / "sd3_mmdit_architecture.png",
        "图 3-1  SD3 的 MM-DiT 双模态联合注意力结构。",
        "SD3 原论文第 5 页；本地 PDF: papers/02_image_vae/sd3_2403.03206.pdf。",
        max_height=6.2,
    )
    add_figure(
        doc,
        FIGURES / "paper_figures" / "crops" / "sd3_autoencoder_table.png",
        "图 3-2  SD3 原论文的 autoencoder 架构比较与 16 通道选择。",
        "SD3 原论文第 7 页 Table 3；source_code/sd3/sd3_impls.py:97-107,247-310。",
    )

    heading(doc, "3.2 FLUX.1 AE：公开结构与未公开训练配方", 2)
    add_para(
        doc,
        "FLUX.1 官方源码公开了与 SD3 同属 2D ResNet autoencoder 家族的结构：in/out RGB=3，base channel=128，ch_mult=[1,2,4,4]，两层 residual blocks，z_channels=16，四个 resolution levels 对应 f8；latent 归一化为 z_norm = 0.3611 * (z - 0.1159)。源码可核验 encoder、decoder 与归一化，但没有公开完整 AE 训练数据、损失权重、判别器策略与优化日程。",
    )
    equation(doc, "z_model = 0.3611 * (z_ae - 0.1159) ;  z_ae = z_model / 0.3611 + 0.1159", "FLUX.1 latent normalization")
    add_table(
        doc,
        ["维度", "SD3", "FLUX.1", "工程含义"],
        [
            ["空间压缩", "f8", "f8（由源码层级推得）", "token 网格规模相近"],
            ["latent channels", "16", "16", "比经典 SD-VAE 的 4ch 更高容量"],
            ["归一化", "scale 1.5305 / shift 0.0609", "scale 0.3611 / shift 0.1159", "权重与 latent 不可交叉混用"],
            ["AE 训练配方", "论文未完整披露", "未公开", "只能复用架构，不能声称复现官方训练"],
        ],
        [1.25, 1.55, 1.65, 2.05],
        font_size=9,
    )
    add_source_note(doc, "source_code/flux/modules/autoencoder.py；source_code/flux/util.py:539-548；证据状态 code-verified+undisclosed。")
    callout(
        doc,
        "关键边界",
        "把 LDM 或 SD3 的 LPIPS/GAN/KL 权重复制给 FLUX autoencoder，可以作为复现实验起点，但必须标注为 engineering baseline，不能写成 FLUX 官方事实。",
        fill=LIGHT_GOLD,
        label_color=GOLD,
    )

    heading(doc, "3.3 TiTok：从二维网格到一维少量 token", 2)
    add_para(
        doc,
        "TiTok 不再把 bottleneck 固定成二维格点，而是让 ViT encoder 同时处理图像 patches 与 K 个可学习 latent tokens，再量化为 K 个一维 tokens。旗舰设置在 256px 使用 32 tokens、codebook 4096。它主要服务 MaskGIT 路线，但对 DiT 的启示非常直接：token 预算不一定只能依靠更大空间压缩获得，也可以改变 token 拓扑。[9]",
    )
    add_figure(
        doc,
        FIGURES / "paper_figures" / "titok_framework.png",
        "图 3-3  TiTok：K 个一维 latent tokens 的编码与解码框架。",
        "TiTok 原文 ar5iv 快照 Fig.3；text/ar5iv_2406.07550.txt。",
        max_height=5.2,
    )
    add_para(
        doc,
        "训练采用两阶段策略：先预测 MaskGIT-VQGAN 的 proxy codes，降低直接 RGB/GAN 优化难度；再冻结 encoder 与 quantizer，仅 fine-tune decoder，使用 VQGAN 风格的像素、感知与对抗目标。该方法说明 decoder 可以被视为独立的高频重建器进行后训练。",
    )

    heading(doc, "3.4 DC-AE：高压缩不是多堆几层 stride", 2)
    add_para(
        doc,
        "DC-AE 把空间压缩推进到 f32/f64/f128，代表一种明显的系统转向：把 token 压缩从 DiT patchify 前移到 autoencoder。本质难点不是网络能否 downsample，而是如何在高压缩下保持全局结构、局部细节和可优化性。DC-AE 使用 residual autoencoding、space-to-channel shortcut，以及低分辨率预训练、高分辨率 latent adaptation、局部 GAN refinement 的分阶段方案。[10]",
    )
    add_figure(
        doc,
        FIGURES / "paper_figures" / "dcae_architecture.png",
        "图 3-4  DC-AE 的 residual autoencoding 与高压缩结构。",
        "DC-AE 原文 ar5iv 资产；text/ar5iv_2410.10733.txt。",
        max_height=5.7,
    )
    add_figure(
        doc,
        FIGURES / "paper_figures" / "dcae_compression_results.png",
        "图 3-5  DC-AE 展示不同压缩等级对重建与生成效率的影响。",
        "DC-AE 原文 Fig.1 资产。",
        max_height=2.0,
    )
    callout(
        doc,
        "工程判断",
        "高压缩 VAE 适合高分辨率生成和长视频，但不天然适合严格复原。压缩越激进，decoder 越可能依靠生成先验补纹理，忠实度风险越高。",
        fill=LIGHT_GREEN,
        label_color=GREEN,
    )


def video_models(doc):
    heading(doc, "4. 主流视频编解码器", 1)
    add_para(
        doc,
        "视频 tokenizer 必须额外处理时间压缩、第一帧策略、因果卷积、长视频缓存和帧间一致性。名义上相同的 4x8x8 压缩，如果 first-frame policy、normalization 或 cache contract 不同，也不能互换。",
    )
    add_table(
        doc,
        ["模型", "类型", "压缩", "关键特征"],
        [
            ["MAGVIT / v2", "离散", "配置相关", "3D token grid；v2 使用 causal conv 与 LFQ"],
            ["CogVideoX", "连续 causal VAE", "4x8x8", "context parallelism；Frame Pack；多通道配置"],
            ["Open-Sora 1.2", "2D+3D hybrid", "4x8x8", "冻结 SDXL VAE；逐步从 feature 到 RGB 目标"],
            ["WF-VAE", "连续 causal VAE", "常见视频倍率", "wavelet energy flow；causal cache"],
            ["LTX-Video", "denoising VAE", "8x32x32", "decoder 兼任最后去噪步骤"],
            ["Cosmos", "连续+离散", "4/8x8x8；8x16x16", "统一图像视频；可选 diffusion decoder"],
            ["Wan", "连续 causal VAE", "4x8x8, 16ch", "首帧+4 帧块；feature cache；FlashVSR 基础"],
        ],
        [1.3, 1.45, 1.2, 2.55],
        font_size=8.8,
    )

    heading(doc, "4.1 MAGVIT 与 MAGVIT-v2：离散视频 token 的祖先路线", 2)
    add_para(
        doc,
        "MAGVIT 通过 3D VQ tokenizer 统一图像与视频任务，为后续视频 token 建模提供重要基线。[13] MAGVIT-v2 进一步采用 temporally causal convolution 和 lookup-free quantization (LFQ)，用二值因子化方式扩展码本，减少传统 VQ lookup 的瓶颈。[14] 即使现代视频 DiT 多用连续 VAE，这条路线仍深刻影响 causal convolution、first-frame handling 和 tokenizer benchmark。",
    )
    add_figure(doc, FIGURES / "paper_figures" / "crops" / "magvit_pipeline.png", "图 4-1  MAGVIT 的统一视频 tokenizer 与生成管线。", "MAGVIT 原论文第 3 页 Fig.2；本地 PDF: papers/03_video_vae/magvit_2212.05199.pdf。")
    add_figure(doc, FIGURES / "paper_figures" / "magvit2_architecture.png", "图 4-2  MAGVIT-v2 的 causal tokenizer 与 LFQ 结构。", "MAGVIT-v2 原文 ar5iv 资产；text/ar5iv_2310.05737.txt。", max_height=5.6)

    heading(doc, "4.2 CogVideoX 3D VAE：生产级 4x8x8 causal latent", 2)
    add_para(
        doc,
        "CogVideoX 采用 temporally causal 3D convolutional VAE，把时间、空间分别压缩 4、8、8 倍。论文同时研究 4/8/16/32 latent channels，提醒复现者不能只凭模型家族名假设所有 checkpoint 的通道数一致。官方实现还引入 context parallelism，解决大视频编码解码的显存瓶颈。[16]",
    )
    add_figure(doc, FIGURES / "paper_figures" / "cogvideox_3dvae.jpg", "图 4-3  CogVideoX 3D causal VAE 与压缩路径。", "CogVideoX 原文 Fig.4；source_code/cogvideox/sat/vae_modules/cp_enc_dec.py:743-998。", max_height=3.4)
    add_figure(doc, FIGURES / "paper_figures" / "cogvideox_framepack.jpg", "图 4-4  CogVideoX Frame Pack：混合时长训练的 batch 组织方式。", "CogVideoX 原文图；text/ar5iv_2408.06072.txt。", max_height=2.5)
    add_para(
        doc,
        "数据侧，CogVideoX 把 caption density 当作模型质量杠杆，使用多阶段质量与 caption pipeline，并在末期使用高质量 cooling subset。VAE 的 self-reconstruction 数据与 DiT 的 video-caption 数据相关但目标不同，应分别管理。",
    )

    heading(doc, "4.3 Open-Sora 1.2：用冻结图像 VAE 启动视频 VAE", 2)
    add_para(
        doc,
        "Open-Sora 1.2 先冻结 SDXL 2D VAE，再叠加 MAGVIT-v2 风格的 3D VAE。训练不是一次性从 RGB 到 RGB，而是三阶段改变 target space：先重建 2D-VAE features 并做 identity alignment，再强化 feature reconstruction，最后直接重建 RGB 视频。其公开日程为 0-380K、380K-640K、640K-1.2M steps。[19]",
    )
    add_figure(doc, FIGURES / "paper_figures" / "opensora_3dvae.png", "图 4-5  Open-Sora 1.2 的 2D VAE + 3D VAE 混合结构。", "Open-Sora 原文 Fig.5 的 ar5iv SVG 渲染；text/ar5iv_2412.20404.txt。", max_height=4.9)
    callout(doc, "训练启示", "当视频 VAE 难以从头稳定训练时，可以先对齐成熟图像 VAE 的 feature space，再逐步切换到直接 RGB 视频监督。", fill=LIGHT_GREEN, label_color=GREEN)

    heading(doc, "4.4 WF-VAE：小波能量流与 block-wise 解码", 2)
    add_para(
        doc,
        "WF-VAE 在 causal 3D backbone 之外加入 multi-level wavelet energy-flow path，让低频结构和高频细节以显式能量流进入编码/解码；同时用 causal cache 支持分块推理。论文消融中 wavelet loss 权重 lambda_WL=0.1，且 L1 在其设置下优于 L2。[17]",
    )
    add_figure(doc, FIGURES / "paper_figures" / "crops" / "wfvae_architecture.png", "图 4-6  WF-VAE 的 wavelet-driven energy flow 架构。", "WF-VAE 原论文第 3 页 Fig.2；本地 PDF: papers/03_video_vae/wfvae_2411.17459.pdf。")
    add_figure(doc, FIGURES / "paper_figures" / "crops" / "wfvae_reconstruction.png", "图 4-7  WF-VAE 原文重建效果对比。", "WF-VAE 原论文第 6 页；同上本地 PDF。", max_height=5.9)

    heading(doc, "4.5 VidTok：把训练变量变成可控实验", 2)
    add_para(
        doc,
        "VidTok 提供 continuous KL 与 discrete FSQ 等多种 bottleneck，用统一协议研究架构、量化、帧率和分辨率。其重要结论之一是：先低分辨率训练完整 tokenizer，再只 fine-tune 高分辨率 decoder，可以显著降低高分辨率训练成本；较低 FPS 不是简单减少 I/O，而是让模型看到更大的帧间变化，改善 motion representation。[18]",
    )
    add_figure(doc, FIGURES / "paper_figures" / "vidtok_overview.png", "图 4-8  VidTok 的模块化视频 tokenizer 研究框架。", "VidTok 原文 Fig.2；text/ar5iv_2412.13061.txt。", max_height=2.3)

    heading(doc, "4.6 LTX-Video：decoder 同时承担最后一步去噪", 2)
    add_para(
        doc,
        "LTX-Video 把压缩提高到时间 8 倍、空间 32 倍，总 pixel-to-latent ratio 约 1:192。它把 patchification 前移进 VAE，并让 timestep-conditioned decoder 接收多层噪声注入，在 latent-to-pixel 重建时完成最后一步去噪。[21] 这打破了经典边界：decoder 不再只是固定 inverse transform，而成为生成轨迹的一部分。",
    )
    add_figure(doc, FIGURES / "paper_figures" / "ltx_vae_encoder.png", "图 4-9  LTX-Video VAE encoder 的高压缩路径。", "LTX-Video 原文 Fig.2 资产；text/ar5iv_2501.00103.txt。", max_height=5.8)
    add_figure(doc, FIGURES / "paper_figures" / "ltx_denoising.png", "图 4-10  LTX-Video 将最后一步 denoising 合入 decoder。", "LTX-Video 原文 Fig.4 资产。", max_height=2.4)
    callout(doc, "适用边界", "这种 decoder 适合端到端实时生成，但与需要严格复用独立 VAE latent contract 的系统兼容性更弱。", fill=LIGHT_GOLD, label_color=GOLD)

    heading(doc, "4.7 Cosmos：连续、离散与 diffusion decoder 共存", 2)
    add_para(
        doc,
        "Cosmos tokenizer family 同时支持 continuous latent diffusion 与 discrete autoregressive modeling。其 encoder 使用 3D Haar wavelet、causal residual/downsampling 与 spatiotemporal attention；离散 token 还可通过 diffusion decoder 恢复到连续表征再生成视频。[22] 这说明未来 tokenizer 可能不再绑定单一生成范式。",
    )
    add_figure(doc, FIGURES / "paper_figures" / "cosmos_tokenizer_architecture.png", "图 4-11  Cosmos tokenizer 的因果时空架构。", "Cosmos 原文网络架构图；text/ar5iv_2501.03575.txt。", max_height=4.2)
    add_figure(doc, FIGURES / "paper_figures" / "cosmos_diffusion_decoder.png", "图 4-12  Cosmos 的 discrete-to-continuous diffusion decoder。", "Cosmos 原文 diffusion decoder 图。", max_height=2.5)

    heading(doc, "4.8 Wan VAE：FlashVSR 所依赖的完整 causal contract", 2)
    add_para(
        doc,
        "Wan VAE 使用 16-channel continuous causal latent，时间与空间压缩为 4x8x8。官方源码按首帧 1 帧、后续每 4 帧处理，并在 encoder/decoder 中维护 feature cache；mean/std 为 16 通道向量。FlashVSR 的 WanDecoder 指的就是这条完整 decoder path。[23]",
    )
    add_figure(doc, FIGURES / "paper_figures" / "wan_vae_architecture.png", "图 4-13  Wan VAE 的 4x8x8 causal video latent 架构。", "Wan 原文 Fig.5 资产；source_code/wan/wan/modules/vae.py:483-660。", max_height=2.3)
    add_para(
        doc,
        "实现层面，WanVAE_ 同时实例化 Encoder3d 与 Decoder3d；encode 将视频拆成 1,4,4,... 帧块，decode 按 latent time step 逐块恢复，并在每次调用前后 clear_cache。这个 cache contract 是流式正确性的组成部分，不能只复制卷积权重而忽略。",
    )


def training_and_data(doc):
    heading(doc, "5. 编解码器与 DiT 的训练方法", 1)
    heading(doc, "5.1 三类训练问题必须分开", 2)
    add_table(
        doc,
        ["训练问题", "输入/目标", "更新参数", "核心风险"],
        [
            ["Tokenizer/VAE 预训练", "RGB/Video -> same RGB/Video", "encoder + bottleneck + decoder，可能含 discriminator", "压缩失真、GAN 幻觉、码本坍塌、时序闪烁"],
            ["冻结 latent 的 DiT 训练", "latent + condition -> noise/velocity", "DiT/Flow Transformer", "tokenizer 缺陷无法被主干修复"],
            ["复原任务适配", "LQ observation + noisy latent -> HR latent/pixel", "adapter/LoRA/condition projector/decoder", "忠实度与感知生成冲突"],
        ],
        [1.45, 1.75, 1.65, 1.65],
        font_size=9,
    )

    heading(doc, "5.2 Tokenizer/VAE 预训练的 LOSS 组合", 2)
    equation(doc, "L_tokenizer = lambda_rec L_rec + lambda_perc L_LPIPS + lambda_reg L_KL/VQ + lambda_adv L_GAN + lambda_temp L_temp + lambda_freq L_freq", "通用分析式，不代表任一模型官方权重")
    for text in [
        "L_rec：L1 通常比 L2 更少鼓励平均化；Charbonnier 可在异常值下更稳。严格复原优先保证这一项。",
        "L_LPIPS：改善语义和局部纹理，但会允许像素偏移。对人脸、文字、细线结构应配合 ROI 或 edge 约束。",
        "L_KL / L_VQ：决定 latent 规则性与可用容量。KL 太强导致信息丢失，VQ commitment 太弱导致 encoder 漂移。",
        "L_GAN：改善高频统计，但最容易引入伪纹理。通常在 reconstruction 稳定后延迟启用，或仅在高分辨率局部 patch 使用。",
        "L_temp：可由 optical-flow warp consistency、feature consistency、first/last-frame consistency 等组成；必须避免错误光流把运动边界抹平。",
        "L_freq：小波、FFT 或高频 band loss 可补充边缘，但不能替代像素与感知目标。WF-VAE 的 wavelet loss 是代表案例。",
    ]:
        bullet(doc, text)
    callout(
        doc,
        "推荐起点，不是定论",
        "面向真实复原的连续 VAE 可先以 L1/Charbonnier 为主，加入较小 LPIPS 和 KL；GAN 延迟启用并从小权重开始。所有权重都应通过 gradient-norm logging、重建/感知双指标和消融实验校准，不能只复制论文常数。",
        fill=LIGHT_GOLD,
        label_color=GOLD,
    )

    heading(doc, "5.3 冻结 latent 的 DiT / Flow Matching", 2)
    add_para(
        doc,
        "经典 DiT 在冻结 VAE latent 上做 diffusion noise prediction；SD3/FLUX 路线使用 rectified flow / flow matching，学习从噪声到数据的 velocity field。encoder 和 decoder 通常不接收梯度，latent 可以预编码缓存，显著减少训练开销。[6][8]",
    )
    equation(doc, "x_t = (1-t) x_0 + t x_1 ;  v_target = x_1 - x_0 ;  L_FM = E ||v_theta(x_t,t,c) - v_target||_2^2", "线性概率路径上的 flow matching 简化式")
    add_para(
        doc,
        "预编码的好处是吞吐高、训练稳定；代价是 tokenizer 固定成为质量天花板。更换 VAE 后，即使空间尺寸相同，只要 latent channel、shift/scale 或统计分布不同，DiT 权重也不能直接复用。",
    )

    heading(doc, "5.4 高压缩和视频 VAE 为什么需要 curriculum", 2)
    add_table(
        doc,
        ["策略", "代表模型", "优化目的"],
        [
            ["proxy code -> RGB", "TiTok", "先学 token 语义，再学高频像素"],
            ["低分辨率全模型 -> 高分辨率 decoder-only", "VidTok", "把高分辨率算力集中到解码器"],
            ["低分辨率 -> HR latent adaptation -> local GAN", "DC-AE", "稳定 f64/f128 高压缩训练"],
            ["2D feature target -> RGB video target", "Open-Sora", "从成熟图像表征逐步过渡到视频像素"],
            ["decoder 融合最后去噪", "LTX-Video", "让 tokenizer 与最终生成轨迹共同优化"],
        ],
        [2.05, 1.65, 2.8],
        font_size=9.1,
    )

    heading(doc, "6. 数据集制作与模型输入", 1)
    heading(doc, "6.1 图像 tokenizer 数据", 2)
    add_para(
        doc,
        "图像 VAE 的目标是覆盖将来生成或复原分布的视觉统计，而不是盲目追求海量 caption。基础流程包括：来源合法性检查、去重、质量筛选、尺寸和长宽比分桶、颜色空间统一、动态范围与 alpha 处理、随机 crop/resize，以及固定归一化到模型范围。生成数据集与 tokenizer 数据集可以不同，但 tokenizer 数据分布如果缺少文字、小脸、夜景或细纹理，decoder 会形成系统性盲区。",
    )
    add_table(
        doc,
        ["阶段", "处理", "必须记录的元数据"],
        [
            ["入库", "来源、许可、SHA/感知去重", "source_id、license、resolution、color profile"],
            ["质量", "美学、清晰度、压缩、NSFW、异常图", "quality scores、filter reason"],
            ["采样", "resolution/aspect buckets、global/local crop", "crop box、resize kernel、bucket id"],
            ["输入", "RGB range、dtype、VAE encode", "normalization、latent scale/shift、encoder version"],
        ],
        [1.0, 2.7, 2.8],
        font_size=9,
    )

    heading(doc, "6.2 视频 tokenizer 数据", 2)
    add_para(
        doc,
        "视频数据需要在图像清洗之外管理 shot boundary、帧率、运动量、duration、首帧策略和解码一致性。低 FPS 可以提高相邻训练帧的运动幅度，帮助模型学习 motion；但过低会产生非连续跳变。长视频通常应先切 shot，再按帧数或 token budget 分桶，并记录原始 FPS 与采样 FPS。",
    )
    for text in [
        "统一张量约定：推荐在数据层固定 [B,3,T,H,W]，模型内部若使用 [B,T,C,H,W] 必须显式转换并写单元测试。",
        "first-frame policy：首帧是否单独编码、是否重复 padding、后续按 4/8 帧块处理，必须进入数据与 cache contract。",
        "mixed-duration batching：按 token 数而非 clip 数控制 batch，避免长视频样本挤爆显存。CogVideoX Frame Pack 是代表方案。",
        "时序清洗：检测重复帧、冻结帧、强转场、字幕条、极端抖动和可变帧率；这些问题会直接污染 causal decoder。",
    ]:
        bullet(doc, text)
    add_figure(doc, FIGURES / "paper_figures" / "cogvideox_caption_pipeline.jpg", "图 6-1  CogVideoX 的质量与 caption 数据处理管线。", "CogVideoX 原文 Fig.6-7 相关资产。", max_height=2.5)

    heading(doc, "6.3 超分/复原配对数据", 2)
    add_para(
        doc,
        "复原任务不能只靠 self-reconstruction。需要从 HQ/GT 生成或采集 LQ observation，并保留退化参数、对齐质量和任务 mask。真实 SR 的退化链应覆盖光学、传感器、ISP、resize、压缩和二次处理，而不是只有 bicubic。FlashVSR 使用 RealBasicVSR degradation 从 HQ 构造 LR-HR 配对。[24]",
    )
    equation(doc, "x_LQ = JPEG(Sharpen(Denoise(Resize(Blur(x_HQ)) + Noise)))", "示意退化链；实际顺序需随机化并记录参数")
    add_table(
        doc,
        ["数据类型", "输入", "目标", "关键风险"],
        [
            ["合成 paired SR", "退化后的 LR/LQ", "原始 HQ", "退化分布过窄，模型只会识别脚本痕迹"],
            ["真实 paired capture", "多焦段/多曝光/多设备 LQ", "高质量参考帧", "亚像素、rolling shutter、曝光和 ISP 不一致"],
            ["unpaired real LQ", "真实低质图像/视频", "无严格 GT", "需 teacher、cycle、GAN 或无参考奖励，幻觉风险高"],
            ["mask-aware restoration", "LQ + face/text/edge/region mask", "区域化 HQ/identity target", "小区域梯度被全图平均，需要面积归一化"],
        ],
        [1.45, 1.75, 1.7, 1.6],
        font_size=8.8,
    )
    callout(
        doc,
        "数据审计",
        "每个训练 sample 应至少能追溯 HQ source、LQ degradation seed/parameters、crop coordinates、frame indices、FPS、mask version、alignment score 和 tokenizer version。没有这些字段，后续失败样本无法定位到数据还是模型。",
        fill=LIGHT_GREEN,
        label_color=GREEN,
    )

    heading(doc, "6.4 可直接落地的输入检查", 2)
    for text in [
        "检查 H、W 是否能被 VAE spatial factor 与 DiT patch size 整除；视频还要检查 T 与 first-frame/temporal factor 的兼容关系。",
        "把 latent mean/std、scale/shift、dtype、range 写入 checkpoint metadata；禁止靠代码默认值猜测。",
        "encode-decode round trip 单独评测：PSNR/SSIM/LPIPS + face/text/edge 子集 + 视频 flicker/warp 指标。",
        "训练前可视化 64 个 batch：原图、crop、LQ、mask、VAE reconstruction、latent channel statistics。",
        "对 causal 模型做 prefix invariance test：截断未来帧后，过去输出应保持一致；否则 padding/normalization/cache 可能泄漏未来信息。",
    ]:
        numbered(doc, text)


def flashvsr_case(doc):
    heading(doc, "7. FlashVSR 深入案例：LQ_proj_in、WanDecoder 与 TCDecoder", 1)
    add_para(
        doc,
        "FlashVSR 的关键贡献不是简单给 Wan 增加 SR condition，而是把实时流式 VSR 拆成三个可独立优化的系统瓶颈：条件输入、生成主干和像素解码。[24] 三个模块名称接近编解码概念，但边界完全不同。",
    )
    add_figure(doc, FIGURES / "explanatory" / "flashvsr_module_boundary.png", "图 7-1  FlashVSR 三类模块的边界与数据流。", "解释图；依据 FlashVSR 原论文、utils.py、TCDecoder.py 和 Wan VAE 源码绘制。", max_height=5.8)
    add_table(
        doc,
        ["模块", "真实类别", "输入 -> 输出", "推理职责"],
        [[r["module"], r["category"], f"{r['input']} -> {r['output']}", r["inference_role"]] for r in FLASH],
        [1.15, 1.35, 2.3, 1.7],
        font_size=8.6,
    )

    heading(doc, "7.1 VSR-120K 与三阶段训练", 2)
    add_para(
        doc,
        "FlashVSR 从 Videvo、Pexels、Pixabay 收集初始 600K 视频和 220K 图像，使用 LAION-Aesthetic、MUSIQ 和 RAFT 做质量与运动过滤，最终得到 120K 视频（平均超过 350 帧）和 180K HQ 图像。LR-HR 配对由 RealBasicVSR degradation 合成。Stage 1 把图像视为 T=1 视频联合训练；Stage 2 转为 block-sparse causal attention；Stage 3 用 DMD 压到一步生成。",
    )
    add_figure(doc, FIGURES / "paper_figures" / "crops" / "flashvsr_training_pipeline.png", "图 7-2  FlashVSR 的数据、条件投影与三阶段训练流程。", "FlashVSR 原论文第 4 页 Fig.2；本地 PDF: papers/05_flashvsr_case/flashvsr_2510.12747.pdf。")
    equation(doc, "L_stage3 = L_DMD + lambda_FM L_FM + lambda_MSE L_MSE + 2 * L_LPIPS", "论文公开的 Stage 3 组合；其他系数按原文/实现核验")
    add_para(
        doc,
        "实现设置包括 LoRA rank 384、AdamW、learning rate 1e-5、weight decay 0.01、batch size 32。这里训练的是 restoration DiT/adapter，不是从头训练 Wan VAE。",
    )

    heading(doc, "7.2 LQ_proj_in：条件投影，不是完整 encoder", 2)
    add_para(
        doc,
        "源码中的 Buffer_LQ4x_Proj 先用 PixelShuffle3d-style rearrangement 把 16x16 空间块搬到 channel，再用两个 causal Conv3d 将时间连续下采样两次，并通过 RMSNorm、SiLU 与逐层 Linear 映射到 DiT feature space。forward 还复制首帧并按 4 帧块处理，以对齐 Wan causal latent。",
    )
    equation(doc, "[B,3,T,H,W] -> rearrange(1,16,16) -> causal Conv3d x2 -> per-block DiT condition features", "source_code/flashvsr/utils.py:54-105")
    add_para(
        doc,
        "它替代的是 LR condition 的 VAE encoding 路径，而不是生成一个可由标准 Wan decoder 独立解码的 posterior。它没有 mu/logvar、reparameterization，也不是一个完整 tokenizer。训练监督来自 restoration flow objectives 间接传递。",
    )
    callout(doc, "常见误读", "把 LQ_proj_in 称为 LQ encoder 会掩盖它与 Wan latent posterior 的统计差异。更准确的名称是 latent-aligned condition projector。", fill=LIGHT_GOLD, label_color=GOLD)

    heading(doc, "7.3 WanDecoder：完整、高质量但昂贵的 causal VAE decoder", 2)
    add_para(
        doc,
        "WanDecoder 接收 16-channel Wan latent，经 conv2、Decoder3d、时空上采样与 feature cache 恢复 RGB 视频。官方源码的 decode 按 latent time step 迭代，并在每一步复用 causal feature cache。FlashVSR 论文分析中，完整 Wan decoder 在 768x1408 设置下约占 70% 推理时间，因此即使 DiT 已经一步化，decoder 仍是系统瓶颈。",
    )
    add_source_note(doc, "source_code/wan/wan/modules/vae.py:483-660；source_code/flashvsr/wan_video_vae.py；FlashVSR p5 Sec.3.4。")

    heading(doc, "7.4 TCDecoder：decoder-only 条件蒸馏器", 2)
    add_para(
        doc,
        "TCDecoder 在源码中是 TAEHV decoder-only：输入 predicted Wan latent，使用 MemBlock 保存前帧状态，经过三次空间上采样、两次时间增长，最后输出 RGB；PixelShuffle3d(4,8,8) 负责把 Wan 压缩尺度重排到 decoder 可消费的时空结构。它还接收 LR conditional features，以减少小 decoder 完全依赖 latent 猜测细节的压力。",
    )
    add_table(
        doc,
        ["源码模块", "锚点", "作用与训练含义"],
        [
            ["PixelShuffle3d", "TCDecoder.py:73-89", "按 (4,8,8) 重排时空压缩维度，建立 Wan latent 到轻量 decoder 的输入接口"],
            ["TAEHV", "TCDecoder.py:170-213", "decoder-only 主体；组合条件特征、残差块与多级上采样并输出 RGB"],
            ["MemBlock", "causal state", "保存前帧特征状态，维持流式解码中的因果一致性与跨帧连续性"],
            ["TGrow", "temporal upsampling", "逐级增长时间维，恢复 Wan 4 倍时间压缩对应的帧序列"],
            ["Dual supervision", "GT + Wan teacher", "对 GT 和 WanDecoder output 分别计算 MSE + LPIPS，兼顾真实目标与 latent-compatible teacher prior"],
        ],
        [1.25, 1.45, 3.8],
        font_size=8.2,
    )
    add_figure(doc, FIGURES / "paper_figures" / "crops" / "flashvsr_tcdecoder_pipeline.png", "图 7-3  FlashVSR 原论文的 TCDecoder 条件解码与蒸馏路径。", "FlashVSR 原论文第 5 页 Fig.4；source_code/flashvsr/TCDecoder.py:73-213。", max_height=4.7)
    equation(doc, "L_TC = (L_MSE(y_tiny,y_GT)+L_LPIPS(y_tiny,y_GT)) + (L_MSE(y_tiny,y_Wan)+L_LPIPS(y_tiny,y_Wan))", "同时监督 GT 与 WanDecoder teacher")
    add_para(
        doc,
        "双目标设计的意义是互补：GT 监督纠正 teacher 的系统偏差，Wan decoder output 则提供更平滑、与 latent contract 一致的蒸馏目标。论文在 61-frame 384x384 clips 上单独训练约两天，并报告 decoder 接近 7 倍加速。这个速度是 decoder 子系统的论文结果，不应误写成整个端到端系统 7 倍。",
    )

    heading(doc, "7.5 Full 与 Tiny 的质量-速度边界", 2)
    add_figure(doc, FIGURES / "paper_figures" / "crops" / "flashvsr_qualitative_results.png", "图 7-4  FlashVSR 原文定性结果：流式超分的细节与时序表现。", "FlashVSR 原论文第 7 页定性对比。", max_height=5.8)
    add_para(
        doc,
        "Full path 保留完整 Wan decoder，适合离线高质量；Tiny path 用 TCDecoder 换取实时性。产品上应把两者视为同一 latent generator 的两个 pixel renderer，而不是两个独立 SR 模型。切换 decoder 时要分别评测：纹理真实感、文字/人脸忠实度、运动边界、首帧冷启动、cache reset 和长视频漂移。",
    )
    callout(
        doc,
        "对 FlashVSR 的系统洞察",
        "一步 DiT 并不等于实时系统。真正的端到端 latency 由 condition projection、attention、latent update、decoder、I/O 和 cache 管理共同决定；FlashVSR 的 TCDecoder 证明 decoder-only 替换可以成为独立的加速方向。",
        fill=LIGHT_GREEN,
        label_color=GREEN,
    )


def synthesis(doc):
    heading(doc, "8. 架构选型与复现建议", 1)
    heading(doc, "8.1 按任务选择 latent contract", 2)
    add_table(
        doc,
        ["任务", "优先选择", "谨慎选择", "理由"],
        [
            ["高保真图像复原", "f8 continuous VAE；4-16ch；弱 GAN", "f64/f128 高压缩", "需要保留文字、人脸和细线结构"],
            ["高分辨率图像生成", "SD3/FLUX 16ch f8；DC-AE f32/f64", "极低容量离散 token", "平衡 token 数与 decoder 细节"],
            ["长视频生成", "causal 3D VAE + first-frame/cache contract", "非因果 3D tokenizer", "流式、分块和长时一致性"],
            ["实时视频超分", "Wan latent + 条件 projector + tiny conditional decoder", "完整 VAE encoder 处理 LR", "条件输入和 pixel renderer 都需轻量化"],
            ["端侧部署", "低通道/低 token、decoder-only distillation、量化友好卷积", "复杂 attention decoder", "内存、启动和功耗比单次 FLOPs 更关键"],
        ],
        [1.35, 2.0, 1.55, 1.6],
        font_size=8.8,
    )

    heading(doc, "8.2 评测不应只看重建 PSNR", 2)
    for text in [
        "Tokenizer 层：PSNR/SSIM、LPIPS/DISTS、rFID、face identity、OCR/line fidelity、color delta、temporal flicker。",
        "生成系统层：FID/FVD、prompt alignment、motion quality、sampling steps、peak VRAM、encode/decode latency。",
        "复原层：paired fidelity、no-reference quality、hallucination audit、区域化指标、真实相机数据盲测。",
        "流式层：first-frame latency、steady-state FPS、cache memory、chunk boundary consistency、随机 seek/reset 行为。",
        "兼容层：latent scale/shift、channel order、first-frame policy、padding、normalization、checkpoint metadata。",
    ]:
        bullet(doc, text)

    heading(doc, "8.3 最小可复现检查表", 2)
    for text in [
        "固定 tokenizer checkpoint、commit SHA、latent channel、spatial/temporal factor、scale/shift 和输入 range。",
        "先跑 encode-decode round trip，确认重建与官方样例一致，再训练 DiT。",
        "单独记录 tokenizer 与 DiT 的数据清单，避免把 caption 数据误当 VAE 训练数据。",
        "所有 LOSS 权重输出 gradient norm；至少做 rec-only、+perceptual、+GAN、+temporal/frequency 的消融。",
        "视频模型测试不同 T：1、2、4、5、17、61 帧，覆盖 first-frame 和非整除边界。",
        "decoder 替换时冻结 latent 输入，比较完整 decoder 与轻量 decoder 的逐样本差异。",
        "所有速度数据写明硬件、dtype、分辨率、帧数、batch、warm-up、是否含 VAE 和 I/O。",
    ]:
        numbered(doc, text)

    heading(doc, "9. 综合洞察与研究空白", 1)
    insights = [
        ("Tokenizer 正在从预处理变成主架构变量", "TiTok、DC-AE、LTX-Video 和 Cosmos 都在改变 token 数、拓扑或 decoder 职责。未来论文需要把 tokenizer 与 DiT 联合报告，而不是只写 backbone 参数量。"),
        ("高压缩与严格恢复存在结构性冲突", "越强的压缩越依赖 decoder 生成缺失高频。生成任务可接受 plausibility，真实相机复原却需要 evidence-grounded fidelity，因此应提供可控感知档位和幻觉审计。"),
        ("视频因果性是一份完整协议", "causal convolution 不足以保证流式正确；padding、normalization、first-frame、cache reset、chunk overlap 都必须一致。"),
        ("Decoder 是独立的性能优化对象", "FlashVSR 与 LTX-Video 从不同方向说明 decoder 可以承担条件恢复、最后去噪或独立蒸馏。模型加速不能只盯 Transformer。"),
        ("公开结构不等于公开训练", "FLUX.1 和 Wan 的架构可由官方源码核验，但完整 AE/VAE 数据与损失仍未披露。复现报告必须把 architecture reproduction 与 recipe reproduction 分开。"),
        ("数据工程决定可学到的边界", "分辨率 curriculum、FPS、motion filtering、first-frame 和退化链不是外围实现，而是模型方法的一部分。"),
    ]
    for title, body in insights:
        callout(doc, title, body, fill=LIGHT_BLUE if "公开" not in title else LIGHT_GOLD, label_color=DARK_BLUE if "公开" not in title else GOLD)

    heading(doc, "9.1 值得继续研究的问题", 2)
    for text in [
        "面向真实相机复原的可逆或近可逆高压缩 tokenizer，能否同时保留 RAW/ISP 证据与生成先验？",
        "能否用 mask/uncertainty map 控制 decoder 只在证据不足区域生成细节，在可靠区域保持像素忠实？",
        "视频 decoder distillation 如何避免长时漂移，并把 teacher 的 cache behavior 一并蒸馏？",
        "是否可以为同一 latent generator 训练多个 renderer：faithful、aesthetic、mobile、cloud，并共享可校准的质量尺度？",
        "encoder/decoder 的量化误差如何与 flow trajectory、latent normalization 和一步蒸馏共同建模？",
        "公开 benchmark 应如何同时测量重建、生成、因果流式和端侧成本，而不是单一 rFID/FVD？",
    ]:
        bullet(doc, text)


def appendices(doc):
    page_break(doc)
    heading(doc, "附录 A：架构总表", 1)
    add_para(doc, "以下为正文使用的紧凑对照。完整字段、证据锚点和备注请查阅同目录 Excel 的“架构比较”Sheet。")
    rows = []
    for item in ARCH:
        rows.append([
            item["model"],
            item["family"],
            item["compression"],
            item["latent_channels_or_tokens"],
            item["evidence_status"],
        ])
    table = add_table(doc, ["模型", "类别", "压缩", "通道/Token", "证据"], rows, [1.25, 1.45, 1.35, 1.55, 0.9], font_size=7.8)
    set_repeat_table_header(table.rows[0])

    heading(doc, "附录 B：LOSS 与训练阶段总表", 1)
    rows = [[r["model"], r["training_regime"], r["losses"], r["caveat"]] for r in TRAINING]
    table = add_table(doc, ["模型", "训练范式", "损失", "边界"], rows, [1.2, 1.4, 2.2, 1.7], font_size=7.7)
    set_repeat_table_header(table.rows[0])

    heading(doc, "附录 C：术语中英对照", 1)
    rows = [[r["term_en"], r["term_zh"], r["definition"], r["pitfall"]] for r in TERMS]
    table = add_table(doc, ["English", "中文", "Definition", "Pitfall"], rows, [1.25, 1.2, 2.05, 2.0], font_size=8)
    set_repeat_table_header(table.rows[0])

    heading(doc, "参考文献与本地状态", 1)
    add_para(
        doc,
        "采用按研究谱系排序的编号。报告正文不嵌入浏览器 URL；可核验地址、本地路径、SHA-256 与源码 commit 见《编解码器论文与来源索引.xlsx》。",
    )
    for index, paper in enumerate(PAPERS, start=1):
        source = SOURCE_BY_KEY.get(paper["key"], {})
        status = source.get("evidence_status", "indexed")
        local = source.get("local_path", "")
        local_note = local if local else "本地完整 PDF 待补；来源索引已记录"
        p = add_para(doc, after=4, line=1.18)
        add_text(p, f"[{index}] ", size=9, bold=True, color=DARK_BLUE)
        add_text(p, f"{paper['title']}. {paper['venue']}. arXiv:{paper['arxiv_id']}. ", size=9)
        add_text(p, f"Status: {status}. {local_note}", size=8.5, italic=True, color=MUTED)

    heading(doc, "资料库文件说明", 2)
    add_table(
        doc,
        ["文件/目录", "用途"],
        [
            ["papers/", "已下载并校验的完整 PDF；未下载论文不伪造文件"],
            ["text/ 与 metadata/raw/", "PDF 文本证据和 ar5iv 全文快照"],
            ["source_code/", "官方源码快照，code_manifest.json 记录 commit 与 SHA-256"],
            ["figures/paper_figures/", "原论文架构图、效果图和 PDF 页面裁剪"],
            ["figures/explanatory/", "本报告的解释图，SVG 与 PNG 均保留"],
            ["编解码器论文与来源索引.xlsx", "全量审计表：来源、架构、训练、数据、FlashVSR 模块、术语"],
        ],
        [2.35, 4.15],
        font_size=9.2,
    )


def structural_qa(doc_path: Path):
    document = Document(doc_path)
    headings = []
    paragraphs = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
        if paragraph.style and paragraph.style.name.startswith("Heading") and text:
            headings.append({"style": paragraph.style.name, "text": text})
    relationships = document.part.rels.values()
    image_relationships = [rel for rel in relationships if "image" in rel.reltype]
    report = {
        "document": str(doc_path),
        "paragraph_count": len(document.paragraphs),
        "nonempty_paragraph_count": len(paragraphs),
        "heading_count": len(headings),
        "headings": headings,
        "table_count": len(document.tables),
        "image_relationship_count": len(image_relationships),
        "section_count": len(document.sections),
        "contains_flashvsr": any("FlashVSR" in text for text in paragraphs),
        "contains_lq_proj_in": any("LQ_proj_in" in text for text in paragraphs),
        "contains_tcdecoder": any("TCDecoder" in text for text in paragraphs),
        "contains_wandecoder": any("WanDecoder" in text for text in paragraphs),
        "contains_undisclosed_boundary": any("undisclosed" in text for text in paragraphs),
    }
    (METADATA / "report_structure_qa.json").write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def build():
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    cover(document)
    abstract_and_method(document)
    foundations(document)
    interface_and_image_models(document)
    video_models(document)
    training_and_data(document)
    flashvsr_case(document)
    synthesis(document)
    appendices(document)
    document.save(REPORT)
    qa = structural_qa(REPORT)
    print(json.dumps({"report": str(REPORT), **qa}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    build()
