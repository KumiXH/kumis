"""Delivery tests for the illustrated priority-effect report."""

from __future__ import annotations

import json
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "daily" / "20260827_录像特效调研"
PRIORITY_PATH = PROJECT / "metadata" / "priority_effects.jsonl"
MD_PATH = PROJECT / "report" / "手机录像特效重点玩法图文洞察_20260827.md"
DOCX_PATH = PROJECT / "report" / "手机录像特效重点玩法图文洞察_20260827.docx"


def read_jsonl(path: Path) -> list[dict]:
    return [
        json.loads(line)
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]


class EffectReportTests(unittest.TestCase):
    def test_markdown_contains_every_priority_case_and_visual(self) -> None:
        text = MD_PATH.read_text(encoding="utf-8")
        priorities = read_jsonl(PRIORITY_PATH)
        self.assertEqual(len(priorities), 50)
        for row in priorities:
            with self.subTest(priority_id=row["priority_id"]):
                self.assertEqual(text.count(row["priority_id"]), 1)
        self.assertEqual(text.count("本项目概念分镜"), 50)
        self.assertIn("Scalable Diffusion Models with Transformers", text)
        self.assertIn("Apple Cinematic mode", text)

    def test_docx_contains_every_priority_case_and_media(self) -> None:
        document = Document(DOCX_PATH)
        paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        priorities = read_jsonl(PRIORITY_PATH)
        for row in priorities:
            with self.subTest(priority_id=row["priority_id"]):
                self.assertEqual(paragraph_text.count(row["priority_id"]), 1)
        self.assertGreaterEqual(len(document.inline_shapes), 64)

        with ZipFile(DOCX_PATH) as archive:
            self.assertIsNone(archive.testzip())
            media = [name for name in archive.namelist() if name.startswith("word/media/")]
            self.assertGreaterEqual(len(media), 64)

    def test_docx_body_and_table_text_is_not_smaller_than_ten_points(self) -> None:
        document = Document(DOCX_PATH)
        too_small = []
        paragraphs = list(document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        for paragraph in paragraphs:
            for run in paragraph.runs:
                if run.text.strip() and run.font.size and run.font.size.pt < 10:
                    too_small.append((run.text[:40], run.font.size.pt))
        self.assertFalse(too_small[:20], too_small[:20])


if __name__ == "__main__":
    unittest.main()
