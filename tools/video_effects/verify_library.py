"""Cross-artifact verification for the mobile video-effects research library."""

from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree as ET
from zipfile import ZipFile

from docx import Document
from PIL import Image

from tools.video_effects import schema


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "daily" / "20260827_录像特效调研"
NOTES_DIR = PROJECT / "notes"
AUDIT_MD = NOTES_DIR / "final_audit.md"
AUDIT_JSON = NOTES_DIR / "final_audit.json"
EXPECTED_SHEETS = ["总览", "特效原子", "完整玩法", "组合配方", "重点50", "真实参考", "按特效族", "按触发方式", "按生成程度", "字段字典"]


def read_jsonl(path: Path) -> list[dict]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def count_docx_media(path: Path) -> dict:
    with ZipFile(path) as archive:
        corrupt = archive.testzip()
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    return {
        "zip_corrupt_member": corrupt,
        "media_files": len(media),
        "image_relationships": relationships.count("relationships/image"),
        "page_breaks": len(re.findall(r'w:type="page"', document_xml)),
        "table_width_9360_occurrences": document_xml.count('w:w="9360"'),
        "table_indent_120_occurrences": document_xml.count('w:w="120"'),
        "microsoft_yahei_occurrences": document_xml.count("Microsoft YaHei") + styles_xml.count("Microsoft YaHei"),
    }


def xlsx_structure(path: Path) -> dict:
    namespaces = {
        "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
        "rel": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pkg": "http://schemas.openxmlformats.org/package/2006/relationships",
    }
    with ZipFile(path) as archive:
        corrupt = archive.testzip()
        workbook = ET.fromstring(archive.read("xl/workbook.xml"))
        rels = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
        rel_map = {
            relationship.attrib["Id"]: relationship.attrib["Target"]
            for relationship in rels.findall("pkg:Relationship", namespaces)
        }
        rows = {}
        formulas = 0
        names = []
        for sheet in workbook.findall("main:sheets/main:sheet", namespaces):
            name = sheet.attrib["name"]
            names.append(name)
            relationship_id = sheet.attrib[f"{{{namespaces['rel']}}}id"]
            target = rel_map[relationship_id].replace("\\", "/")
            if target.startswith("/"):
                archive_path = target.lstrip("/")
            elif target.startswith("xl/"):
                archive_path = target
            else:
                archive_path = f"xl/{target}"
            root = ET.fromstring(archive.read(archive_path))
            rows[name] = len(root.findall("main:sheetData/main:row", namespaces))
            formulas += len(root.findall(".//main:f", namespaces))
    return {
        "zip_corrupt_member": corrupt,
        "sheet_names": names,
        "sheet_rows": rows,
        "formula_cells": formulas,
    }


def verify() -> dict:
    atoms = read_jsonl(PROJECT / "metadata" / "effect_atoms.jsonl")
    ideas = read_jsonl(PROJECT / "metadata" / "effect_ideas.jsonl")
    recipes = read_jsonl(PROJECT / "metadata" / "effect_recipes.jsonl")
    priorities = read_jsonl(PROJECT / "metadata" / "priority_effects.jsonl")
    references = read_jsonl(PROJECT / "references" / "reference_manifest.jsonl")
    storyboards = read_jsonl(PROJECT / "figures" / "effect_storyboards" / "storyboard_manifest.jsonl")

    atom_ids = {row["atom_id"] for row in atoms}
    idea_ids = {row["effect_id"] for row in ideas}
    for row in atoms:
        schema.validate_atom(row)
    for row in ideas:
        schema.validate_idea(row, atom_ids)
    for row in recipes:
        schema.validate_recipe(row, atom_ids, idea_ids)
    for row in priorities:
        schema.validate_priority(row, idea_ids)
    for row in references:
        schema.validate_reference(row)

    expected_counts = {"atoms": 120, "ideas": 300, "recipes": 200, "priorities": 50, "references": 14, "storyboards": 50}
    actual_counts = {"atoms": len(atoms), "ideas": len(ideas), "recipes": len(recipes), "priorities": len(priorities), "references": len(references), "storyboards": len(storyboards)}
    assert actual_counts == expected_counts, actual_counts

    reference_coverage = Counter(effect_id for row in references for effect_id in row["effect_ids"])
    missing_reference_coverage = sorted(row["effect_id"] for row in priorities if not reference_coverage[row["effect_id"]])
    assert not missing_reference_coverage, missing_reference_coverage

    storyboard_checks = []
    for row in storyboards:
        path = ROOT / row["image_path"]
        assert path.exists(), path
        with Image.open(path) as image:
            storyboard_checks.append((image.width, image.height, image.mode))
            assert image.width >= 1500 and image.height >= 600 and image.mode == "RGB"
        assert row["visual_status"] == "本项目概念分镜"

    markdown_path = PROJECT / "report" / "手机录像特效重点玩法图文洞察_20260827.md"
    markdown = markdown_path.read_text(encoding="utf-8")
    for row in priorities:
        assert markdown.count(row["priority_id"]) == 1, row["priority_id"]
    markdown_images = re.findall(r"!\[[^\]]*\]\(([^)]+)\)", markdown)
    missing_markdown_images = []
    for image_path in markdown_images:
        resolved = (markdown_path.parent / image_path).resolve()
        if not resolved.exists():
            missing_markdown_images.append(str(resolved))
    assert len(markdown_images) == 64, len(markdown_images)
    assert not missing_markdown_images, missing_markdown_images

    docx_path = PROJECT / "report" / "手机录像特效重点玩法图文洞察_20260827.docx"
    document = Document(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for row in priorities:
        assert paragraph_text.count(row["priority_id"]) == 1, row["priority_id"]
    too_small = []
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if run.text.strip() and run.font.size and run.font.size.pt < 10:
                too_small.append((run.text[:40], run.font.size.pt))
    assert not too_small, too_small[:10]
    docx = count_docx_media(docx_path)
    assert docx["zip_corrupt_member"] is None
    assert len(document.inline_shapes) == 64
    assert docx["media_files"] == 64
    assert docx["image_relationships"] == 64
    assert docx["microsoft_yahei_occurrences"] > 0

    workbook_path = PROJECT / "matrix" / "手机录像特效玩法库_20260827.xlsx"
    workbook = xlsx_structure(workbook_path)
    assert workbook["zip_corrupt_member"] is None
    assert workbook["sheet_names"] == EXPECTED_SHEETS, workbook["sheet_names"]
    expected_rows = {"特效原子": 121, "完整玩法": 301, "组合配方": 201, "重点50": 51, "真实参考": 15, "按特效族": 13, "按生成程度": 4, "字段字典": 13}
    for sheet, row_count in expected_rows.items():
        assert workbook["sheet_rows"][sheet] == row_count, (sheet, workbook["sheet_rows"][sheet])
    assert workbook["formula_cells"] >= 40

    render_status_path = PROJECT / "report" / "qa" / "render_qa_status.json"
    render_status = json.loads(render_status_path.read_text(encoding="utf-8")) if render_status_path.exists() else {"status": "not_run", "reason": "render QA script has not been run"}

    return {
        "status": "passed_with_rendering_limitation" if render_status["status"] != "rendered" else "passed",
        "counts": actual_counts,
        "reference_coverage": {"priority_effects_covered": len(priorities), "missing": missing_reference_coverage},
        "storyboards": {"checked": len(storyboard_checks), "unique_dimensions": sorted({f"{w}x{h}:{mode}" for w, h, mode in storyboard_checks})},
        "markdown": {"priority_ids": len(priorities), "image_references": len(markdown_images), "missing_images": missing_markdown_images},
        "docx": {**docx, "paragraphs": len(document.paragraphs), "tables": len(document.tables), "inline_shapes": len(document.inline_shapes), "sections": len(document.sections), "explicit_runs_below_10pt": len(too_small)},
        "xlsx": workbook,
        "render_qa": render_status,
    }


def write_audit(result: dict) -> None:
    NOTES_DIR.mkdir(parents=True, exist_ok=True)
    AUDIT_JSON.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    render = result["render_qa"]
    lines = [
        "# 手机录像特效研究库最终审计",
        "",
        f"- 审计状态：`{result['status']}`",
        f"- 数据数量：原子 {result['counts']['atoms']}，完整玩法 {result['counts']['ideas']}，组合配方 {result['counts']['recipes']}，重点案例 {result['counts']['priorities']}，参考 {result['counts']['references']}，概念分镜 {result['counts']['storyboards']}",
        f"- 重点案例参考覆盖：{result['reference_coverage']['priority_effects_covered']}/50",
        f"- Markdown：50 个重点 ID，{result['markdown']['image_references']} 个图片引用，缺图 {len(result['markdown']['missing_images'])}",
        f"- DOCX：{result['docx']['paragraphs']} 个段落，{result['docx']['tables']} 个表格，{result['docx']['inline_shapes']} 张嵌入图，媒体文件 {result['docx']['media_files']}，显式小于 10 磅的文本 {result['docx']['explicit_runs_below_10pt']}",
        f"- XLSX：{len(result['xlsx']['sheet_names'])} 个工作表，公式单元格 {result['xlsx']['formula_cells']}，ZIP 完整性通过",
        "",
        "## 已验证",
        "",
        "- 所有 JSONL 通过 schema 验证，ID 引用与数量一致",
        "- 50 张概念分镜均存在、为 RGB，且显式标注“本项目概念分镜”",
        "- 14 张参考卡均存在，参考清单明确写出能够证明与不能证明的内容",
        "- Markdown 的 64 个图片路径全部有效",
        "- DOCX ZIP、媒体关系、重点 ID、微软雅黑字体和最小字号规则通过",
        "- XLSX 工作表名、明细行数、公式单元格和 ZIP 结构通过",
        "",
        "## 未验证",
        "",
        f"- Word 逐页 PNG 视觉渲染状态：`{render['status']}`",
        f"- 原因：{render['reason'] or '无'}",
        "- 因当前主机没有 LibreOffice/soffice 或 Microsoft Word，未能检查 Word 页面级裁切、分页和图文重叠；已以结构、媒体、字体和源图联系表审计作为替代，但不能等同于完整视觉门禁",
        "",
    ]
    AUDIT_MD.write_text("\n".join(lines), encoding="utf-8", newline="\n")


if __name__ == "__main__":
    audit = verify()
    write_audit(audit)
    print(json.dumps(audit, ensure_ascii=False, indent=2))
