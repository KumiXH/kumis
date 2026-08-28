"""Build the illustrated Markdown and Word companion for 50 priority effects."""

from __future__ import annotations

import json
import math
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[2]
PROJECT = ROOT / "daily" / "20260827_录像特效调研"
REPORT_DIR = PROJECT / "report"
MD_PATH = REPORT_DIR / "手机录像特效重点玩法图文洞察_20260827.md"
DOCX_PATH = REPORT_DIR / "手机录像特效重点玩法图文洞察_20260827.docx"

NAVY = "16324F"
BLUE = "1B6B93"
DARK_BLUE = "17465E"
INK = "243447"
MUTED = "5E6D7E"
LIGHT_BLUE = "E8F1F6"
LIGHT_GREEN = "E6F4EF"
LIGHT_GOLD = "FFF2CC"
LIGHT_RED = "FCE8E6"
LIGHT_GRAY = "F3F6F8"
WHITE = "FFFFFF"
FONT = "Microsoft YaHei"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

FAMILY_LABELS = {
    "light_trails_optics": "光轨与可编程光学",
    "body_motion_clones": "身体运动与分身",
    "face_gaze_expression": "面部、视线与表情",
    "time_editing": "局部时间编辑",
    "spatial_portals": "空间入口与传送门",
    "virtual_light_shadow": "虚拟光影",
    "material_morph": "材质变形",
    "particles_weather": "粒子与天气",
    "world_style": "世界风格化",
    "audio_lyrics": "音频与歌词",
    "effect_cinematography": "特效摄影",
    "multi_person_interaction": "多人互动",
}

FAMILY_ENGINEERING = {
    "light_trails_optics": {
        "data": "室内外手部/人体/移动光源轨迹视频；同步相机姿态、触摸采样、手部关键点和可选深度；合成数据覆盖不同笔画、遮挡、曝光与平面姿态。",
        "loss": "L_path（轨迹点与速度）、L_anchor（平面/世界锚点）、L_occ（遮挡边界）、L_render（发光层重建）与 L_temp（亮度和位置时序连续）按任务分阶段组合。",
        "metrics": "轨迹重投影误差、遮挡 F1、笔画断裂率、锚点漂移、原画面高光保真、用户起笔/闭合成功率。",
        "constraint": "光效可夸张，但原始人物、场景几何和真实高光必须保留可回退版本。",
    },
    "body_motion_clones": {
        "data": "多动作人体视频、姿态与人体 mask、长短遮挡片段、不同服装和背景；通过时间偏移与轨迹重排生成分身监督。",
        "loss": "L_mask、L_pose、L_id、L_flow、L_occ 与 L_temp；对分身外区域使用背景重建约束，对身份敏感 ROI 使用特征一致性。",
        "metrics": "身份相似度、姿态误差、mask 边界 F1、遮挡顺序错误率、克隆闪烁率和动作时间偏差。",
        "constraint": "不得把历史分身误写成现场真实人物；录制元数据应明确副本时间戳。",
    },
    "face_gaze_expression": {
        "data": "多机位或屏幕标定的人脸视频；包含视线目标、头姿、虹膜/眼睑关键点、眨眼、眼镜反光、侧脸和多人对话；需有身份授权。",
        "loss": "L_gaze（角度）、L_landmark、L_iris、L_identity、L_background、L_blink 与 L_temporal；只在眼部小 ROI 内优化，外部区域强约束为原帧。",
        "metrics": "视线角误差、身份余弦相似度、眼睑关键点误差、眨眼保持率、双眼一致性、用户自然度评分。",
        "constraint": "低置信度、强侧脸、闭眼或反光时回退原帧；不得生成新的虹膜身份特征。",
    },
    "time_editing": {
        "data": "具有对象 mask、光流、事件边界和动作阶段标注的视频；用正放、倒放、冻结、局部循环和时间偏移自动构造监督。",
        "loss": "L_recon、L_flow、L_boundary、L_cycle、L_order 与 L_seam；非编辑区域施加严格保持损失。",
        "metrics": "循环接缝误差、边界泄漏、对象身份交换率、光流一致性、时间事件可读性和原区域保真。",
        "constraint": "时间效果必须明确是创作性改写，原始时间线和事件顺序需保留。",
    },
    "spatial_portals": {
        "data": "带深度、平面、相机位姿、镜面/地面/手掌 mask 的移动视频；配对另一场景或合成门户内容，并标注进入/退出事件。",
        "loss": "L_plane、L_depth、L_warp、L_border、L_occ、L_camera 与 L_temp；生成内容另加身份和几何约束。",
        "metrics": "边框重投影误差、门户漂移、穿越遮挡正确率、视差一致性、空洞率和时间闪烁。",
        "constraint": "门户另一侧是生成或重放内容，不能被解释为真实摄像头已经观察到的空间。",
    },
    "virtual_light_shadow": {
        "data": "多光源人像和场景视频；线性 RGB/HDR、主体 mask、法线、深度、光源方向和真实阴影；可用渲染器合成多样光照监督。",
        "loss": "L_relight、L_chroma、L_edge、L_normal、L_shadow、L_identity 与 L_temp；高光区增加饱和保护。",
        "metrics": "受光方向一致性、肤色偏差、阴影接触误差、边缘溢光、时序闪烁和原始动态范围保持。",
        "constraint": "虚拟光影属于感知增强；不得覆盖脸部身份细节或改变真实安全信息。",
    },
    "material_morph": {
        "data": "物体/服装材质视频、实例 mask、表面法线、光流与可控材质参数；渲染合成玻璃、金属、全息和溶解状态序列。",
        "loss": "L_mask、L_surface、L_normal、L_reflection、L_material、L_identity 与 L_temp；边界和非编辑区施加重建损失。",
        "metrics": "材质区域 IoU、表面附着误差、反射方向一致性、边缘泄漏、时序稳定和身份保持。",
        "constraint": "材质可生成，但目标物体轮廓、人体身份和交互接触点必须稳定。",
    },
    "particles_weather": {
        "data": "带深度与人物 mask 的真实/合成粒子视频；雨、花瓣、尘埃、歌词粒子与光束，覆盖不同景深、风向和遮挡。",
        "loss": "L_depth_order、L_particle_motion、L_occ、L_density、L_beat 与 L_temp；人物边缘使用抠像/遮挡监督。",
        "metrics": "前后景顺序正确率、粒子轨迹平滑度、节拍同步偏差、边缘穿透率和密度稳定性。",
        "constraint": "粒子不能遮挡关键面部、字幕或安全区域；必要时自动降低密度。",
    },
    "world_style": {
        "data": "多场景长视频及风格参考；相机轨迹、深度、语义 mask、结构边缘和色彩统计；需覆盖昼夜、季节和复杂运动。",
        "loss": "L_structure、L_content、L_style、L_semantic、L_identity、L_flow 与 L_temp；生成式区域加入事实和几何保持门控。",
        "metrics": "结构边缘保持、身份相似度、风格一致性、光流扭曲误差、闪烁率和事实改写人工审查。",
        "constraint": "这是生成式或感知式世界改写；必须保留原片并显式标注创作效果。",
    },
    "audio_lyrics": {
        "data": "多说话人/演唱视频、分离音轨、说话人时间戳、ASR 字幕、嘴型关键点、声源方向和人物轨迹。",
        "loss": "L_sync、L_speaker、L_lip、L_layout、L_occ、L_readability 与 L_temp；错误 ASR 使用低置信度隐藏策略。",
        "metrics": "音画同步偏差、说话人归属准确率、字幕可读性、遮挡正确率、嘴型关联和切换抖动。",
        "constraint": "不得伪造说话内容；歌词/字幕必须与原音频绑定并保留置信度。",
    },
    "effect_cinematography": {
        "data": "带焦点、深度、人物 mask、相机内参与触发事件的素材；可用自动裁切生成擦镜、分屏、冲击变焦和焦点脉冲监督。",
        "loss": "L_frame、L_focus、L_mask、L_transition、L_camera 与 L_temp；保持主体关键区域不被转场裁断。",
        "metrics": "触发时刻误差、主体保留率、过渡接缝、焦点命中率、构图偏差和观感评分。",
        "constraint": "只保留具有明显可见效果的摄影玩法；不把普通防抖或常规运镜作为主功能。",
    },
    "multi_person_interaction": {
        "data": "双人到多人交互视频、每人身份/姿态/mask、接触点、说话人和遮挡顺序；覆盖交叉、出入框和身份交换难例。",
        "loss": "L_multi_id、L_contact、L_pose、L_assignment、L_occ、L_sync 与 L_temp；用匈牙利匹配或显式轨迹 ID 保持人物归属。",
        "metrics": "身份交换率、接触点误差、多人遮挡正确率、同步误差、互动完成率和失败回退率。",
        "constraint": "任何人的效果归属都不能串到另一人；身份不确定时停用互动层。",
    },
}


def read_jsonl(path: Path) -> list[dict]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def repo_path(path: Path) -> str:
    return path.resolve().relative_to(ROOT.resolve()).as_posix()


def references_for_effect(references: list[dict], effect_id: str) -> list[dict]:
    return [reference for reference in references if effect_id in reference["effect_ids"]]


def md_escape(text: str) -> str:
    return text.replace("|", "\\|").replace("\n", " ")


def build_markdown(priorities: list[dict], ideas: dict[str, dict], references: list[dict], storyboards: dict[str, dict]) -> str:
    family_counts = Counter(ideas[row["effect_id"]]["family"] for row in priorities)
    reference_numbers = {row["reference_id"]: index for index, row in enumerate(references, 1)}
    lines = [
        "# 手机录像特效重点玩法图文洞察",
        "",
        "_50 个交互型录像特效的产品体验、实现链路、数据、Loss、风险与可核实参考 | 2026-08-27_",
        "",
        "---",
        "",
        "## 研究摘要",
        "",
        "本手册从 300 个完整玩法中抽取 50 个重点案例，覆盖 12 个特效族。研究对象是用户能够在取景器或成片中直接看见、触发和调节的录像特效，不把普通防抖、常规降噪、HDR、超分或一般运镜当作主玩法。每个案例同时给出三帧概念分镜、交互状态、模块级数据流、建议训练数据、建议 Loss、移动端预览边界、录后精修和失败降级。",
        "",
        "> **证据边界：** 14 份官方产品/软件或学术论文参考只用于证明基础能力。概念分镜不是既有产品截图，训练和 Loss 组合是工程建议，不是厂商未公开事实。性能数字只有在目标设备实测后才能进入结论。",
        "",
        "| 资产 | 数量 | 用途 |",
        "| --- | ---: | --- |",
        "| 特效原子 | 120 | 可组合的最低层视觉、交互与时序能力 |",
        "| 完整玩法 | 300 | 具有可见结果、触发、控制、预览和录后流程 |",
        "| 组合配方 | 200 | 跨原子和玩法的显式组合 |",
        "| 重点案例 | 50 | 本手册逐项深拆 |",
        "| 可核实参考 | 14 | 说明能够证明与不能证明的边界 |",
        "",
        "## 技术总线",
        "",
        "```mermaid",
        "flowchart LR",
        "    accTitle: Mobile Video Effect Pipeline",
        "    accDescr: A two-stage mobile video effect pipeline that creates a low-cost preview during recording, stores editable metadata, and recomputes a high-quality result after recording with fallbacks to the original video",
        "",
        "    capture[采集 RGB HDR 深度 IMU 音频] --> perceive[感知 检测 分割 跟踪 姿态]",
        "    perceive --> event[交互事件与状态机]",
        "    event --> preview[低成本 ROI 预览合成]",
        "    preview --> record[原片与结构化元数据]",
        "    record --> refine[录后高分辨率重跟踪与重渲染]",
        "    refine --> audit{置信度和真实性检查}",
        "    audit -->|通过| output[可编辑成片]",
        "    audit -->|失败| fallback[原片或局部回退]",
        "",
        "    classDef source fill:#e8f1f6,stroke:#1b6b93,stroke-width:2px,color:#16324f",
        "    classDef process fill:#e6f4ef,stroke:#00866a,stroke-width:2px,color:#145c4a",
        "    classDef decision fill:#fff2cc,stroke:#a66a00,stroke-width:2px,color:#6b4d00",
        "    classDef output_style fill:#fce8e6,stroke:#b43c35,stroke-width:2px,color:#7a2420",
        "    class capture source",
        "    class perceive,event,preview,record,refine process",
        "    class audit decision",
        "    class output,fallback output_style",
        "```",
        "",
        "## 特效族分布",
        "",
        "| 特效族 | 重点案例数 | 主要技术难点 |",
        "| --- | ---: | --- |",
    ]
    for family, count in family_counts.items():
        lines.append(f"| {FAMILY_LABELS[family]} | {count} | {md_escape(FAMILY_ENGINEERING[family]['constraint'])} |")

    lines.extend([
        "",
        "## 产品与研究原则",
        "",
        "1. **预览负责交互，录后负责质量。** 取景器只运行 ROI、代理分辨率、分级更新和有限实例数；原片与结构化元数据必须同时保存。",
        "2. **可见失败必须有明确降级。** 跟踪丢失、遮挡错误、身份不确定或深度漂移时，不继续扩大生成结果，而是冻结、减弱、回退屏幕空间或恢复原帧。",
        "3. **生成内容与真实观察分离。** 空间入口、世界风格和材质重绘属于创作结果，不能被写成摄像头真实观察。",
        "4. **人物身份优先于视觉刺激。** 人脸、瞳孔、嘴型、人体归属和多人轨迹必须有单独的身份/几何约束。",
        "5. **性能结论必须设备实测。** 本手册不虚构毫秒、FPS、功耗、内存或温升数据。",
        "",
        "## 重点案例",
        "",
    ])

    for index, priority in enumerate(priorities, 1):
        idea = ideas[priority["effect_id"]]
        family = idea["family"]
        engineering = FAMILY_ENGINEERING[family]
        relevant_refs = references_for_effect(references, priority["effect_id"])
        storyboard = storyboards[priority["priority_id"]]
        image_rel = Path("..") / Path(storyboard["image_path"]).relative_to(PROJECT.relative_to(ROOT))
        lines.extend([
            f"### 案例 {index:02d}：{idea['name_zh']}",
            "",
            f"`{priority['priority_id']}`",
            "",
            f"![{idea['name_zh']}三帧概念分镜]({image_rel.as_posix()})",
            f"_图 {index}：本项目概念分镜。依次展示录制前、触发中和录后成片；不代表既有产品、论文实验结果或已测性能。_",
            "",
            f"**研究问题。** {priority['problem']}",
            "",
            f"**目标体验。** {priority['experience_story']}",
            "",
            f"**可见创新。** {idea['novelty']} {idea['visible_effect']}",
            "",
            "**交互时间线**",
            "",
        ])
        for item in priority["interaction_timeline"]:
            lines.append(f"- {item}")
        lines.extend(["", "**模块和信号链**", ""])
        for item in priority["module_pipeline"]:
            lines.append(f"- {item}")
        lines.extend([
            "",
            f"`{priority['tensor_or_signal_flow']}`",
            "",
            "**数据与训练建议**",
            "",
            f"- 数据：{engineering['data']}",
            f"- Loss：{engineering['loss']}",
            f"- 指标：{engineering['metrics']}",
            f"- 真实性/身份约束：{engineering['constraint']}",
            "",
            "**端侧预览、录后精修与控制**",
            "",
            f"- 预览：{priority['preview_budget']}",
            f"- 录后：{priority['post_refinement']}",
            f"- 元数据：{'；'.join(priority['recorded_metadata'])}",
            f"- 参数：`{'`, `'.join(priority['adjustable_parameters'])}`",
            f"- 产品形态：{priority['mobile_product_form']}",
            "",
            "**失败与降级**",
            "",
        ])
        for item in priority["failure_and_fallback"]:
            lines.append(f"- {item}")
        ref_text = "；".join(
            f"{reference['title']}[^${reference_numbers[reference['reference_id']]}]".replace("^$", "^")
            for reference in relevant_refs
        )
        lines.extend([
            "",
            f"**参考边界。** {ref_text}。这些来源只证明对应的基础能力；不能据此声称本完整特效已经量产。",
            "",
        ])

        if index == 1:
            lines.extend([
                "#### 光绘专门分析",
                "",
                "实时光绘的关键不是画一条二维线，而是把一次手部动作变成可编辑事件：起笔、轨迹、平面绑定、闭合确认、遮挡和结束。预览可用短历史缓冲与低分辨率发光层；录后再用高分辨率手部重跟踪、平面重投影和遮挡重建修补断笔。训练时不应把所有模块端到端强行统一：手部关键点、深度/平面、轨迹滤波和渲染器可以分别预训练，再通过轨迹与合成损失联合微调。",
                "",
                "```mermaid",
                "stateDiagram-v2",
                "    accTitle: Spatial Light Painting States",
                "    accDescr: The interaction states for spatial light painting from hand detection through stroke creation, world anchoring, confirmation, persistence, and fallback",
                "",
                "    [*] --> Ready: 检测手部和平面",
                "    Ready --> Drawing: 指尖超过起笔阈值",
                "    Drawing --> Anchored: 轨迹绑定可信平面",
                "    Drawing --> ScreenFallback: 平面置信度低",
                "    Anchored --> Confirmed: 闭合手势或触摸确认",
                "    Confirmed --> Persistent: 保留空间光字",
                "    Persistent --> Editing: 擦除或再次触碰",
                "    Editing --> Persistent: 参数更新",
                "    ScreenFallback --> [*]: 输出屏幕空间短拖尾",
                "    Persistent --> [*]: 录制结束",
                "```",
                "",
                "建议的联合目标可写为：`L = λp L_path + λa L_anchor + λo L_occ + λr L_render + λt L_temporal`。权重不是固定事实，必须通过梯度量级、消融实验和目标设备数据重新确定。",
                "",
            ])
        elif index == 2:
            lines.extend([
                "#### 视线矫正专门分析",
                "",
                "身份保持的核心不是让输出眼睛看起来更锐，而是限制允许修改的自由度。建议把编辑变量限定为小幅虹膜中心位移、局部眼睑补偿和 alpha 融合；脸型、皮肤、眉毛、睫毛、虹膜纹理和闭眼状态默认来自原帧。身份表征可由 ArcFace 类人脸识别特征、眼周局部特征和关键点几何共同计算，但身份损失只能是代理指标，不能替代人类对自然度和个人特征的判断。",
                "",
                "建议目标：`L = λg L_gaze + λlm L_landmark + λid L_identity + λbg L_outside + λb L_blink + λt L_temporal`。其中 `L_outside` 对眼部 ROI 外区域施加强保持；`L_blink` 避免闭眼被错误补成睁眼；低置信度样本不进入生成式补眼分支。",
                "",
                "训练数据应覆盖屏幕中心、镜头中心、左右提示点、多头姿、不同镜片反光和不同眼型，并记录每只眼的可见性。评估不能只看视线角误差，还要同时看身份相似度、眼睑关键点、眨眼保持、双眼一致性和用户自然度。",
                "",
            ])

    lines.extend([
        "## 可核实参考",
        "",
        "参考卡中的“能够证明”和“不能证明”应成对阅读。桌面后期软件、论文原型和静态摄影能力不能自动升级为手机录像实时量产结论。",
        "",
    ])
    for index, reference in enumerate(references, 1):
        image_rel = Path("..") / "figures" / "real_references" / f"{index:02d}_{reference['reference_id'].lower()}.png"
        lines.extend([
            f"### {reference['title']}",
            "",
            f"![{reference['title']}可核实参考卡]({image_rel.as_posix()})",
            f"_参考卡 {index}：{reference['publisher']}，{reference['year']}，边界 `{reference['implementation_boundary']}`。这不是产品效果截图。_",
            "",
            f"- 能够证明：{reference['demonstrates']}",
            f"- 不能证明：{reference['does_not_prove']}",
            f"- 本地证据：{'；'.join(reference['local_files']) if reference['local_files'] else '仅核验元数据'}",
            "",
        ])

    lines.extend([
        "## 结论",
        "",
        "手机录像特效的产品机会不在于把桌面后期完整搬进取景器，而在于建立一条可编辑录像资产链：实时感知和交互负责让用户知道特效将发生在哪里，结构化元数据保留对象、事件、深度和参数，录后端侧或云端负责高质量重算。最值得优先验证的是实时光绘、视线矫正、局部时间编辑、空间入口、虚拟光影和多人互动，因为这些方向的用户可见价值强，同时能复用现有跟踪、分割、深度、音频和生成模型能力。",
        "",
        "后续原型评估应同时报告成功率、失败降级、身份/事实保持、设备性能和用户观感，不应只展示成功样片。",
        "",
        "## References",
        "",
    ])
    for index, reference in enumerate(references, 1):
        lines.append(f"[^{index}]: {reference['publisher']}. ({reference['year']}). \"{reference['title']}.\" {reference['original_source']}")
    lines.extend(["", f"_Last updated: {date.today().isoformat()}_", ""])
    return "\n".join(lines)


def set_run_font(run, size: float = 11, bold: bool | None = None, italic: bool | None = None, color: str = INK) -> None:
    run.font.name = FONT
    run.font.size = Pt(max(size, 10))
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag, width in (("tblW", CONTENT_WIDTH_DXA), ("tblInd", TABLE_INDENT_DXA)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DARK_BLUE),
    ):
        style = document.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("Mobile Video Effects Research | 2026-08-27"), 10, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("ReadPaper research companion | Page "), 10, color=MUTED)
    add_page_number(footer)


def add_paragraph(document: Document, text: str = "", *, size=11, bold=False, italic=False, color=INK, align=None, before=0, after=6, keep=False):
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.keep_with_next = keep
    if text:
        set_run_font(paragraph.add_run(text), size, bold=bold, italic=italic, color=color)
    return paragraph


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    set_run_font(paragraph.add_run(text), 10)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, {1: 16, 2: 13, 3: 12}[level], bold=True, color=BLUE if level < 3 else DARK_BLUE)


def add_callout(document: Document, label: str, text: str, fill=LIGHT_BLUE, label_color=BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    set_run_font(paragraph.add_run(f"{label}："), 10, bold=True, color=label_color)
    set_run_font(paragraph.add_run(text), 10, color=INK)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.add_run(header), 10, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            if row_index % 2:
                set_cell_shading(cell, LIGHT_GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_run_font(cell.paragraphs[0].add_run(str(value)), 10)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def image_size(path: Path, max_width=6.45, max_height=4.15) -> tuple[Inches, Inches]:
    with Image.open(path) as image:
        width, height = image.size
    scale = min(max_width / width, max_height / height)
    return Inches(width * scale), Inches(height * scale)


def add_figure(document: Document, path: Path, caption: str, source: str, *, max_width=6.45, max_height=4.15) -> None:
    width, height = image_size(path, max_width=max_width, max_height=max_height)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(path), width=width, height=height)
    shape._inline.docPr.set("descr", caption)
    caption_paragraph = add_paragraph(document, caption, size=10, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, keep=True)
    add_paragraph(document, f"证据/来源：{source}", size=10, italic=True, color=MUTED, after=5)


def page_break(document: Document) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_cover(document: Document) -> None:
    add_paragraph(document, before=70, after=20)
    add_paragraph(document, "RESEARCH COMPANION", size=10, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_paragraph(document, "手机录像特效重点玩法图文洞察", size=28, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_paragraph(document, "50 Interactive Mobile Video Effects", size=16, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    add_paragraph(document, "产品体验 · 交互状态 · 模块链路 · 数据与 Loss · 失败降级 · 可核实证据", size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=60)
    add_table(document, ["资产", "数量", "说明"], [
        ["重点案例", "50", "12 个特效族，逐项包含三帧概念分镜"],
        ["完整玩法库", "300", "另含 120 个原子与 200 个组合配方"],
        ["可核实参考", "14", "官方产品/软件与论文，均写明证据边界"],
        ["真实性原则", "原片优先", "概念分镜不是产品截图；生成结果必须可回退"],
    ], [1800, 1200, 6360])
    add_paragraph(document, "版本：2026-08-27", size=10, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, before=34, after=4)
    add_paragraph(document, "研究目录：daily/20260827_录像特效调研", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    page_break(document)


def add_overview(document: Document, priorities: list[dict], ideas: dict[str, dict], references: list[dict]) -> None:
    add_heading(document, "摘要与研究口径", 1)
    add_paragraph(document, "本报告面向手机录像特效预研，从 300 个完整玩法中选择 50 个重点案例。主线不是传统画质修复，而是用户能够看见、触发、控制并愿意分享的特效：光绘、视线矫正、分身、局部时间、传送门、虚拟光影、材质、粒子、世界风格、音频歌词、特效摄影和多人互动。")
    add_callout(document, "证据边界", "14 份参考只证明基础能力。概念分镜是本项目视觉化表达，训练数据、Loss 和移动端拆分是工程建议；未公开的厂商实现和未实测性能不写成事实。", fill=LIGHT_GOLD, label_color="8A5A00")
    family_counts = Counter(ideas[row["effect_id"]]["family"] for row in priorities)
    add_table(document, ["特效族", "重点数", "真实性/身份主约束"], [
        [FAMILY_LABELS[family], str(count), FAMILY_ENGINEERING[family]["constraint"]]
        for family, count in family_counts.items()
    ], [2300, 1100, 5960])
    add_heading(document, "统一实现框架", 2)
    for text in (
        "采集层保留原始 RGB/HDR、相机内参、IMU、音频和可选深度；不把低成本预览当成最终母版。",
        "感知层输出人物/物体轨迹、语义 mask、深度/平面、姿态、视线、节拍和交互事件。",
        "预览层使用 ROI、代理分辨率、分级更新、历史缓存和有限实例数，让用户理解特效位置与触发状态。",
        "录制层保存原片和结构化元数据，使轨迹、遮挡、参数和事件都能录后重新计算。",
        "精修层进行高分辨率重跟踪、边界细化、时序优化与真实性审计；不通过时局部或整体回退原片。",
    ):
        add_bullet(document, text)
    add_heading(document, "评价方法", 2)
    add_paragraph(document, "每个原型至少同时评估交互成功率、时序稳定、身份/事实保持、失败降级、设备性能和用户观感。不能只选择成功样片，也不能用单一感知指标替代对象归属、几何和真实性检查。")
    add_callout(document, "性能口径", "报告没有写入未经目标设备实测的毫秒、FPS、功耗、内存或温升数字。预览预算只描述分辨率、ROI、更新层级、实例上限、缓存和降级顺序。", fill=LIGHT_GREEN, label_color="147A62")
    page_break(document)


def add_case(document: Document, index: int, priority: dict, idea: dict, references: list[dict], storyboard: dict) -> None:
    if index > 1:
        page_break(document)
    add_heading(document, f"案例 {index:02d}：{idea['name_zh']}", 1)
    add_paragraph(document, priority["priority_id"], size=10, bold=True, color=BLUE, after=4)
    add_paragraph(document, f"{FAMILY_LABELS[idea['family']]} | {idea['generation_level']} | 端侧难度：{idea['edge_difficulty']}", size=10, color=MUTED, after=8)
    image_path = ROOT / storyboard["image_path"]
    add_figure(
        document,
        image_path,
        f"图 {index}  本项目概念分镜：录制前、触发中、录后成片。",
        "本项目绘制；不代表既有产品、论文实验结果或已测性能。",
    )
    add_heading(document, "问题、体验与创新", 2)
    add_paragraph(document, f"研究问题：{priority['problem']}")
    add_paragraph(document, f"目标体验：{priority['experience_story']}")
    add_paragraph(document, f"可见创新：{idea['novelty']} {idea['visible_effect']}")
    add_heading(document, "交互状态", 2)
    for item in priority["interaction_timeline"]:
        add_bullet(document, item)
    add_heading(document, "模块与信号链", 2)
    for item in priority["module_pipeline"]:
        add_bullet(document, item)
    add_callout(document, "张量/信号流", priority["tensor_or_signal_flow"], fill=LIGHT_BLUE)

    engineering = FAMILY_ENGINEERING[idea["family"]]
    add_heading(document, "数据、Loss 与评价建议", 2)
    add_paragraph(document, f"数据制作：{engineering['data']}")
    add_paragraph(document, f"Loss 组合：{engineering['loss']}")
    add_paragraph(document, f"评价指标：{engineering['metrics']}")
    add_callout(document, "真实性/身份边界", engineering["constraint"], fill=LIGHT_GOLD, label_color="8A5A00")

    if index == 1:
        add_heading(document, "实时光绘专项拆解", 2)
        add_paragraph(document, "二维拖尾升级为空间光绘的核心，是把手指动作记录为可编辑事件序列：起笔、轨迹、平面绑定、闭合确认、遮挡、擦除和结束。手部关键点、平面/深度、轨迹滤波和发光合成器可以分模块预训练，再用轨迹和合成目标联合微调。")
        add_callout(document, "建议目标", "L = λp L_path + λa L_anchor + λo L_occ + λr L_render + λt L_temporal。权重必须由梯度量级和消融实验确定，不是固定事实。", fill=LIGHT_GREEN, label_color="147A62")
        add_table(document, ["状态", "进入条件", "输出/回退"], [
            ["Ready", "检测到手部和稳定背景", "显示候选起笔点"],
            ["Drawing", "指尖速度超过阈值", "保存屏幕轨迹并预览光核"],
            ["Anchored", "平面和深度置信度通过", "轨迹转为世界锚点"],
            ["Confirmed", "闭合手势或触摸确认", "固化笔画并允许节拍调制"],
            ["Fallback", "平面、遮挡或轨迹不可信", "退回屏幕短拖尾或原片"],
        ], [1500, 3360, 4500])
    elif index == 2:
        add_heading(document, "视线矫正与身份保持专项拆解", 2)
        add_paragraph(document, "允许编辑的自由度应尽量小：小幅虹膜中心位移、必要的眼睑补偿和局部 alpha 融合。脸型、皮肤、眉毛、睫毛、虹膜纹理、闭眼状态和眼部 ROI 外画面默认来自原帧。身份表征可以由 ArcFace 类全脸特征、眼周局部特征和关键点几何共同计算，但它们都是代理量，不能替代人类自然度评估。")
        add_callout(document, "建议目标", "L = λg L_gaze + λlm L_landmark + λid L_identity + λbg L_outside + λb L_blink + λt L_temporal。低置信度、强侧脸、闭眼或镜片反光时不进入生成式补眼。", fill=LIGHT_GREEN, label_color="147A62")
        add_table(document, ["保护对象", "建议约束", "失败处理"], [
            ["视线方向", "角度损失 + 标定偏差", "超过最大偏移时减弱或关闭"],
            ["身份", "全脸/眼周特征一致性", "身份相似度下降时回退原 ROI"],
            ["眨眼与眼睑", "关键点 + 闭眼状态保持", "不把闭眼补成睁眼"],
            ["双眼几何", "左右眼一致性与头姿条件", "不一致时零偏移"],
            ["眼外区域", "强重建保持", "禁止联动改变脸型和皮肤"],
        ], [1700, 3860, 3800])

    add_heading(document, "预览、录后与产品控制", 2)
    add_paragraph(document, f"预览预算：{priority['preview_budget']}")
    add_paragraph(document, f"录后精修：{priority['post_refinement']}")
    add_paragraph(document, f"录制元数据：{'；'.join(priority['recorded_metadata'])}")
    add_paragraph(document, f"可调参数：{'；'.join(priority['adjustable_parameters'])}")
    add_paragraph(document, f"产品形态：{priority['mobile_product_form']}")
    add_heading(document, "失败与降级", 2)
    for item in priority["failure_and_fallback"]:
        add_bullet(document, item)

    relevant = references_for_effect(references, priority["effect_id"])
    ref_summary = "；".join(f"{item['title']}（{item['implementation_boundary']}）" for item in relevant)
    add_callout(document, "参考边界", f"{ref_summary}。这些来源只证明相关基础能力，不能证明本完整特效已经量产。", fill=LIGHT_RED, label_color="A23B34")


def add_reference_appendix(document: Document, references: list[dict]) -> None:
    page_break(document)
    add_heading(document, "可核实参考附录", 1)
    add_paragraph(document, "每张参考卡都把来源能够证明和不能证明的内容并列。官方产品、桌面后期和论文原型不应被混写为手机录像实时实现。")
    for index, reference in enumerate(references, 1):
        if index > 1:
            page_break(document)
        add_heading(document, reference["title"], 2)
        card_path = PROJECT / "figures" / "real_references" / f"{index:02d}_{reference['reference_id'].lower()}.png"
        add_figure(
            document,
            card_path,
            f"参考卡 {index}  {reference['publisher']}，{reference['year']}。",
            "本项目证据卡；不是产品效果截图。",
            max_height=3.2,
        )
        add_table(document, ["字段", "内容"], [
            ["来源类型", reference["source_type"]],
            ["实现边界", reference["implementation_boundary"]],
            ["能够证明", reference["demonstrates"]],
            ["不能证明", reference["does_not_prove"]],
            ["原始来源", reference["original_source"]],
            ["本地证据", "；".join(reference["local_files"]) if reference["local_files"] else "仅核验元数据"],
            ["绑定重点玩法", f"{len(reference['effect_ids'])} 项"],
        ], [1800, 7560])


def set_properties(document: Document) -> None:
    props = document.core_properties
    props.title = "手机录像特效重点玩法图文洞察"
    props.subject = "50 个交互型手机录像特效的产品、算法、数据、Loss 与风险研究"
    props.author = "ReadPaper Research Workspace"
    props.keywords = "手机录像, 视频特效, 光绘, 视线矫正, 时间编辑, 生成式视频"
    props.comments = "包含 50 张本项目概念分镜和 14 张可核实参考卡。"


def build() -> dict:
    priorities = read_jsonl(PROJECT / "metadata" / "priority_effects.jsonl")
    ideas = {row["effect_id"]: row for row in read_jsonl(PROJECT / "metadata" / "effect_ideas.jsonl")}
    references = read_jsonl(PROJECT / "references" / "reference_manifest.jsonl")
    storyboards = {row["priority_id"]: row for row in read_jsonl(PROJECT / "figures" / "effect_storyboards" / "storyboard_manifest.jsonl")}

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    markdown = build_markdown(priorities, ideas, references, storyboards)
    MD_PATH.write_text(markdown, encoding="utf-8", newline="\n")

    document = Document()
    configure_document(document)
    set_properties(document)
    add_cover(document)
    add_overview(document, priorities, ideas, references)
    for index, priority in enumerate(priorities, 1):
        add_case(document, index, priority, ideas[priority["effect_id"]], references, storyboards[priority["priority_id"]])
    add_reference_appendix(document, references)
    document.save(DOCX_PATH)

    return {
        "markdown": repo_path(MD_PATH),
        "docx": repo_path(DOCX_PATH),
        "priority_cases": len(priorities),
        "references": len(references),
        "inline_images_expected": len(priorities) + len(references),
        "docx_bytes": DOCX_PATH.stat().st_size,
    }


if __name__ == "__main__":
    print(json.dumps(build(), ensure_ascii=False, indent=2))
