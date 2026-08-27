from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
ATLAS_ROOT = ROOT / "daily" / "20260826_后处理调研"
CORE_PATH = ATLAS_ROOT / "metadata" / "idea_universe" / "core_ideas.jsonl"
PAGE_ROOT = ATLAS_ROOT / "report" / "idea_atlas_pages"
DOCX_PATH = ATLAS_ROOT / "report" / "手机录像后处理_IDEA全量图文图鉴_20260827.docx"


def read_ids() -> list[str]:
    rows = []
    for line in CORE_PATH.read_text(encoding="utf-8").splitlines():
        if line.strip():
            rows.append(json.loads(line)["idea_id"])
    return rows


def verify() -> dict:
    if not DOCX_PATH.exists():
        raise AssertionError(f"missing DOCX: {DOCX_PATH}")
    idea_ids = read_ids()
    pages = sorted(PAGE_ROOT.glob("*.md"))
    document = Document(DOCX_PATH)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    missing_ids = [idea_id for idea_id in idea_ids if idea_id not in paragraph_text]
    duplicate_ids = [idea_id for idea_id in idea_ids if paragraph_text.count(idea_id) != 1]
    heading_counts = {
        name: sum(paragraph.style.name == name for paragraph in document.paragraphs)
        for name in ("Heading 1", "Heading 2", "Heading 3")
    }

    with ZipFile(DOCX_PATH) as archive:
        archive.testzip()
        names = archive.namelist()
        relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        media_files = [name for name in names if name.startswith("word/media/")]
        image_relationships = relationships.count('Target="media/image')

    checks = {
        "core_ideas": len(idea_ids),
        "cluster_pages": len(pages),
        "heading_1": heading_counts["Heading 1"],
        "heading_2": heading_counts["Heading 2"],
        "heading_3": heading_counts["Heading 3"],
        "inline_shapes": len(document.inline_shapes),
        "media_files": len(media_files),
        "image_relationships": image_relationships,
        "missing_idea_ids": len(missing_ids),
        "duplicate_idea_ids": len(duplicate_ids),
        "docx_bytes": DOCX_PATH.stat().st_size,
    }
    assert len(idea_ids) == 1154, checks
    assert len(pages) == 44, checks
    assert heading_counts["Heading 3"] == 1154, checks
    assert len(document.inline_shapes) == 46, checks
    assert len(media_files) == 46, checks
    assert image_relationships == 46, checks
    assert not missing_ids, missing_ids[:10]
    assert not duplicate_ids, duplicate_ids[:10]
    return checks


if __name__ == "__main__":
    print(json.dumps(verify(), ensure_ascii=False, indent=2))
