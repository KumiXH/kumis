"""Build a Chinese-English Markdown and DOCX research report for the ISP video study."""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(r"D:\Repository\ReadPaper\daily\20260826_后处理调研")
REPORT_DIR = ROOT / "report"
REPORT_DIR.mkdir(parents=True, exist_ok=True)
MD_PATH = REPORT_DIR / "手机录像创新功能与ISP技术机会洞察.md"
DOCX_PATH = REPORT_DIR / "手机录像创新功能与ISP技术机会洞察.docx"


def load_jsonl(path: Path) -> list[dict]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


opportunities = load_jsonl(ROOT / "metadata" / "opportunities.jsonl")
deep = load_jsonl(ROOT / "metadata" / "deep_dive_30.jsonl")
priority = load_jsonl(ROOT / "metadata" / "priority_10.jsonl")
manifest = json.loads((ROOT / "sources" / "source_manifest.json").read_text(encoding="utf-8"))
papers = load_jsonl(ROOT / "sources" / "papers" / "paper_records.jsonl")
core_papers = load_jsonl(ROOT / "sources" / "papers" / "core_paper_records.jsonl")
datasets = load_jsonl(ROOT / "sources" / "datasets" / "dataset_records.jsonl")
patents = load_jsonl(ROOT / "sources" / "patents" / "patent_records.jsonl")


def family_counts() -> list[tuple[str, int]]:
    counts: dict[str, int] = {}
    for row in opportunities:
        counts[row["family_zh"]] = counts.get(row["family_zh"], 0) + 1
    return list(counts.items())


def source_by_id(source_id: str) -> dict:
    return next((x for x in manifest if x.get("source_id") == source_id), {})


def evidence_name(source_id: str) -> str:
    src = source_by_id(source_id)
    if not src:
        return source_id
    return f"{src.get('title', source_id)} | {src.get('publisher_or_authors', '')} | {src.get('evidence_level', '')}"


def md_report() -> str:
    lines = [
        "# 手机录像创新功能与 ISP 后处理技术机会洞察",
        "",
        f"> 研究日期：{date.today()}  | 研究目录：`daily/20260826_后处理调研`",
        "> 研究对象：手机录像、ISP 后处理、计算摄影、相机玩法、端侧视频算法和生成式视频编辑",
        "",
        "## 摘要",
        "",
        "本报告面向 ISP 视频和手机影像预研，采用“产品创意池 + 技术可实现性评估 + 证据图谱”的组织方式。研究不把录像能力局限为传统降噪、锐化和防抖，而是将专业相机、电影机、运动相机、无人机、短视频 App、后期软件、学术论文和公开专利中的能力，重新映射到手机录像链路。",
        "",
        f"本阶段建立 {len(opportunities)} 条机会记录，覆盖 {len(family_counts())} 个技术能力族；完成 {len(deep)} 张技术深度卡和 {len(priority)} 个组合创新概念；来源清单包含 {len(manifest)} 条产品/论文/本地资料，筛出 {len(core_papers)} 篇核心论文、{len(datasets)} 个数据集，并保留 OpenAlex 初筛论文元数据 {len(papers)} 条。机会记录不是产品承诺，其中主体为 E5（本报告推演），外部证据与视频化推演在表格中分栏。",
        "",
        "## 1. 研究口径与证据边界",
        "",
        "### 1.1 录像优先",
        "",
        "拍照能力只有在进一步说明连续帧、运动、曝光、遮挡、功耗和交互之后，才进入录像机会图谱。实时预览、录制在线处理、录后设备处理和云端重处理分别标注，不把照片功能直接改名为视频功能。",
        "",
        "### 1.2 证据等级",
        "",
        """- **E1 已量产**：官方产品页、说明书或支持文档明确描述。
- **E2 公开演示/有限发布**：官方白皮书、SDK、Demo 或发布材料。
- **E3 学术原型**：论文、补充材料、开源代码或实验报告。
- **E4 专利储备**：公开专利文本，仅证明提出过方案。
- **E5 本报告推演**：基于已有证据提出的录像化创新，不能当作行业事实。""",
        "",
        "### 1.3 真实性边界",
        "",
        """- **Faithful enhancement 忠实增强**：尽量恢复真实内容，例如夜景降噪、去模糊、滚动快门校正和跨摄色彩连续。
- **Perceptual enhancement 感知增强**：允许改变质感和风格，例如胶片颗粒、虚拟镜头、局部 LUT 和受控人像增强。
- **Generative creation 生成式创作**：允许补全、替换或重构，但必须保存原片、生成 mask、编辑参数和可回退版本。""",
        "",
        "## 2. 技术机会总览",
        "",
        "本阶段把机会拆为 14 个能力族，每个能力族先形成 8 条候选记录。",
        "",
        """| 能力族 | 候选数 | 录像落点 |
|---|---:|---|
""" + "\n".join(f"| {name} | {count} | 实时预览、在线录像或录后处理，详见 Opportunity_Map |" for name, count in family_counts()),
        "",
        "### 2.1 三个最值得持续投入的主线",
        "",
        """1. **视频真实性增强主线**：夜景 HDR、运动复原、滚动快门、跨摄连续、局部人像增强。这些方向产品价值稳定，难点是时序和系统协同，适合作为 ISP/NPU 联合优化对象。
2. **电影化计算摄影主线**：动态星芒、虚拟景深、虚拟镜头、语义 LUT、虚拟运镜。它们更容易形成用户可感知差异，但需要把效果参数化，不能依赖逐帧贴图。
3. **可信生成式录像主线**：天空/天气/画外延展/局部背景编辑。它们创新性最高，但必须把“生成了什么”纳入产品交互和文件元数据。""",
        "",
        "## 3. 从相机与软件玩法迁移到手机录像",
        "",
        "### 3.1 手机原生影像已经验证的能力",
        "",
        "Apple Cinematic mode 证明了焦点转移和视频景深可以成为手机原生体验；Google Pixel Cinematic Blur 和 Video Boost 分别代表视频景深和云端高质量视频处理；Samsung Super Steady 代表手机级稳定；DJI、GoPro、Insta360 则把稳定、地平线锁定、主体跟踪和自动构图推到运动影像产品中。Blackmagic Camera、DaVinci Resolve、After Effects 和 CapCut 说明，移动拍摄和后期编辑之间正在形成连续工作流。",
        "",
        "### 3.2 迁移的关键不是功能名称，而是时序状态",
        "",
        "相机和后期软件中很多功能在手机录像化时都需要新增状态：对象 track ID、关键帧、光源 ID、曝光轨迹、镜头状态、深度层、生成 mask、置信度和回退信息。没有这些状态，模型容易出现 flicker、texture crawl、identity drift、mask leakage 和 exposure breathing。",
        "",
        "## 4. 端到端录像后处理系统框架",
        "",
        """```text
Sensor / RAW / Multi-camera / Depth / IMU / Audio
                    |
        Front-end ISP and synchronization
                    |
     Detection: subject / light / motion / depth / quality
                    |
       Temporal memory: track / keyframe / flow / state
                    |
 Restoration / relighting / rendering / generative editing
                    |
    Temporal consistency + confidence + rollback metadata
                    |
        Color management / encoder / preview / export
```""",
        "",
        "### 4.1 实时与录后不是二选一",
        "",
        "建议采用分级链路：实时预览只做低分辨率分析、粗粒度效果和质量提示；录像在线处理使用因果模型和有限参考帧；录后处理保留 proxy、关键帧、深度、mask、IMU 和曝光状态，再使用双向模型或生成式模型重算。FlashVSR、视频 VAE 和 DiT/FLUX 资料说明，latent 表征、特征传播和蒸馏可以把复杂模型压缩到更接近视频应用的推理形态，但具体手机 FPS、内存和功耗仍必须在目标 SoC 上实测。",
        "",
        "## 5. 核心论文与数据集路线",
        "",
        "### 5.1 核心论文",
        "",
        "下表不是全量综述，而是本阶段用于支撑视频后处理功能设计的核心入口。`Paper_Discovery_381` 中的宽搜结果只作为发现池，不能替代论文正文核验。",
        "",
        "| 论文 | 年份/会议 | 方向 | 对手机录像的价值 |\n|---|---|---|---|\n" + "\n".join(f"| {row['title']} | {row['year']} / {row['venue']} | {row['family']} | {row['why_relevant']} |" for row in core_papers),
        "",
        "推荐先读 `VRT` 和 `FlashVSR` 理解视频复原、特征传播和低步数推理；再读实时背景抠图、视频抠图与 `MODNet` 理解人像/环境 mask 的时序问题；随后读可变光圈散景、人像视频数据和视频 VAE/DiT，建立电影化效果与生成式编辑的基础。",
        "",
        "### 5.2 数据集与数据制作入口",
        "",
        "| 数据集 | 用途 | 访问/许可边界 |\n|---|---|---|\n" + "\n".join(f"| {row['name']} | {row['task']} | {row['license_or_access']} |" for row in datasets),
        "",
        "数据集不能简单混合使用。静态人脸数据适合身份和细节预训练，视频人脸数据适合时序与身份稳定，真实退化/挑战数据适合建立评测集。对夜景、光效、多摄、IMU、曝光和镜头切换，本项目仍需要自行采集手机连续视频，并保留 sensor/ISP metadata。",
        "",
        "### 5.3 专利检索状态",
        "",
        f"当前仅登记 {len(patents)} 条待核验专利检索主题，尚未获得可稳定引用的公开专利号和正文，因此不将其写成 E4 已核验事实。后续应按申请人、公开日、权利要求和法律状态逐条补齐。",
        "",
        "## 6. 30 个技术深度方向",
        "",
    ]
    for row in deep:
        lines += [
            f"### {row['deep_dive_id']} {row['title']} / {row['english']}",
            f"- **能力族 / 模式 / 边界**：{row['family']} / {row['mode']} / {row['truth']}",
            f"- **来源**：" + "; ".join(evidence_name(x) for x in row["source_ids"]),
            f"- **研究问题**：{row['problem']}",
            f"- **技术方案**：{row['solution']}",
            f"- **输入信号**：{row['inputs']}",
            f"- **模型链路**：{row['model']}",
            f"- **训练数据**：{row['training']}",
            f"- **LOSS / objective**：{row['loss']}",
            f"- **时序策略**：{row['temporal']}",
            f"- **端侧落点**：{row['edge']}",
            f"- **风险**：{row['risks']}",
            f"- **MVP**：{row['mvp']}",
            "",
        ]
    lines += [
        "## 7. 10 个组合创新概念",
        "",
        "组合创新不是现有产品清单，而是将已经存在的产品/论文能力重新组织为手机录像预研方向。",
        "",
    ]
    for item in priority:
        lines += [
            f"### {item['concept_id']} {item['name']} / {item['english']}",
            f"- **组成**：{', '.join(item['components'])}",
            f"- **用户故事**：{item['user_story']}",
            f"- **交互**：{item['interaction']}",
            f"- **系统链路**：{item['pipeline']}",
            f"- **MVP**：{item['mvp']}",
            f"- **数据**：{item['data']}",
            f"- **指标**：{item['metrics']}",
            f"- **风险**：{item['risk']}",
            f"- **真实性边界**：{item['truth_boundary']}",
            "",
        ]
    lines += [
        "## 8. 训练数据与 LOSS 设计原则",
        "",
        "### 8.1 数据制作",
        "",
        """1. **真实连续视频优先**：训练时必须保留真实运动、曝光变化、压缩、滚动快门和镜头切换，不要只把独立图片随机拼成视频。
2. **退化链路可解释**：将 sensor noise、ISO、motion blur、compression、downsampling、color shift、lens flare 和遮挡按场景组合，记录每个退化参数。
3. **多源监督并存**：恢复类使用 GT/LQ 配对；人像使用 parsing、landmark、identity embedding；视频编辑使用 mask/depth/flow；生成式方案使用 protected/editable region。
4. **难例要单独建集**：夜景高光、玻璃反射、眼镜、发丝、多人遮挡、快速运动和镜头切换不能只混在随机训练集中，必须有独立评测集。""",
        "",
        "### 8.2 LOSS 组合",
        "",
        "建议按功能而不是按网络结构组织损失：",
        "",
        """```text
L_total = L_reconstruction
        + lambda_p L_perceptual
        + lambda_t L_temporal_warp
        + lambda_i L_identity_or_semantic
        + lambda_b L_boundary_or_mask
        + lambda_c L_color_or_exposure
        + lambda_g L_generation_control
        + lambda_e L_edge_latency_power
```""",
        "",
        "损失权重只能作为起始假设。实际工程中要检查每个 loss 的梯度量级、不同区域的有效像素数、ROI 面积归一化和 ablation 结果。对于人像身份，identity loss 应当约束结构稳定而不是把所有帧强行拉向同一个表情；对于生成式编辑，protected-region loss 应优先于视觉风格损失。",
        "",
        "## 9. 端侧可实现性与系统建议",
        "",
        """- **ISP 前端**适合做 RAW/YUV 预处理、曝光/颜色状态、镜头切换和低延迟几何校正。
- **DSP**适合持续运行 IMU、音频事件、轻量质量检测和环形缓存。
- **NPU/GPU**适合人像 ROI、光流、深度、视频复原和局部渲染。
- **VPU/编码器前**适合将中间状态与 proxy 视频配套保存。
- **录后任务**适合双向时序模型、FLUX/扩散教师蒸馏、局部生成和高质量超分。

建议所有复杂功能提供至少三档：preview、online recording、offline rerender。温度感知调度器必须使用 hysteresis，避免模型在相邻帧之间频繁升降级。""",
        "",
        "## 10. 主要研究空白",
        "",
        """1. **生成式视频效果的长期身份和结构稳定**：单帧效果已经很多，但手机真实长视频的 track、遮挡和镜头切换仍缺少统一方案。
2. **光照效果的物理可解释控制**：动态打光、星芒、halation 和 flare 需要参数化、可回退、可跨镜头连续。
3. **多摄与生成模型协同**：多摄提供深度和真实细节，生成模型提供补全，但两者的置信度和视差融合还不成熟。
4. **端侧质量-功耗-时序联合优化**：论文常报告画质或速度的一面，真实手机需要同时优化热、存储、后台任务、编码和用户等待。
5. **可信生成式录像文件标准**：生成区域、源帧 hash、编辑参数和可回退版本应进入产品级元数据，而不只是 UI 提示。""",
        "",
        "## 11. 建议的后续研究顺序",
        "",
        "第一阶段建议先做无生成或弱生成的高价值方向：跨摄连续、夜景 HDR、滚动快门与 EIS 联合、视频人像身份锚定、语义 LUT 和数字变焦 ROI 超分。第二阶段再做动态光学、录后景深、语义长曝光和旅行净景。第三阶段进入有限视角扩展、天空/天气编辑和可信视频扩散。每个阶段都保留原片、代理、mask、深度、运动和评测日志，形成可持续的数据闭环。",
        "",
        "## 附录：本地资料与交付物",
        "",
        "- Excel：`matrix/手机录像创新功能机会库.xlsx`",
        "- 全量机会：`metadata/opportunities.jsonl`",
        "- 30 张技术卡：`metadata/deep_dive_30.jsonl`",
        "- 10 个组合概念：`metadata/priority_10.jsonl`",
        "- 官方与本地来源：`sources/source_manifest.json`",
        "- OpenAlex 原始查询：`sources/papers/openalex_raw/`",
        "",
        "报告中凡是“建议”“可以”“应当”的句子均属于研究判断或预研假设；产品、论文和专利事实以来源清单和证据等级为准。",
    ]
    return "\n".join(lines) + "\n"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "1F2933", size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color="FFFFFF", size=8)
        set_cell_shading(cell, "0B3A53")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8)
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[i], "F2F6F8")
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [("Title", 22, "0B3A53"), ("Heading 1", 16, "0B3A53"), ("Heading 2", 13, "136F8A"), ("Heading 3", 11, "1F637A")]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("手机录像创新功能与 ISP 后处理技术机会洞察")
    p = doc.add_paragraph()
    p.add_run("Mobile Video Innovation and ISP Post-Processing Opportunity Map\n").bold = True
    p.add_run(f"研究日期：{date.today()}\n研究目录：daily/20260826_后处理调研\n面向：ISP 视频、手机影像与端侧算法预研")
    doc.add_paragraph("本报告是个人研究手册，不是产品发布材料。外部证据、论文结论和本报告推演严格分开。")

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(f"本阶段建立 {len(opportunities)} 条机会记录，覆盖 14 个能力族；完成 {len(deep)} 个技术深度方向和 {len(priority)} 个组合创新概念。研究将手机、相机、运动相机、无人机、短视频应用、专业后期软件和学术论文中的能力，统一映射到手机录像的实时预览、在线录像、录后设备处理和云端重处理链路。")
    add_table(doc, ["指标", "数量", "说明"], [["候选机会", str(len(opportunities)), "14 个能力族，每族 8 条"], ["技术深挖", str(len(deep)), "输入、模型、数据、LOSS、时序、端侧和风险"], ["组合创新", str(len(priority)), "面向预研立项的组合方案"], ["核心论文", str(len(core_papers)), "已筛选的主要阅读入口"], ["数据集", str(len(datasets)), "已有本地来源与访问边界"], ["论文发现池", str(len(papers)), "OpenAlex 初筛元数据，需继续正文核验"]], [1.5, 0.8, 4.7])

    doc.add_heading("1. 研究口径与证据边界", level=1)
    doc.add_heading("1.1 录像优先", level=2)
    doc.add_paragraph("拍照能力只有在说明连续帧、运动、曝光、遮挡、功耗和交互之后，才进入录像机会图谱。实时预览、录像在线处理、录后设备处理和云端重处理分别标注，不把照片功能直接改名为视频功能。")
    doc.add_heading("1.2 证据等级", level=2)
    for text in ["E1 已量产：官方产品页、说明书或支持文档明确描述。", "E2 公开演示/有限发布：官方白皮书、SDK、Demo 或发布材料。", "E3 学术原型：论文、补充材料、开源代码或实验报告。", "E4 专利储备：公开专利文本，仅证明提出过方案。", "E5 本报告推演：基于已有证据提出的录像化创新，不是行业事实。"]:
        add_bullet(doc, text)
    doc.add_heading("1.3 真实性边界", level=2)
    add_table(doc, ["边界", "含义", "典型方向"], [["Faithful", "尽量恢复真实内容", "夜景降噪、滚动快门、跨摄连续"], ["Perceptual", "受控改变质感/风格", "胶片颗粒、虚拟镜头、人像增强"], ["Generative", "补全、替换或重构内容", "天空、天气、画外延展、局部编辑"]], [1.2, 4.2, 1.6])

    doc.add_heading("2. 技术机会总览", level=1)
    add_table(doc, ["能力族", "候选数", "主要录像价值"], [[name, str(count), "见 Opportunity_Map；需按实时/录后继续拆分"] for name, count in family_counts()], [2.4, 0.8, 3.8])
    doc.add_heading("2.1 三条主线", level=2)
    for text in ["视频真实性增强：夜景 HDR、运动复原、滚动快门、跨摄连续和人像身份锚定，产品价值稳定，重点是时序与系统协同。", "电影化计算摄影：动态星芒、虚拟景深、虚拟镜头、语义 LUT 和虚拟运镜，差异化强，但必须做参数化和跨帧状态。", "可信生成式录像：天空、天气、画外延展和局部背景编辑，创新性最高，但必须呈现生成边界并保留原片。"]:
        add_bullet(doc, text)

    doc.add_heading("3. 录像后处理系统框架", level=1)
    doc.add_paragraph("Sensor / RAW / Multi-camera / Depth / IMU / Audio\n↓\nFront-end ISP and synchronization\n↓\nSubject / light / motion / depth / quality analysis\n↓\nTemporal memory: track / keyframe / flow / state\n↓\nRestoration / relighting / rendering / generative editing\n↓\nTemporal consistency + confidence + rollback metadata\n↓\nColor management / encoder / preview / export")
    doc.add_paragraph("建议采用分级链路：实时预览做低分辨率分析和粗粒度效果；在线录像使用因果模型和有限参考帧；录后保留 proxy、关键帧、深度、mask、IMU 和曝光状态，再使用双向模型或生成式模型重算。FlashVSR、视频 VAE、DiT/FLUX 资料可作为时序 latent、特征传播和教师蒸馏的技术入口，但手机 FPS、内存和功耗必须在目标 SoC 上实测。")

    doc.add_heading("4. 核心论文与数据集路线", level=1)
    doc.add_heading("4.1 核心论文", level=2)
    add_table(doc, ["论文", "年份/会议", "方向", "对手机录像的价值"], [[row["title"], f"{row['year']} / {row['venue']}", row["family"], row["why_relevant"]] for row in core_papers], [2.2, 1.2, 1.2, 2.4])
    doc.add_paragraph("推荐先读 VRT 和 FlashVSR 理解视频复原、时序对齐和低步数推理；再读实时背景抠图、视频抠图与 MODNet 理解人像/环境 mask；随后阅读可变光圈散景、人像视频数据和视频 VAE/DiT，建立电影化效果与生成式编辑基础。")
    doc.add_heading("4.2 数据集", level=2)
    add_table(doc, ["数据集", "用途", "访问/许可边界"], [[row["name"], row["task"], row["license_or_access"]] for row in datasets], [1.5, 2.8, 2.7])
    doc.add_paragraph("静态人脸数据适合身份和细节预训练，视频人脸数据适合时序和身份稳定，真实退化/挑战数据适合评测。夜景、光效、多摄、IMU、曝光和镜头切换仍需要自采连续手机视频，并保留 sensor/ISP metadata。")
    doc.add_heading("4.3 专利检索状态", level=2)
    doc.add_paragraph(f"当前登记 {len(patents)} 条待核验专利检索主题，尚未获得可稳定引用的公开专利号和正文，因此不写成 E4 已核验事实。")

    doc.add_heading("5. 30 个技术深度方向", level=1)
    for row in deep:
        doc.add_heading(f"{row['deep_dive_id']} {row['title']}", level=2)
        doc.add_paragraph(f"{row['english']} | {row['family']} | {row['mode']} | {row['truth']}")
        add_table(doc, ["项目", "分析"], [["证据", "; ".join(evidence_name(x) for x in row["source_ids"])], ["研究问题", row["problem"]], ["技术方案", row["solution"]], ["输入信号", row["inputs"]], ["模型链路", row["model"]], ["训练数据", row["training"]], ["LOSS / objective", row["loss"]], ["时序策略", row["temporal"]], ["端侧落点", row["edge"]], ["风险", row["risks"]], ["MVP", row["mvp"]]], [1.35, 5.65])

    doc.add_heading("6. 10 个组合创新概念", level=1)
    for item in priority:
        doc.add_heading(f"{item['concept_id']} {item['name']}", level=2)
        doc.add_paragraph(item["english"])
        add_table(doc, ["项目", "分析"], [["组成", ", ".join(item["components"])], ["用户故事", item["user_story"]], ["交互", item["interaction"]], ["系统链路", item["pipeline"]], ["MVP", item["mvp"]], ["数据", item["data"]], ["指标", item["metrics"]], ["风险", item["risk"]], ["真实性边界", item["truth_boundary"]]], [1.35, 5.65])

    doc.add_heading("7. 训练数据与 LOSS 设计", level=1)
    doc.add_heading("7.1 数据制作", level=2)
    for text in ["真实连续视频优先：保留真实运动、曝光变化、压缩、滚动快门和镜头切换。", "退化链路可解释：记录 sensor noise、ISO、motion blur、compression、downsampling、color shift、flare 和遮挡参数。", "多源监督并存：恢复类使用 GT/LQ；人像使用 parsing/landmark/identity；视频编辑使用 mask/depth/flow；生成式方案使用 protected/editable region。", "难例独立建集：夜景高光、玻璃反射、眼镜、发丝、多人遮挡、快速运动和镜头切换应单独评测。"]:
        add_bullet(doc, text)
    doc.add_heading("7.2 LOSS 组合", level=2)
    doc.add_paragraph("L_total = L_reconstruction + λp L_perceptual + λt L_temporal_warp + λi L_identity/semantic + λb L_boundary/mask + λc L_color/exposure + λg L_generation_control + λe L_edge_latency_power")
    doc.add_paragraph("损失权重只能作为起始假设。工程上需要检查每个 loss 的梯度量级、ROI 面积归一化、区域有效像素数和 ablation。人像 identity loss 应约束结构稳定而不是把帧强行拉向同一表情；生成式编辑应让 protected-region loss 优先于风格损失。")

    doc.add_heading("8. 端侧可实现性与研究空白", level=1)
    add_table(doc, ["模块", "适合承担的职责"], [["ISP", "RAW/YUV、曝光/颜色状态、切镜和低延迟几何校正"], ["DSP", "IMU、音频事件、质量检测、环形缓存"], ["NPU/GPU", "人像 ROI、光流、深度、视频复原、局部渲染"], ["编码器前/VPU", "proxy、metadata、中间 latent 配套存储"], ["录后任务", "双向时序、FLUX/扩散教师蒸馏、局部生成、高质量超分"]], [1.5, 5.5])
    doc.add_paragraph("主要研究空白包括：生成式视频效果的长期身份和结构稳定、光照效果的物理可解释控制、多摄与生成模型的置信度融合、质量-功耗-时序联合优化，以及生成区域/源帧 hash/编辑参数的可信视频元数据。")

    doc.add_heading("9. 后续研究顺序", level=1)
    doc.add_paragraph("第一阶段先做无生成或弱生成的高价值方向：跨摄连续、夜景 HDR、滚动快门与 EIS 联合、视频人像身份锚定、语义 LUT 和数字变焦 ROI 超分。第二阶段做动态光学、录后景深、语义长曝光和旅行净景。第三阶段进入有限视角扩展、天空/天气编辑和可信视频扩散。每个阶段都保留原片、代理、mask、深度、运动和评测日志，形成可持续数据闭环。")

    doc.add_heading("附录：本地资料", level=1)
    for text in ["Excel：matrix/手机录像创新功能机会库.xlsx", "全量机会：metadata/opportunities.jsonl", "30 张技术卡：metadata/deep_dive_30.jsonl", "10 个组合概念：metadata/priority_10.jsonl", "来源清单：sources/source_manifest.json", "OpenAlex 原始查询：sources/papers/openalex_raw/"]:
        add_bullet(doc, text)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("ISP Video Innovation Research | ReadPaper")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(8)
    doc.save(DOCX_PATH)


MD_PATH.write_text(md_report(), encoding="utf-8")
build_docx()
print(json.dumps({"markdown": str(MD_PATH), "docx": str(DOCX_PATH), "opportunities": len(opportunities), "deep": len(deep), "priority": len(priority)}, ensure_ascii=False, indent=2))
