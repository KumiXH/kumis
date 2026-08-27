from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
ATLAS_ROOT = ROOT / "daily" / "20260826_后处理调研"
REPORT_ROOT = ATLAS_ROOT / "report"
PAGE_ROOT = REPORT_ROOT / "idea_atlas_pages"
OUTPUT = REPORT_ROOT / "手机录像后处理_IDEA全量图文图鉴_20260827.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2933"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft YaHei"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 italic: bool | None = None, color: str | None = None,
                 latin: str = FONT_LATIN, cjk: str = FONT_CJK) -> None:
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), cjk)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_hyperlink(paragraph, text: str, target: str, internal: bool = False,
                  bold: bool = False, italic: bool = False,
                  color: str = BLUE) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    if internal:
        hyperlink.set(qn("w:anchor"), target)
    else:
        relation_id = paragraph.part.relate_to(
            target,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    rpr.append(fonts)
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    rpr.append(run_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    if bold:
        rpr.append(OxmlElement("w:b"))
    if italic:
        rpr.append(OxmlElement("w:i"))
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "19")
    rpr.append(size)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\)|\*[^*]+?\*)")


def resolve_link(target: str, source_path: Path) -> str:
    if target.startswith(("http://", "https://", "mailto:")):
        return target
    local_path = (source_path.parent / target).resolve()
    return local_path.as_uri()


def add_inline_markdown(paragraph, text: str, source_path: Path,
                        default_size: float = 9.5,
                        default_color: str = INK) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, default_size, color=default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, default_size, bold=True, color=default_color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, default_size, color=DARK_BLUE, latin="Consolas", cjk=FONT_CJK)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), LIGHT_GRAY)
            run._r.get_or_add_rPr().append(shading)
        elif token.startswith("["):
            label, target = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, resolve_link(target, source_path))
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, default_size, italic=True, color=MUTED)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, default_size, color=default_color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.1

    styles = {
        "Title": (28, INK, 0, 8),
        "Subtitle": (14, MUTED, 0, 12),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (10.5, DARK_BLUE, 7, 2),
        "Caption": (8.5, MUTED, 2, 7),
    }
    for name, (size, color, before, after) in styles.items():
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.bold = True
    doc.styles["Caption"].font.italic = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(9.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(1)
        style.paragraph_format.line_spacing = 1.08

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("手机录像后处理 IDEA 全量图文图鉴")
    set_run_font(run, 8.5, color=MUTED)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D9E1E8")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("2026-08-27  |  ")
    set_run_font(run, 8.5, color=MUTED)
    add_page_field(fp)


def add_cover(doc: Document, pages: list[Path]) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(58)
    p.add_run("手机录像后处理\nIDEA 全量图文图鉴")

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("移动影像后处理、计算摄影与生成式录像创意参考手册")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("1,154 条基础 IDEA  ·  44 个创意簇  ·  44 张概念图")
    set_run_font(run, 11, bold=True, color=BLUE)

    collage = build_collage(pages)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.add_run().add_picture(collage, width=Inches(5.8))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("个人研究与创意探索资料  |  2026-08-27")
    set_run_font(run, 9.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("图片均为概念示意，不是论文原图、真实产品截图或实测结果。")
    set_run_font(run, 8.5, italic=True, color=MUTED)
    doc.add_page_break()


def build_collage(pages: list[Path]) -> io.BytesIO:
    page_images = []
    for page_path in pages[:12]:
        text = page_path.read_text(encoding="utf-8")
        match = re.search(r"!\[[^\]]*\]\(([^)]+)\)", text)
        if not match:
            continue
        image_path = (page_path.parent / match.group(1)).resolve()
        image = Image.open(image_path).convert("RGB")
        image.thumbnail((300, 300))
        page_images.append(image.copy())
    canvas = Image.new("RGB", (1200, 900), f"#{WHITE}")
    draw = ImageDraw.Draw(canvas)
    for index, image in enumerate(page_images):
        row, col = divmod(index, 4)
        x = col * 300
        y = row * 300
        canvas.paste(image.resize((288, 288)), (x + 6, y + 6))
        draw.rectangle((x + 5, y + 5, x + 294, y + 294), outline="#D9E1E8", width=2)
    output = io.BytesIO()
    canvas.save(output, format="JPEG", quality=86, optimize=True)
    output.seek(0)
    return output


def add_callout(doc: Document, text: str, source_path: Path) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_inline_markdown(p, text, source_path, default_size=9.5, default_color=DARK_BLUE)


def add_overview(doc: Document, pages: list[Path]) -> None:
    doc.add_heading("使用说明与边界", level=1)
    intro = (
        "这是一份个人阅读用的概念图鉴。它把手机录像后处理创意与视觉示意放在一起，"
        "帮助读者先理解“用户会看到什么”，再回看算法、输入信号、数据和风险。"
    )
    p = doc.add_paragraph()
    add_inline_markdown(p, intro, REPORT_ROOT / "手机录像后处理_IDEA图文图鉴_20260827.md", 10)
    boundaries = [
        "全量基础 IDEA：1,154 条；创意簇：44 个。",
        "每个创意簇配置一张概念示意图，并保留该簇全部基础 IDEA。",
        "图像由统一的概念视觉母图裁切得到，不是论文原图，也不是厂商真实产品截图。",
        "旧 112 条机会的来源证据与新创意的概念推演分开表达；新增创意统一标记为 idea_only。",
        "实时预览、录制在线、录后端侧、云端、30/60 fps、ROI 等属于实现变体；相关数据索引保留在文末。",
    ]
    for item in boundaries:
        p = doc.add_paragraph(style="List Bullet")
        add_inline_markdown(p, item, REPORT_ROOT / "手机录像后处理_IDEA图文图鉴_20260827.md")

    doc.add_heading("总体图谱", level=1)
    flow = build_flow_diagram()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(flow, width=Inches(6.0))
    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run("图 0：从手机录像输入到可重算录像资产的总体能力图谱。")

    doc.add_heading("章节目录", level=1)
    p = doc.add_paragraph()
    add_inline_markdown(
        p,
        "以下目录可点击跳转到对应章节。Word 的导航窗格也可按标题浏览完整结构。",
        REPORT_ROOT / "手机录像后处理_IDEA图文图鉴_20260827.md",
        9.5,
    )
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [720, 3540, 5100])
    headers = ("序号", "创意簇", "内容概述")
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        run = p.add_run(value)
        set_run_font(run, 9.5, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for index, page_path in enumerate(pages, 1):
        title, summary, count = page_metadata(page_path)
        cells = table.add_row().cells
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{index:02d}")
        set_run_font(run, 9)
        p = cells[1].paragraphs[0]
        add_hyperlink(p, title, f"cluster_{index:02d}", internal=True, bold=True)
        p = cells[2].paragraphs[0]
        add_inline_markdown(p, f"{summary}（{count} 条 IDEA）", page_path, 8.8)
    doc.add_page_break()


def build_flow_diagram() -> io.BytesIO:
    canvas = Image.new("RGB", (1800, 780), f"#{WHITE}")
    draw = ImageDraw.Draw(canvas)
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    font_bold_path = Path("C:/Windows/Fonts/msyhbd.ttc")
    font = ImageFont.truetype(str(font_path), 32) if font_path.exists() else ImageFont.load_default()
    font_bold = ImageFont.truetype(str(font_bold_path), 34) if font_bold_path.exists() else font

    boxes = {
        "input": (60, 300, 320, 450, "手机录像输入"),
        "state": (420, 300, 760, 450, "时序状态与\n语义对象"),
        "restore": (900, 40, 1260, 160, "恢复与增强"),
        "optics": (900, 220, 1260, 340, "光学与时间重构"),
        "gen": (900, 400, 1260, 520, "生成式编辑与\n事实保护"),
        "delivery": (900, 580, 1260, 700, "编码、功耗与交付"),
        "asset": (1440, 300, 1740, 450, "可重算录像资产"),
    }
    for key, (x1, y1, x2, y2, label) in boxes.items():
        fill = "#E8EEF5" if key not in ("input", "asset") else "#DCEAF7"
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline="#2E74B5", width=4)
        bbox = draw.multiline_textbbox((0, 0), label, font=font_bold, spacing=8, align="center")
        tx = (x1 + x2 - (bbox[2] - bbox[0])) / 2
        ty = (y1 + y2 - (bbox[3] - bbox[1])) / 2
        draw.multiline_text((tx, ty), label, font=font_bold, fill="#1F4D78", spacing=8, align="center")

    def arrow(start: tuple[int, int], end: tuple[int, int]) -> None:
        draw.line((start, end), fill="#667788", width=6)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) > abs(ey - sy):
            points = [(ex, ey), (ex - 20, ey - 12), (ex - 20, ey + 12)]
        else:
            points = [(ex, ey), (ex - 12, ey - 20), (ex + 12, ey - 20)]
        draw.polygon(points, fill="#667788")

    arrow((320, 375), (420, 375))
    for y in (100, 280, 460, 640):
        draw.line((760, 375, 820, 375, 820, y, 900, y), fill="#667788", width=6)
        arrow((820, y), (900, y))
    for y in (100, 280, 460, 640):
        draw.line((1260, y, 1350, y, 1350, 375, 1440, 375), fill="#667788", width=6)
    arrow((1350, 375), (1440, 375))
    output = io.BytesIO()
    canvas.save(output, format="PNG", optimize=True)
    output.seek(0)
    return output


def page_metadata(page_path: Path) -> tuple[str, str, int]:
    text = page_path.read_text(encoding="utf-8")
    title = re.search(r"^#\s+\d+\.\s+(.+)$", text, re.MULTILINE).group(1)
    reading = re.search(r"## 看图理解\s+(.+?)\s+本簇收录", text, re.DOTALL)
    summary = re.sub(r"\s+", " ", reading.group(1)).strip() if reading else "详见本章图文说明"
    count = len(re.findall(r"^###\s+", text, re.MULTILINE))
    return title, summary, count


def compressed_image(path: Path, max_pixels: int = 1200, quality: int = 88) -> io.BytesIO:
    image = Image.open(path).convert("RGB")
    image.thumbnail((max_pixels, max_pixels))
    output = io.BytesIO()
    image.save(output, format="JPEG", quality=quality, optimize=True)
    output.seek(0)
    return output


def add_cluster_page(doc: Document, page_path: Path, index: int, bookmark_id: int) -> int:
    lines = page_path.read_text(encoding="utf-8").splitlines()
    title_match = re.match(r"^#\s+(.+)$", lines[0])
    title = title_match.group(1) if title_match else page_path.stem
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.page_break_before = True
    add_bookmark(heading, f"cluster_{index:02d}", bookmark_id)
    bookmark_id += 1

    blank_pending = False
    for line in lines[1:]:
        stripped = line.strip()
        if not stripped:
            blank_pending = True
            continue
        if stripped.startswith("[返回图文图鉴总览]"):
            continue
        if stripped.startswith("> "):
            add_callout(doc, stripped[2:], page_path)
        elif stripped.startswith("!["):
            match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
            if match:
                alt, rel = match.groups()
                image_path = (page_path.parent / rel).resolve()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(5)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run()
                inline = run.add_picture(compressed_image(image_path), width=Inches(5.15))
                doc_pr = inline._inline.docPr
                doc_pr.set("descr", alt)
        elif stripped.startswith("*图 ") and stripped.endswith("*"):
            p = doc.add_paragraph(style="Caption")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_markdown(p, stripped, page_path, 8.5, MUTED)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            p.paragraph_format.keep_with_next = True
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.keep_together = True
            add_inline_markdown(p, stripped[2:], page_path, 9.2)
        else:
            p = doc.add_paragraph()
            if blank_pending:
                p.paragraph_format.space_before = Pt(1)
            add_inline_markdown(p, stripped, page_path, 9.5)
        blank_pending = False
    return bookmark_id


def add_appendix(doc: Document) -> None:
    doc.add_heading("数据、索引与图片说明", level=1)
    source = REPORT_ROOT / "手机录像后处理_IDEA图文图鉴_20260827.md"
    items = [
        ("基础 IDEA 全量 JSONL", ATLAS_ROOT / "metadata" / "idea_universe" / "core_ideas.jsonl"),
        ("单轴变体全量 JSONL", ATLAS_ROOT / "metadata" / "idea_universe" / "idea_variants.jsonl"),
        ("基础 IDEA 全量纯文本报告", REPORT_ROOT / "手机录像后处理_IDEA全量宇宙_20260827.md"),
        ("实现变体全量纯文本报告", REPORT_ROOT / "手机录像后处理_IDEA变体全量_20260827.md"),
        ("Excel 全量数据库", ATLAS_ROOT / "matrix" / "手机录像后处理_IDEA全量宇宙_20260827.xlsx"),
        ("视觉资产目录", ATLAS_ROOT / "figures" / "idea_atlas"),
    ]
    for label, path in items:
        p = doc.add_paragraph(style="List Bullet")
        if path.exists():
            add_hyperlink(p, label, path.resolve().as_uri())
        else:
            add_inline_markdown(p, f"{label}（当前路径不存在）", source)

    doc.add_heading("图片说明", level=2)
    paragraphs = [
        "本图鉴中的图像用于建立视觉直觉：动态星芒、虚拟打光、对象级快门、多摄切镜、声音驱动、语义编码等功能分别对应用户可观察的画面或交互变化。涉及编解码、功耗、可信生成和多光谱的图像使用可视化界面或分层表达，因为这些能力不能仅靠普通摄影画面直接呈现。",
        "每张图片都应与章节中的 idea_id、来源状态、输入信号、真实性边界和风险一起阅读；不能把概念图解读为已经实现的性能、量产能力或论文实验结果。",
    ]
    for text in paragraphs:
        p = doc.add_paragraph()
        add_inline_markdown(p, text, source, 10)


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "手机录像后处理 IDEA 全量图文图鉴"
    props.subject = "移动影像后处理、计算摄影与生成式录像创意参考手册"
    props.author = "ReadPaper Research Workspace"
    props.keywords = "手机录像, ISP, 计算摄影, 视频后处理, 生成式视频, 创意图鉴"
    props.comments = "包含 1,154 条基础 IDEA、44 个创意簇和 44 张概念示意图。"


def build() -> Path:
    pages = sorted(PAGE_ROOT.glob("*.md"), key=lambda path: int(path.name.split("_", 1)[0]))
    if len(pages) != 44:
        raise RuntimeError(f"expected 44 cluster pages, found {len(pages)}")
    doc = Document()
    style_document(doc)
    set_core_properties(doc)
    add_cover(doc, pages)
    add_overview(doc, pages)
    bookmark_id = 1
    for index, page_path in enumerate(pages, 1):
        bookmark_id = add_cluster_page(doc, page_path, index, bookmark_id)
    doc.add_page_break()
    add_appendix(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
