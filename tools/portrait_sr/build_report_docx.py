import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\Repository\ReadPaper\daily\PortraitSR")
OUTPUT = ROOT / "report" / "人像超分与人脸细节恢复_阶段性洞察_20260806.docx"
FIGURES = {
    "tiger": ROOT / "figures" / "representative_pages" / "tiger_page_05.png",
    "svfr": ROOT / "figures" / "representative_pages" / "svfr_page_04.png",
    "osdface": ROOT / "figures" / "representative_pages" / "osdface_page_03.png",
    "authface": ROOT / "figures" / "representative_pages" / "authface_page_04.png",
    "hdrface": ROOT / "figures" / "representative_pages" / "hdrface_page_04.png",
    "heads_up": ROOT / "figures" / "representative_pages" / "heads_up_page_03.png",
    "heads_up_results": ROOT / "figures" / "representative_pages" / "heads_up_results_page_06.png",
    "heads_up_training": ROOT / "figures" / "representative_pages" / "heads_up_training_page_16.png",
    "heads_up_dataset_details": ROOT / "figures" / "representative_pages" / "heads_up_dataset_details_page_17.png",
    "iconface": ROOT / "figures" / "representative_pages" / "iconface_page_02.png",
    "iconface_results": ROOT / "figures" / "representative_pages" / "iconface_results_page_05.png",
    "geomar": ROOT / "figures" / "representative_pages" / "geomar_page_03.png",
    "cfrnet": ROOT / "figures" / "representative_pages" / "cfrnet_page_03.png",
    "authface_data": ROOT / "figures" / "representative_pages" / "authface_data_page_03.png",
    "tiger_dataset": ROOT / "figures" / "representative_pages" / "tiger_dataset_page_08.png",
    "tiger_training_ablation": ROOT / "figures" / "representative_pages" / "tiger_training_ablation_page_12.png",
    "restorerid_degradation": ROOT / "figures" / "representative_pages" / "restorerid_degradation_page_04.png",
    "restorerid_training": ROOT / "figures" / "representative_pages" / "restorerid_training_page_08.png",
    "brushnet_architecture": ROOT / "figures" / "mask_conditioning" / "brushnet_architecture_page_06.png",
    "powerpaint_tasks": ROOT / "figures" / "mask_conditioning" / "powerpaint_tasks_page_05.png",
    "anydoor_architecture": ROOT / "figures" / "mask_conditioning" / "anydoor_architecture_page_03.png",
    "cosmicman_parsing": ROOT / "figures" / "mask_conditioning" / "cosmicman_parsing_page_05.png",
    "stableviton_architecture": ROOT / "figures" / "mask_conditioning" / "stableviton_architecture_page_04.png",
    "idm_vton_architecture": ROOT / "figures" / "mask_conditioning" / "idm_vton_architecture_page_05.png",
    "sapiens_parsing": ROOT / "figures" / "mask_conditioning" / "sapiens_parsing_page_06.png",
    "matanyone_architecture": ROOT / "figures" / "mask_conditioning" / "matanyone_architecture_page_04.png",
    "matanyone_results": ROOT / "figures" / "mask_conditioning" / "matanyone_results_page_06.png",
    "synthlight_architecture": ROOT / "figures" / "mask_conditioning" / "synthlight_architecture_page_04.png",
    "synthlight_results": ROOT / "figures" / "mask_conditioning" / "synthlight_results_page_06.png",
    "compose_architecture": ROOT / "figures" / "mask_conditioning" / "compose_architecture_page_04.png",
    "compose_results": ROOT / "figures" / "mask_conditioning" / "compose_results_page_11.png",
    "softshadow_architecture": ROOT / "figures" / "mask_conditioning" / "softshadow_architecture_page_03.png",
    "softshadow_results": ROOT / "figures" / "mask_conditioning" / "softshadow_results_page_07.png",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_text(paragraph, text, bold=False, italic=False, size=10.5, color=None):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    run.italic = italic
    return run


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, text)
    return paragraph


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(2)
    add_text(paragraph, text)
    return paragraph


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "DCE8F2")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(paragraph, header, bold=True, size=9, color=(22, 50, 79))
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            add_text(paragraph, str(value), size=8.5)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(document, key, caption):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(FIGURES[key]), width=Cm(16.2))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(cap, caption, italic=True, size=9, color=(70, 70, 70))


def add_heading(document, text, level):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        set_run_font(run, size={1: 16, 2: 13, 3: 11}.get(level, 10.5), bold=True, color=(22, 50, 79))
    return paragraph


def add_paragraph(document, text, bold_prefix=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(paragraph, bold_prefix, bold=True)
        add_text(paragraph, text[len(bold_prefix):])
    else:
        add_text(paragraph, text)
    return paragraph


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for style_name in ["List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_header_footer(document):
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text(header, "ReadPaper · PortraitSR · 2026-08-07", size=8.5, color=(90, 105, 115))
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(footer, "人像超分与人脸细节恢复阶段性洞察", size=8.5, color=(90, 105, 115))


def main():
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(44)
    add_text(title, "人像超分与人脸细节恢复", bold=True, size=24, color=(22, 50, 79))
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(subtitle, "2024-2026 工业界与前沿研究阶段性洞察", bold=True, size=15, color=(47, 111, 159))
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(18)
    add_text(meta, "ReadPaper PortraitSR Library\n证据截止日：2026-08-07", size=11, color=(80, 90, 100))
    document.add_paragraph()
    note = document.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(note.cell(0, 0), "FFF4D6")
    p = note.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "45 篇 PortraitSR 候选 · 新增 11 篇 Mask 扩展 PDF · 14 条 Mask 证据 · 33 张原文证据页", bold=True, size=11, color=(107, 79, 0))
    document.add_page_break()

    add_heading(document, "摘要", 1)
    add_paragraph(document, "本阶段研究聚焦人脸盲修复、全幅人像超分、参考身份保持、视频人脸修复、真实感与幻觉控制，以及一步化和端侧部署。当前资料库保留 45 篇 PortraitSR 核心候选，其中 30 篇发表于 2025-2026；本轮另新增 11 篇 mask conditioning、human parsing、matting、重光照、阴影编辑和 virtual try-on 核心 PDF，建立 14 条逐页证据记录和 15 张原文页面。工业界已由原文确认的代表包括 Xiaomi 的 TIGER 与 MeInTime、Tencent 的 SVFR 与 BrushNet、vivo 参与的 OSDFace/HDRFace/AuthFace、Topaz Labs 的 HeadsUp、Alibaba/Ant Group 的 AnyDoor、Meta 的 Sapiens、Adobe 的 SynthLight 与 COMPOSE，以及 Huawei、Kuaishou、Lenovo Research 等工作。")
    add_paragraph(document, "核心判断是：前沿研究正从“单张、对齐、固定人脸裁剪”转向“身份、几何、生成先验协同的全场景修复”。视频方向强调动态姿态与身份解耦；真实感方向开始引入摄影数据、频率约束和幻觉控制；效率方向从减少采样步数推进到一步流模型、量化友好网络和消费级 NPU 实测。本版进一步系统整理训练目标、LOSS 冲突、数据制作、RAW/ISP 退化、分阶段训练与手机相机域适配。")

    add_heading(document, "1. 研究范围与证据口径", 1)
    add_heading(document, "1.1 三条主研究线", 2)
    add_number(document, "Face restoration：重点恢复五官、脸型、身份和皮肤细节。")
    add_number(document, "Portrait restoration/SR：进一步覆盖头发、服装、身体、背景和全幅构图。")
    add_number(document, "Personalized/reference restoration：使用同一人物参考图、身份嵌入或个性化参数约束恢复结果。")
    add_heading(document, "1.2 证据等级", 2)
    add_table(document, ["等级", "证据", "可用于的结论"], [
        ["A", "已下载原文首页、正文、表格或图", "作者机构、方法结构、数据集、论文内实验数字"],
        ["B", "arXiv API 原始元数据与摘要", "标题、作者、版本日期、摘要声明、公开代码链接"],
        ["C", "官方项目页或数据集页快照", "发布状态、访问入口、许可或下载说明"],
        ["D", "技术归类与业务推断", "研究路线、系统架构和相机业务启示"],
    ], [1.2, 6.0, 8.5])
    add_paragraph(document, "赛事赞助方、部署芯片提供者和论文引用机构不会被当作作者公司。例如 CFRNet 在 HiSilicon Hi3402 NPU 上实测，但原文作者机构不是 Huawei；NTIRE 2026 的 OPPO、Kuaishou 是赛事赞助方，不等于挑战赛论文作者公司。")

    add_heading(document, "2. 主要技术变化", 1)
    add_heading(document, "2.1 几何化条件与人脸专用先验", 2)
    add_paragraph(document, "GeoMAR 针对 codebook 修复在重度退化下的条件歧义与一次预测脆弱问题，提取带空间锚点的人脸组件描述，通过 KV-Q exchange 注入低质量特征，再以 masked autoregressive refinement 逐步修复复杂区域。AuthFace 则从数据分布入手，使用 1,500 张原始分辨率超过 8K 的专业摄影人像和摄影师参与的标注流程微调生成先验，并通过 time-aware latent facial feature loss 约束眼睛和嘴部等关键区域。")
    add_figure(document, "geomar", "图 1  GeoMAR 原文第 3 页：几何对齐条件与掩码自回归修复框架。")
    add_figure(document, "authface", "图 2  AuthFace 原文第 4 页：摄影数据驱动的两阶段人脸专用生成先验。")
    add_paragraph(document, "HDRFace 由 City University of Hong Kong 与 vivo BlueImage Lab 合作提出，核心问题是重度退化使低质量输入缺失身份关键细节，仅靠 LQ condition 的扩散修复容易得到逼真但不忠实的结果。方法先用现成人脸修复器生成结构较可靠的中间结果，再用 DINOv3 从低质量输入和中间结果提取高维视觉表征（high-dimensional visual representations），并通过结构-细节感知自适应融合模块 SDFM 注入条件分支。结构建模阶段更强调原始输入的全局约束，细节合成阶段提高表征先验权重，从而在 structural consistency 与 detail fidelity 之间动态平衡。该模块不修改生成骨干，并分别在 SD V2.1-base 的 U-Net 扩散架构和 Qwen-Image 的 Rectified Flow DiT 架构上验证，说明其研究价值在于跨生成范式的条件注入，而非绑定单一基础模型。")
    add_figure(document, "hdrface", "图 3  HDRFace 原文第 4 页：DINOv3 高维表征、LQ/中间结果双条件与 SDFM 结构-细节自适应融合。")

    add_heading(document, "2.2 从人脸裁剪走向全幅 4K 人像", 2)
    add_paragraph(document, "HeadsUp 由 Texas A&M University 与 Topaz Labs 合作提出，研究对象不是对齐的人脸裁剪，而是同时包含人物、头发、服装和自然背景的全幅人像。工程上的直接痛点是：通用图像超分对人脸不敏感，而“通用 SR + 人脸专家 + 融合”又容易在脸部边缘产生边界伪影和局部风格不一致。HeadsUp 以 OSEDiff 类一步潜空间扩散模型为基础，对 VAE encoder 与 denoiser 施加 LoRA，固定 decoder；输入低质量全图、可选参考脸和脸部 mask，一步预测高质量潜变量。训练时将全图损失与脸部 fidelity、LPIPS 和 identity loss 组合，使小面积脸部获得更高监督权重。")
    add_paragraph(document, "论文同时提出 PortraitSR-4K：30K 张 4K 人像，报告划分为 27K 训练、3K 测试、163K 训练参考对和 190 组人像-参考测试对。该数据集弥补了通用 SR 数据和对齐人脸数据之间的空白，但原文说明其图像来自 LAION2B、Photo Concept Bucket、PD12M 等网络来源；独立许可和再分发权尚未核验，因此当前库只记录论文声明和来源入口，不下载数据。")
    add_figure(document, "heads_up", "图 4  HeadsUp 原文第 3 页：一步全幅人像超分、可选参考身份和脸部专用监督。")
    add_figure(document, "heads_up_results", "图 5  HeadsUp 原文第 6 页：通用 SR、融合式方案与端到端人像 SR 的视觉对比。")

    add_heading(document, "2.3 身份保持与局部细节恢复", 2)
    add_paragraph(document, "极端退化下，低清输入可能已不含足够身份信息。IDFSR 通过遮蔽不可靠人脸区域、对齐参考图风格并拟合身份 embedding，将风格与身份显式解耦。RestorerID 使用单张参考图，以 FIR-Adapter 重新平衡低清内容与参考身份冲突，并随退化程度调整身份注入强度。产品设计上，参考身份功能应单独授权，并提供可见的身份强度控制和失败回退。")
    add_paragraph(document, "IConFace 将身份保持从全局 ArcFace/AdaFace 相似度推进到局部稳定特征。它以 FLUX.2-Klein-base-4B 为骨干，支持 0-3 张同身份参考图：低质量图和参考图都保留为密集视觉 token，AdaFace 提供按可识别性加权的紧凑多参考身份向量，degraded-structure adapter 则用输入残差与全局/局部记忆反复锚定目标姿态和结构。训练采用 flow matching 与参考/目标身份损失。论文还构建人工核验的局部身份细节基准，对 167 个由配对真值和同身份参考共同支持的痣、疤痕、雀斑和色素区域进行保留率评估。该设计揭示一个重要边界：高全局身份相似度并不等于真实保留了人物的局部辨识细节。")
    add_figure(document, "iconface", "图 6  IConFace 原文第 2 页：FLUX.2 密集参考 token、AdaFace 多参考身份引导和结构记忆路径。")
    add_figure(document, "iconface_results", "图 7  IConFace 原文第 5 页：局部痣分布等持久身份细节的恢复对比。")

    add_heading(document, "2.4 视频修复成为独立问题", 2)
    add_paragraph(document, "Xiaomi MiLM Plus 的 TIGER 将视频人脸修复分解为 Identity、Geometry 和 Generative 三类先验：参考图提供静态身份，退化视频提供运动与视角，3D 人脸参数用于跨来源融合；生成端采用一步 rectified flow，并以三阶段训练优化结构、纹理和分布真实感。")
    add_figure(document, "tiger", "图 8  TIGER 原文第 5 页：身份、3D 几何和一步生成先验的三路融合。")
    add_paragraph(document, "Tencent Youtu Lab 的 SVFR 将视频盲修复、补全和着色统一到 Stable Video Diffusion 框架中，使用任务 embedding、Unified Latent Regularization、facial prior learning 与 self-referred refinement 提升任务共享和时序稳定性。FADRA 则使用 text-to-video diffusion 的时序先验，并加入低质量像素对齐、重复残差适配与频率感知约束。")
    add_figure(document, "svfr", "图 9  SVFR 原文第 4 页：统一视频盲修复、补全和着色的训练与推理流程。")

    add_heading(document, "2.5 一步化与端侧部署", 2)
    add_paragraph(document, "OSDFace 由 Shanghai Jiao Tong University 与 vivo 共同完成，通过视觉 tokenizer 与 VQ dictionary 生成 visual prompts，以身份损失约束一致性，并使用 GAN guidance 做分布对齐。论文在其统一实验设置下报告一步推理和 0.10 秒推理时间；该数字不能直接外推到手机端。")
    add_figure(document, "osdface", "图 10  OSDFace 原文第 3 页：VQ 视觉提示、身份损失与对抗引导的一步扩散训练。")
    add_paragraph(document, "CFRNet 采用 2.0M 参数 ResNet-style 网络与 Cycle-Consistent Fixed-Point Training。原文报告在 HiSilicon Hi3402 NPU 上以 INT8 运行，每个循环约 23 ms，循环次数可作为质量旋钮。其边界也很明确：只处理 256×256 人脸裁剪，不能代表全幅高分辨率人像修复。")
    add_figure(document, "cfrnet", "图 11  CFRNet 原文第 3 页：共享权重的循环固定点训练及端侧推理机制。")

    add_heading(document, "3. 工业界代表论文", 1)
    add_table(document, ["公司", "论文", "核心创新", "解决的问题"], [
        ["Xiaomi", "TIGER", "身份、3D 几何、生成先验；一步 rectified flow", "视频身份漂移、视角纠缠、时序真实感"],
        ["Tencent", "SVFR", "视频盲修复、补全、着色统一；SVD 时序先验", "单任务模型割裂、视频数据不足"],
        ["vivo", "OSDFace", "VQ visual prompts、身份损失、GAN guidance", "多步扩散慢且局部细节不协调"],
        ["vivo / Ant Group", "AuthFace", "8K 摄影数据、摄影标注、人脸特征损失", "通用 T2I 先验皮肤过平滑或生成错误细节"],
        ["vivo", "HDRFace", "DINOv3 高维表征、双条件、SDFM；兼容 U-Net/DiT", "重度退化下身份关键细节缺失，兼顾结构与细节忠实度"],
        ["Topaz Labs", "HeadsUp", "一步全幅人像扩散、脸部专用监督、可选参考身份", "通用 SR 人脸失真与双模型融合边界伪影"],
        ["Xiaomi", "MeInTime", "身份/年龄解耦、Gated Residual Fusion、Age-Aware Gradient Guidance", "跨年龄参考导致年龄漂移，适配历史照片与长期相册"],
        ["Huawei", "IQPFR", "无参考图像质量先验、质量条件码本与 Transformer", "训练真值平均质量限制恢复上限"],
        ["Alibaba", "DicFace", "Dirichlet 连续化码本、时空 Transformer", "视频细节恢复中的码本跳变与闪烁"],
        ["Kuaishou", "LRPO", "在线强化学习、复合奖励、似然正则", "在大解空间中提升感知质量并限制偏离真值"],
        ["Lenovo", "InfoBFR", "信息瓶颈压缩与扩散 LoRA 补偿", "预训练修复器的先验偏差、结构/纹理畸变和残留伪影"],
        ["Honor", "NTIRE 2026", "HONORAICamera 挑战赛方案", "真实世界感知质量和身份校验"],
    ], [2.8, 3.2, 5.5, 5.5])

    add_heading(document, "4. 数据集与评测", 1)
    add_table(document, ["数据角色", "代表数据", "用途", "主要风险"], [
        ["高质量静态人脸", "FFHQ、CelebA-HQ、AuthFace 8K set", "皮肤、五官、发丝生成先验", "人像许可、来源版权、分布偏差"],
        ["全幅 4K 人像", "PortraitSR-4K", "脸、头发、服装、背景统一超分", "网络来源、派生许可、身份与再分发权"],
        ["身份数据", "VGGFace2、Celeb-Ref、Reface-HQ", "身份 embedding、参考修复", "生物特征隐私、非商业条款"],
        ["跨年龄身份", "FG-NET、AgeDB、CACD、IMDB-WIKI、UTKFace", "年龄一致性与跨年龄参考评测", "年龄标签噪声、名人偏差、儿童与生物特征隐私"],
        ["局部身份细节", "IConFace 167-trait benchmark", "痣、疤痕、雀斑、色素区域保留率", "人工标注规模小、发布与许可尚待核验"],
        ["高质量视频", "VFHQ、CelebV-HQ、HDTF", "时序修复、姿态变化", "视频来源权利、身份同意"],
        ["真实世界评测", "LFW-Test、WIDER-Test、FOS、NTIRE", "复杂退化与主观质量", "多数无配对 GT、指标偏差"],
    ], [3.0, 5.2, 4.5, 4.5])
    add_paragraph(document, "数据集目录已记录 23 项，20 个官方或论文页面已归档。FFHQ、CelebA、VGGFace2、AgeDB 等条款不同；PortraitSR-4K、IConFace 局部细节基准、AuthFace 8K set、Reface-HQ、Celeb-Ref 与 TIGER 派生视频集必须视为“论文明确使用或提出，但独立许可与完整下载入口仍待核验”。当前失败记录包括 FFHQ 页面请求超时、LFW 官方域名 DNS 失败和旧 VFRx 项目地址返回 404；资料库不会自动下载或再分发受限人脸数据。")
    add_heading(document, "4.1 指标解释", 2)
    for item in [
        "PSNR/SSIM 偏像素失真，可能奖励平滑结果。",
        "LPIPS/DISTS 更接近深层感知差异，但不能验证人物身份。",
        "ArcFace/AdaFace similarity 用于身份一致性，但受姿态、年龄和域偏差影响。",
        "MUSIQ/MANIQA/CLIP-IQA/TOPIQ 适合真实图，但可能偏好更锐、更生成式的纹理。",
        "FVD 和时序指标用于视频一致性，局部眼口闪烁仍需人工视频检查。",
    ]:
        add_bullet(document, item)

    add_heading(document, "5. 训练目标与 LOSS 设计", 1)
    add_paragraph(document, "人像超分训练的核心矛盾不是 LOSS 数量不足，而是多个目标互相冲突：像素重建倾向平滑，感知与对抗目标倾向锐利但可能虚构细节，身份损失只能约束整体可识别性，视频时序约束又可能压制快速运动细节。下式是跨论文归纳，不是任一论文的原始公式：")
    formula = document.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(formula, "L_total = w_task L_task + w_rec L_rec + w_per L_per + w_id L_id + w_local L_local + w_dist L_dist + w_temp L_temp", bold=True, size=10.5, color=(22, 50, 79))
    add_table(document, ["目标", "常见实现", "主要作用", "主要风险"], [
        ["任务基础", "noise/velocity/latent MSE", "保持扩散或流模型的主训练几何", "被附加损失压制后轨迹失真"],
        ["重建", "L1、pixel/latent MSE", "锚定轮廓、颜色和低频结构", "过强会平滑皮肤、睫毛和头发"],
        ["感知", "LPIPS、VGG features", "恢复边缘与视觉纹理", "可能奖励非事实细节"],
        ["身份", "ArcFace/AdaFace cosine", "保持整体人物可识别性", "不能证明痣、疤痕等局部事实"],
        ["局部", "face LPIPS、眼口/landmark、局部 GAN", "防止小脸梯度被背景淹没", "可能局部过锐或产生边界不连续"],
        ["分布", "GAN、VSD、distribution guidance", "减少平均化纹理、提升真实感", "幻觉、震荡和 identity drift"],
        ["时序/频率", "spatiotemporal GAN、DCT、ULR", "抑制闪烁、恢复高频", "可能牺牲快速运动与单帧锐度"],
    ], [2.5, 4.2, 5.1, 5.0])

    add_heading(document, "5.1 从结构锚定到分布真实感", 2)
    add_paragraph(document, "TIGER 的三级训练体现了清晰的目标递进。Stage I 仅以 latent MSE 建立身份和几何一致的潜空间映射；Stage II 联合优化 DiT mapper 与 VAE decoder，加入 LPIPS 和 pixel MSE 修复纹理与边缘；Stage III 再加入时空判别器，并冻结 decoder、只优化 mapper，以降低 mode collapse。原文采用的 LOSS 权重配置为 (lambda1, lambda2, lambda3, lambda4)=(1.0, 2.0, 10.0, 0.2)，但该数值只适用于 TIGER 的模型、数据和损失归一化。")
    add_figure(document, "tiger_training_ablation", "图 12  TIGER 原文第 12 页：三级训练视觉递进与 latent、perceptual、pixel、adversarial 权重消融。")
    add_paragraph(document, "HeadsUp 在全图 MSE、LPIPS、VSD 之外加入 face fidelity、face LPIPS、identity 和 face adversarial loss，原文给出 lambda_fid=1、lambda_face-LPIPS=0.8、lambda_id=4。关键思想是对小面积人脸单独提权，否则全幅人像中的背景、服装和头发会主导梯度。")

    add_heading(document, "5.2 身份、局部、Flow Matching 与稳定训练", 2)
    for item in [
        "IConFace 使用冻结 AdaFace 的参考/目标 cosine loss、stop-gradient identity target 和 timestep weighting；高噪声时降低身份监督，避免干扰 flow matching。",
        "AuthFace 将监督集中到眼睛和嘴巴，并让局部判别器感知 diffusion timestep；原文消融表明静态局部判别器会破坏皮肤或牙齿细节。",
        "FADRA 保留 flow-matching velocity MSE，并在 latent 上做 8x8 DCT 频率监督；同一 clip 共享退化参数，避免训练数据制造随机闪烁。",
        "SVFR 的目标由 diffusion noise MSE、跨任务 Unified Latent Regularization 和 68 点 landmark prior 组成。",
        "CFRNet 将 fixed-point 权重在前 100K 次迭代从 0 提升至 0.5，GAN warm-up 5K，再用 5K 逐步打开对抗和局部组件损失。",
    ]:
        add_bullet(document, item)
    add_paragraph(document, "综合建议：先训练 task/reconstruction 主目标，再打开 identity/local，最后 warm-up 并 ramp LPIPS、GAN 或 VSD；必须记录每个 LOSS 对共享参数的 gradient norm 和 cosine similarity，不能只比较标量值。")

    add_heading(document, "6. 训练数据制作", 1)
    add_heading(document, "6.1 高质量真值与摄影分布", 2)
    add_paragraph(document, "AuthFace 收集 1.5K 张原始分辨率超过 8K 的专业摄影人像，并在普通语义标签之外增加光线、焦点、景深、构图、肤质和妆容等 photographic tags。其目的不是简单扩大数据量，而是修正通用 T2I 数据中皮肤过平滑、焦点错误和非真实人脸纹理。")
    add_figure(document, "authface_data", "图 13  AuthFace 原文第 3 页：摄影标签、高质量人脸数据与 SDXL 微调前后对比。")
    add_paragraph(document, "HeadsUp 的 PortraitSR-4K 从 LAION2B、Photo Concept Bucket、PD12M 等来源筛选候选，要求至少 3840x2160、宽高比 0.6-1.6，经 Q-Align 筛选；双眼距离至少 64 像素，同身份参考对 CVLFace 相似度至少 0.65。论文报告 30K 图、27K/3K train-test、163K 训练参考对和 190 测试参考对。网络来源不等于可自由再分发，本库只保存论文证据和入口。")
    add_figure(document, "heads_up_training", "图 14  HeadsUp 原文第 16 页：PortraitSR-4K 样例、模型输入及两阶段混合训练。")
    add_figure(document, "heads_up_dataset_details", "图 15  HeadsUp 原文第 17 页：4K、宽高比、Q-Align、人脸尺度和身份相似度筛选。")

    add_heading(document, "6.2 样本视图、身份隔离与参考配对", 2)
    add_number(document, "Full portrait view：监督头发、服装、身体、背景和整体色彩一致性。")
    add_number(document, "Aligned face view：监督眼、口、鼻、皮肤和身份，并保存全图到人脸裁剪的可微变换。")
    add_number(document, "Identity reference view：来自同身份不同时间、姿态或场景，不得与目标同帧、近重复或同一次连拍。")
    add_paragraph(document, "训练、验证和测试必须按身份划分，而不是按图片随机划分；视频还应按原始 clip 或拍摄 session 隔离。IConFace 的 0/1/2/3 张参考概率 30/30/20/20 与 HeadsUp 的 reference dropout=0.2，都是为了支持多参考、弱参考和无参考回退。")

    add_heading(document, "6.3 从二阶退化扩展到 RAW/ISP", 2)
    add_paragraph(document, "建议把合成退化拆成光学、传感器、ISP、编码和显示链：optical blur/defocus/motion -> resize/digital zoom -> shot/read noise -> demosaic/AWB/CCM/tone mapping -> denoise/sharpen -> JPEG/HEIF/二次压缩。只在 sRGB 上添加独立高斯噪声，无法覆盖真实手机噪声和 ISP 耦合。")
    add_paragraph(document, "RestorerID 使用 blur-resize-noise-JPEG 二阶退化，并把 sRGB 经 ISP 模型转换到 RAW 域后再添加相机噪声。IConFace 在线以 50% BSRGAN 和 50% Real-ESRGAN 生成 LQ，并按 0-3、4-8、9-16 三个退化强度桶以 0.5/0.3/0.2 采样。")
    add_figure(document, "restorerid_degradation", "图 16  RestorerID 原文第 4 页：二阶退化与 sRGB-to-RAW/ISP 相机噪声模拟。")

    add_heading(document, "6.4 视频筛选、时变退化与数据治理", 2)
    add_paragraph(document, "TIGER 从 OpenHumanVid、CelebV-HQ、VFHQ 按运动幅度、人脸清晰度和人脸占比筛选：人脸中心位移小于 100 像素、clarity score 高于 0.7、人脸最短边大于 320 像素，最终保留 28,491 个训练样本。")
    add_figure(document, "tiger_dataset", "图 17  TIGER 原文第 8 页：运动、清晰度和人脸占比驱动的视频筛选。")
    for item in [
        "同一视频 clip 共享主要 blur/noise/compression 参数，但可加入缓慢时间漂移，模拟 AE、AWB、AF 和防抖波动。",
        "按 tiny/small/medium/large 人脸、姿态、运动、低照度、逆光、眼镜、牙齿、饰品和复杂发型建立难例桶。",
        "多摄训练要覆盖主摄/长焦/超广角之间的视角、色彩、噪声、锐化和畸变域差。",
        "记录来源、许可、商业用途、身份同意、生物特征限制、再分发权、删除入口和派生链；公开可见不等于可训练或可再分发。",
    ]:
        add_bullet(document, item)

    add_heading(document, "7. 推荐训练流程与实验设计", 1)
    add_paragraph(document, "本节是跨论文综合建议，不是单篇论文的原设置。")
    add_table(document, ["阶段", "目标", "可训练模块", "主要 LOSS", "退出条件"], [
        ["A 结构预训练", "稳定 LQ-to-HQ、守轮廓姿态颜色", "adapter/ControlNet/LoRA/mapper", "task + reconstruction", "PSNR/LPIPS/landmark 稳定，真实图无结构漂移"],
        ["B 人脸身份强化", "眼口皮肤发际线与身份", "face branch、ID projector、structure adapter", "task + rec + identity + local", "身份与局部细节提高，轻退化不过修"],
        ["C 感知分布对齐", "减少塑料感、平均纹理和闪烁", "受控 mapper/decoder、discriminator", "+ perceptual + GAN/VSD + temporal", "主观质量提高且事实/身份/时序不回退"],
        ["D 相机域与部署", "RAW/ISP、夜景、变焦、多摄、端侧", "轻量 adapter、QAT/NPU 模块", "域重建、蒸馏、量化/固定点", "真实设备闭环通过并可回退保真路径"],
    ], [2.4, 4.0, 4.3, 4.2, 4.0])

    add_heading(document, "7.1 冻结/解冻与学习率", 2)
    add_paragraph(document, "AuthFace 先全 U-Net 学摄影人脸先验，再冻结 U-Net 训练 ControlNet；RestorerID 先训练基础修复，再只训练 FIR-Adapter；TIGER 先 mapper、再 mapper+decoder、最后冻结 decoder 做对抗对齐。这些方法共同说明：先让模块形成清晰职责，再联合训练或增加高风险目标。")
    add_figure(document, "restorerid_training", "图 18  RestorerID 原文第 8 页：单阶段与基础修复/FIR-Adapter 两阶段训练消融。")
    add_paragraph(document, "建议解冻顺序为：条件分支/adapter -> 高层生成 block -> VAE decoder 纹理层 -> 必要时小学习率全量微调。新初始化的小模块通常用较高学习率，预训练 LoRA 更保守；例如 IConFace 的 base、LoRA、identity projector 学习率分别为 1e-5、5e-7、3e-5。")

    add_heading(document, "7.2 LOSS 搜索起点与最小消融", 2)
    for item in [
        "先对 LOSS 做 batch 内或 EMA 尺度归一化，以 task/rec 为锚点；Stage A 先搜索 task:rec=1:0.5/1/2。",
        "Stage B 的 identity/local 可从 0.05/0.1/0.2 小网格起步，并按参考质量和身份置信度动态衰减。",
        "Stage C 的 LPIPS 可从 0.05/0.1/0.2 起步；GAN/VSD 从 0 warm-up 至 0.01/0.05/0.1 搜索。以上只是搜索起点，不是论文结论或性能承诺。",
        "最少比较 A、A+B、A+B+C、完整 A-D；分别去掉 identity、local、perceptual、GAN/VSD、temporal/frequency。",
        "同时比较 BSRGAN、Real-ESRGAN、RAW/ISP 和真实采集混合退化，以及 0/1/多参考、adapter-only/LoRA/decoder 解冻。",
        "每组实验记录 PSNR/LPIPS、IQA、ArcFace/AdaFace、局部事实保留、视频闪烁、各 LOSS gradient norm 与梯度夹角。",
    ]:
        add_bullet(document, item)

    add_heading(document, "7.3 面向手机相机的训练配方", 2)
    for item in [
        "覆盖不同 ISO、曝光、焦距、数字变焦、运动、肤色、年龄、逆光、夜景、多脸和多摄，并尽可能保存 RAW/YUV/JPEG 或 HEIF 对。",
        "按人物划分 train/val/test；参考图不得与目标同帧或近重复，相册身份参考必须支持授权、撤回和删除。",
        "默认输出低幻觉保真版本；仅在高退化且质量闸门通过时启用生成纹理，参考身份必须由用户主动开启。",
        "端侧轻量模型可使用教师蒸馏、QAT 或 fixed-point/cycle consistency；服务器大模型的速度和质量数字不能直接外推到手机。",
        "上线验收同时覆盖视觉质量、身份一致、事实细节和视频时序，任何单一 IQA 或 ID 分数都不足以放行。",
    ]:
        add_number(document, item)
    add_paragraph(document, "逐篇训练设置、页码和证据等级已结构化保存于 metadata/training_evidence_matrix.json。")

    add_heading(document, "8. Mask 条件驱动的人像增强与可控编辑", 1)
    add_paragraph(document, "本章回答一个工程问题：当输入不仅有低质量人像，还提供人物 mask 或更细的人体语义区域时，模型能否同时完成局部超分、身份细节恢复、皮肤/发丝增强、背景虚化、重光照、衣物编辑、遮挡修复和视频稳定？结论是可以，但 mask 不能只被理解为一个硬二值开关。更稳健的系统需要把它组织为四层空间控制：编辑区域、语义部件、软边界/遮挡，以及置信度与强度。以下“论文报告”来自原文；标注为“综合建议”的公式、权重和训练调度需要通过本项目实验验证。")

    add_heading(document, "8.1 Mask 的四层表示", 2)
    add_table(document, ["层级", "推荐表示", "解决的问题", "典型来源"], [
        ["L0 编辑域", "人物/脸部/目标区域 binary mask", "限定超分、修复、移除或替换的大范围", "HeadsUp、BrushNet、PowerPaint"],
        ["L1 语义域", "皮肤、头发、眼、眉、鼻、口唇、牙齿、耳、衣物、饰品、背景", "让不同区域使用不同条件和 LOSS，避免大面积皮肤淹没小部件", "Sapiens、CosmicMan、GeoMAR"],
        ["L2 过渡域", "alpha matte、trimap、boundary band、occlusion map、soft shadow map", "处理发丝、衣领、眼镜和半影，减少 halo、接缝和颜色泄漏", "MatAnyone、SoftShadow"],
        ["L3 控制域", "confidence、uncertainty、edit-strength、reference-reliability map", "解析错误或遮挡时自动减弱高风险生成", "跨论文综合建议"],
    ], [2.2, 4.8, 6.3, 4.2])
    add_paragraph(document, "Sapiens 的 28 类人体解析覆盖头发、舌头、牙齿、上下唇和躯干等细粒度区域；GeoMAR 则证明 parsing map 可为五官提供空间锚点，但在严重退化或遮挡时必须与 RGB 外观线索交叉验证，并通过 component dropout 降低对完美解析的依赖。因而，训练时应保存 hard label、每类 soft probability、边界距离和置信度，而不是只保留 argmax 后的单通道标签图。")
    add_figure(document, "sapiens_parsing", "图 19  Sapiens 原文第 6 页：细粒度人体部件解析及其监督定义。")
    add_figure(document, "cosmicman_parsing", "图 20  CosmicMan 原文第 5 页：人体 parsing 与分部位文本标签的对应关系。")

    add_heading(document, "8.2 输入人像 Mask 后可以实现的效果", 2)
    add_table(document, ["效果", "核心 Mask", "主要条件/模型", "高风险点"], [
        ["局部超分与五官恢复", "face + eyes/mouth/teeth + confidence", "SR adapter/ControlNet；ROI perceptual；identity", "假睫毛、假牙齿、局部身份漂移"],
        ["皮肤增强", "skin - eyes/lips/hair + blemish map", "低频颜色/曝光校正 + 弱纹理恢复", "磨皮塑料感、伪造毛孔、肤色偏移"],
        ["头发与发际线", "hair alpha + boundary band", "多尺度边缘/方向频率 + matte-guided decoder", "halo、粘连、凭空补发"],
        ["眼口牙齿局部精修", "component ROI + landmark + occlusion", "局部专家/ROI discriminator/feature loss", "小区域过锐、表情改变"],
        ["背景虚化或替换", "person alpha/trimap", "matting + depth/bokeh renderer 或 inpainting", "发丝边缘污染、前后景漏色"],
        ["重光照与阴影控制", "person/skin + shadow soft map + strength", "environment-map condition、relighting diffusion", "身份变化、光向不一致、阴影边界"],
        ["衣物编辑/试穿", "clothes + body/pose + occlusion", "garment reference、cross/self-attention correspondence", "图案变形、手臂遮挡、体型漂移"],
        ["遮挡/杂物移除", "edit mask + protect mask + boundary", "BrushNet/PowerPaint 类 inpainting", "非编辑区域变化、重复物体、上下文误补"],
        ["视频人像增强", "tracked semantic masks + alpha + validity", "memory propagation/temporal adapter", "遮挡后记忆污染、边界闪烁"],
    ], [3.1, 4.7, 6.0, 4.2])
    add_paragraph(document, "BrushNet 把 mask 与 masked-image feature 放到独立分支，并逐层注入冻结生成骨干；AnyDoor 同时使用 shape mask、ID token 和高频 detail map，说明空间位置、对象身份和细节纹理应由不同条件负责。对于人像 SR，可对应为：mask 决定在哪里增强，LQ/参考特征决定恢复什么，edit-strength 决定增强到什么程度。")
    add_figure(document, "brushnet_architecture", "图 21  BrushNet 原文第 6 页：mask、masked image 与 noisy latent 的双分支逐层注入。")
    add_figure(document, "anydoor_architecture", "图 22  AnyDoor 原文第 3 页：shape mask、身份 token 与高频细节图的分工。")
    add_figure(document, "powerpaint_tasks", "图 23  PowerPaint 原文第 5 页：task prompt 与 mask 配合完成填充、移除和形状控制。")

    add_heading(document, "8.3 四种 Mask 注入架构", 2)
    add_number(document, "输入通道拼接（early concatenation）：类似 HeadsUp，把 LQ latent、参考 latent 与 mask 拼接到 denoiser 输入。实现简单、位置明确，但语义和不确定性容易在深层被稀释。")
    add_number(document, "独立条件分支（ControlNet/BrushNet）：mask、masked image、edge/normal/depth 在独立 encoder 中提特征，再逐层注入主干。适合精细边界与多尺度空间控制。")
    add_number(document, "注意力路由（attention routing）：StableVITON/IDM-VTON 用 mask/agnostic map 定义人体结构，再把衣物或参考特征送入 cross/self-attention。适合衣物、发型、妆容等参考驱动任务。")
    add_number(document, "损失与采样门控（loss/sampling gating）：AuthFace 只在脸区计算 Stage-I noise reconstruction；推理阶段还可把不同区域映射为不同 guidance scale、采样步数或生成强度。")
    add_figure(document, "stableviton_architecture", "图 24  StableVITON 原文第 4 页：agnostic map/mask、DensePose 与服装 cross-attention 对应。")
    add_figure(document, "idm_vton_architecture", "图 25  IDM-VTON 原文第 5 页：高层服装语义进入 cross-attention，低层细节进入 self-attention。")

    add_heading(document, "8.4 区域归一化 LOSS：先守编辑边界，再提高区域效果", 2)
    add_paragraph(document, "综合建议的区域归一化形式为：L_region = sum_r w_r * [sum(M_r * error_r) / (sum(M_r) + epsilon)]。分母按区域面积归一化是关键，否则皮肤、衣物和背景等大区域会压制眼睛、牙齿、眼镜和饰品等小区域。M_r 建议使用 soft probability × confidence，而不是始终使用硬 0/1。")
    add_table(document, ["目标", "建议定义", "主要作用", "不宜施加的区域"], [
        ["L_task", "diffusion noise / flow velocity / latent regression", "保持生成骨干的训练几何", "无"],
        ["L_edit", "编辑区 Charbonnier/L1 + perceptual", "确保目标区域真正发生预期变化", "保护区"],
        ["L_keep", "非编辑区 pixel/latent/feature consistency", "防止背景、衣物、脸型被连带改写", "编辑区中心"],
        ["L_boundary", "dilate(M)-erode(M) 上的 gradient/color/alpha", "减少 halo、接缝和颜色泄漏", "远离边界的区域"],
        ["L_identity", "ArcFace/AdaFace + 局部事实特征", "保持人物整体及痣、疤痕等稳定特征", "背景；衣物通常不应强约束"],
        ["L_component", "eyes/mouth/teeth/eyebrow ROI perceptual/landmark", "小部件结构与真实细节", "皮肤大平坦区"],
        ["L_freq", "hair/texture 区域 DCT/FFT/gradient", "恢复发丝和纺织方向性高频", "皮肤上应弱化"],
        ["L_temp", "warped mask/alpha/feature consistency", "视频区域不闪烁且遮挡后可恢复", "镜头切换、无效光流区"],
    ], [2.4, 6.0, 6.0, 4.0])
    add_paragraph(document, "综合建议的总目标可写为 L = L_task + lambda_edit L_edit + lambda_keep L_keep + lambda_boundary L_boundary + lambda_id L_identity + lambda_part L_component + lambda_freq L_region-frequency + lambda_temp L_mask-temporal。该式是实验设计模板，不是任何单篇论文直接给出的最终配方。实际权重应先测量各项梯度范数与夹角，再通过分阶段消融确定。")
    add_paragraph(document, "三个常见冲突需要单独诊断：第一，identity 与 pixel/perceptual 冲突时，模型可能得到可识别但过平滑的脸，HeadsUp 已观察到只用 identity loss 会降低图像质量；第二，GAN/感知损失在皮肤和牙齿区域过强会生成不存在的纹理；第三，强 L_keep 会阻止边界区自然融合，因此应把 mask 拆成 edit core、transition band 与 strict keep 三个区域。")

    add_heading(document, "8.5 软边界、阴影与重光照", 2)
    add_paragraph(document, "人像发丝、衣领、眼镜透明区域和阴影半影都不是二值边界。SoftShadow 把阴影表示为连续灰度 mask：0 表示受光区，1 表示本影，(0,1) 表示半影，并联合优化 shadow-removal、mask reconstruction 与 penumbra gradient constraint。这个思想可迁移到 portrait SR：边界强度不只表示属于/不属于人物，也可表示需要保留多少原像素、允许生成多少新细节。")
    add_figure(document, "softshadow_architecture", "图 26  SoftShadow 原文第 3 页：soft mask 预测、去阴影网络及三类损失。")
    add_figure(document, "softshadow_results", "图 27  SoftShadow 原文第 7 页：半影区域相对硬 mask 方法的边界改善。")
    add_paragraph(document, "COMPOSE 将环境光与可编辑主光源解耦，可控制阴影位置、形状、软硬和强度；SynthLight 则用目标 environment map 条件训练扩散重光照，并通过真实人像多任务数据和 image-conditioned CFG 缓解合成域偏移与身份变化。若用户只提供人物 mask，可以实现前景整体重光照；若进一步提供 skin/hair/clothes 子 mask 与 shadow-strength map，就能给不同材质设不同的高光、阴影和保护强度。")
    add_figure(document, "compose_architecture", "图 28  COMPOSE 原文第 4 页：光照估计、扩散、阴影合成与最终 compositing。")
    add_figure(document, "compose_results", "图 29  COMPOSE 原文第 11 页：阴影柔化、增强、改变大小与位置的可控效果。")
    add_figure(document, "synthlight_architecture", "图 30  SynthLight 原文第 4 页：合成重光照与真实人像多任务联合训练。")
    add_figure(document, "synthlight_results", "图 31  SynthLight 原文第 6 页：不同环境光条件下的人像效果。")

    add_heading(document, "8.6 训练数据如何制作", 2)
    add_number(document, "为每张 HQ 人像离线生成 face/body semantic labels、soft probabilities、person alpha、boundary distance、landmarks、occlusion 和 quality/confidence；hair、glasses、teeth、accessories 保留独立小类。")
    add_number(document, "从 HQ 构造多视图：全幅人像、对齐脸、五官 ROI、同身份参考、衣物参考；所有视图共享 identity 与原始文件 ID，数据划分按身份和拍摄 session 隔离。")
    add_number(document, "构造 task pair：SR 使用真实/合成 LQ；背景任务使用 alpha compositing；重光照使用 OLAT、环境图渲染或可信 relighting teacher；衣物任务使用 person-garment pair 与 agnostic map；移除任务使用实例 mask 和 clean plate/生成真值。")
    add_number(document, "训练时随机破坏 mask：dilation/erosion、边界抖动、孔洞、错位、组件 dropout、类别混淆、低分辨率锯齿、漏标遮挡。GeoMAR 的 component dropout 与 PowerPaint 的随机扩张都说明模型必须见过不完美空间条件。")
    add_number(document, "为保护区构造显式负样本：同一输入改变 mask 但保持 prompt，或保持 mask 但改变 task token，检查模型是否只修改指定区域。BrushNet 的 random/segmentation inside/outside 任务可作为这种对照思路。")
    add_number(document, "视频数据中传播 mask/alpha 时记录 flow validity、遮挡和镜头切换；memory 只在高置信度区域更新。MatAnyone 表明语义 core 与 alpha boundary 应分别建模，避免记忆被相似背景污染。")
    add_figure(document, "matanyone_architecture", "图 32  MatAnyone 原文第 4 页：语义 core、细节 boundary 与一致性 memory propagation。")
    add_figure(document, "matanyone_results", "图 33  MatAnyone 原文第 6 页：语义稳定、边界细节与真实视频效果对比。")

    add_heading(document, "8.7 推荐的多任务训练流程", 2)
    add_table(document, ["阶段", "任务与可训练模块", "开启的 Mask/LOSS", "退出条件"], [
        ["M0 解析/Matting", "冻结或独立训练 Sapiens/SAM/matting", "CE/Dice/alpha/boundary；mask corruption", "小部件与发丝边界可用，置信度校准"],
        ["M1 保真 SR", "restoration adapter/LoRA", "L_task + global rec + area-normalized region rec + L_keep", "轻/中退化不改身份和背景"],
        ["M2 区域专家", "face/hair/eye/mouth branches", "+ L_component + weak L_identity + L_boundary", "小部件提升且无局部过锐"],
        ["M3 效果任务", "relight/inpaint/garment condition branches", "+ task token + edit-strength + effect-specific losses", "一种 mask 对应一种可预测效果，保护区稳定"],
        ["M4 感知生成", "受控 decoder/denoiser + discriminator", "+ LPIPS/GAN/VSD ramp", "主观质量提升且幻觉闸门通过"],
        ["M5 视频/部署", "temporal memory、distillation、QAT", "+ L_temp + fixed-point consistency", "遮挡恢复、无闪烁、设备闭环通过"],
    ], [2.5, 5.2, 6.8, 5.0])
    add_paragraph(document, "多任务采样不要简单平均。建议按任务难度与有效区域面积做 bucket，并保证每个 batch 至少包含一种保真任务；生成式 effect batch 不能连续占据过高比例，否则骨干会遗忘原始人像恢复。task token/segmented mask/strength map 三者要做独立 dropout，使模型能够在缺少某一条件时回退，而不是输出随机效果。")

    add_heading(document, "8.8 最小消融与评测矩阵", 2)
    add_table(document, ["变量", "最小对照", "必须报告"], [
        ["Mask 粒度", "binary person；face；10 类；28 类；semantic + alpha + confidence", "编辑区、保护区、边界区分开指标"],
        ["注入方式", "input concat；ControlNet；dual branch；attention routing；loss-only", "质量、显存、延迟、mask 泄漏"],
        ["边界", "hard；Gaussian blur；trimap；learned alpha；distance transform", "halo、boundary PSNR/LPIPS、发丝主观盲测"],
        ["Mask 误差", "clean；dilate/erode；shift；holes；component dropout；label confusion", "性能-扰动曲线与失败图"],
        ["LOSS", "global；+ region norm；+ keep；+ boundary；+ identity/component；+ GAN", "梯度范数/夹角及每步收益"],
        ["效果强度", "0/0.25/0.5/0.75/1.0 strength map", "单调性、身份、非编辑区漂移"],
        ["视频", "逐帧；flow consistency；memory；memory + uncertainty gate", "warp error、flicker、遮挡恢复"],
    ], [3.0, 8.0, 7.5])
    add_paragraph(document, "评测至少拆成三块：E_edit 只衡量目标区域是否达到预期；E_keep 衡量非编辑区域是否保持；E_boundary 衡量 transition band 是否平滑。对于脸部还要独立报告 identity、landmark、局部事实保留与人工盲测；对于重光照报告光向/阴影一致；对于衣物报告图案、logo 与遮挡关系；对于视频报告时序稳定和遮挡后的重新识别。")

    add_heading(document, "8.9 面向手机相机的具体建议", 2)
    for item in [
        "拍摄链路输出 person/face/hair/skin/clothes 的低分辨率 soft masks 与 confidence 即可；高分辨率边界可在 SR decoder 中联合细化，避免先做昂贵全分辨率 parsing。",
        "默认只开启保真 SR、局部去噪和轻度皮肤/发丝增强；重光照、背景替换、衣物编辑属于显式效果模式，应由用户选择并显示强度。",
        "把脸部 strict protect、五官 recover、皮肤 conservative、头发 boundary-aware、衣物 editable、背景 replaceable 设为产品级策略，而不是只传一个 person mask。",
        "参考身份只作用于 face/ear/hairline 等受控区域，避免迁移参考图的妆容、年龄、光照和衣服；低置信度或多脸冲突时关闭 identity branch。",
        "生成结果保留可回退的 fidelity checkpoint，并保存 mask、强度、模型版本与质量闸门日志，便于复现误修与灰度发布。",
    ]:
        add_number(document, item)
    add_paragraph(document, "逐篇证据、页码、机构和局限已保存于 metadata/mask_conditioning_evidence_matrix.json；本章原文截图来自 figures/mask_conditioning，PDF 位于 papers/07_mask_conditioning。")

    add_heading(document, "9. 对手机相机的启示", 1)
    for item in [
        "端侧默认路径：检测、对齐、去噪、去压缩、温和锐化和低幻觉重建。",
        "生成增强路径：只在严重数字变焦、老照片或用户主动选择时启用。",
        "参考身份路径：单独授权相册参考，明确显示是否使用参考人物信息。",
        "视频路径：复用跟踪、3D 人脸参数和历史帧，避免逐帧独立生成。",
        "质量闸门：检测眼镜、牙齿、文字、饰品、遮挡和多脸冲突，失败时回退到保真结果。",
    ]:
        add_number(document, item)
    add_heading(document, "9.1 当前研究空白", 2)
    for item in [
        "缺少 RAW 域、HDR、多摄融合与生成式人像修复的统一公开基准。",
        "缺少肤色、公平性、年龄变化和儿童人脸的系统身份风险评测。",
        "全幅人像超分明显少于对齐人脸裁剪，头发、服装与背景的一致性不足。",
        "消费 NPU 实测很少，扩散论文通常只报告服务器 GPU 或理论一步数。",
        "参考身份数据的许可、删除权和端云隐私机制常被方法论文忽略。",
    ]:
        add_bullet(document, item)

    add_heading(document, "10. 推荐阅读路线", 1)
    add_table(document, ["阶段", "论文", "学习目标"], [
        ["基础", "CodeFormer、DifFace、NTIRE 2025/2026", "理解 codebook、扩散后验、真实评测和身份检查"],
        ["工业界", "HeadsUp、AuthFace、HDRFace、OSDFace、TIGER、SVFR、CFRNet", "理解全幅人像、摄影数据、高维表征、一步扩散、视频先验和端侧部署"],
        ["身份细节", "IConFace、MeInTime、IDFSR、RestorerID", "理解 FLUX.2 多参考、局部稳定特征、跨年龄与参考冲突"],
        ["质量与优化", "IQPFR、InfoBFR、LRPO、Latent-PMRF", "理解质量先验、信息瓶颈、强化学习和感知-失真折中"],
        ["待补", "Face2Scene、DVFace、DicFace", "场景修复、双先验视频扩散和概率码本"],
    ], [2.2, 6.5, 8.0])

    add_heading(document, "11. 局限与下一步", 1)
    add_paragraph(document, "本报告是阶段性研究手册，不是完整系统综述。45 篇种子池中仍有 27 篇只有元数据或 HTML 原文证据、等待 PDF 下载；部分 2026 工作是预印本，标题、作者或实验可能随版本更新。HeadsUp 与 IConFace 已完成 PDF、页数、SHA-256 和全文抽取；Face2Scene、DVFace、MeInTime、IQPFR、DicFace、LRPO、InfoBFR 等仍待逐篇补齐 PDF。官方页快照只证明页面在检索日的公开状态，不代表获得再分发授权。")
    add_paragraph(document, "本轮进一步补充 11 篇 mask conditioning、人体解析、matting、重光照、阴影编辑与虚拟试穿核心 PDF，形成 14 篇逐页证据矩阵和 15 张原文图页；并新增 Mask 分层、区域归一化 LOSS、保护区/边界区约束、mask corruption、多任务训练、消融和手机相机产品策略。下一轮应把综合建议转化为可执行配置，在真实人像 SR 基线中验证梯度冲突、权重范围和设备性能。")

    add_heading(document, "参考文献与核验入口", 1)
    references = [
        "NTIRE 2026: https://arxiv.org/abs/2604.10532",
        "GeoMAR: https://arxiv.org/abs/2608.03923",
        "AuthFace: https://arxiv.org/abs/2410.09864",
        "HDRFace: https://arxiv.org/abs/2605.14821",
        "HeadsUp: https://arxiv.org/abs/2510.09924",
        "IConFace: https://arxiv.org/abs/2605.02814",
        "MeInTime: https://arxiv.org/abs/2603.18645",
        "IQPFR: https://arxiv.org/abs/2503.09294",
        "DicFace: https://arxiv.org/abs/2506.13355",
        "LRPO: https://arxiv.org/abs/2509.23339",
        "InfoBFR: https://arxiv.org/abs/2501.15443",
        "IDFSR: https://arxiv.org/abs/2508.10937",
        "RestorerID: https://arxiv.org/abs/2411.14125",
        "TIGER: https://arxiv.org/abs/2606.24336",
        "SVFR: https://arxiv.org/abs/2501.01235",
        "FADRA: https://arxiv.org/abs/2607.06389",
        "OSDFace: https://arxiv.org/abs/2411.17163",
        "CFRNet: https://arxiv.org/abs/2606.06850",
        "NTIRE 2025: https://arxiv.org/abs/2504.14600",
        "CodeFormer: https://arxiv.org/abs/2206.11253",
        "DifFace: https://arxiv.org/abs/2212.06512",
        "Latent-PMRF: https://arxiv.org/abs/2507.00447",
        "BrushNet: https://arxiv.org/abs/2403.06976",
        "PowerPaint: https://arxiv.org/abs/2312.03594",
        "AnyDoor: https://arxiv.org/abs/2307.09481",
        "CosmicMan: https://arxiv.org/abs/2404.01294",
        "StableVITON: https://arxiv.org/abs/2312.01725",
        "IDM-VTON: https://arxiv.org/abs/2403.05139",
        "Sapiens: https://arxiv.org/abs/2408.12569",
        "MatAnyone: https://arxiv.org/abs/2501.14677",
        "SynthLight: https://arxiv.org/abs/2501.09756",
        "COMPOSE: https://arxiv.org/abs/2406.12013",
        "SoftShadow: https://arxiv.org/abs/2409.07041",
        "SAM 2: https://arxiv.org/abs/2408.00714",
        "Matte Anything: https://arxiv.org/abs/2306.04121",
        "Text2Relight: https://arxiv.org/abs/2412.13734",
        "Generative Portrait Shadow Removal: https://arxiv.org/abs/2410.05525",
    ]
    for reference in references:
        add_bullet(document, reference)

    add_header_footer(document)
    document.core_properties.title = "人像超分与人脸细节恢复阶段性洞察"
    document.core_properties.subject = "Portrait super-resolution and face detail restoration research"
    document.core_properties.author = "ReadPaper"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
