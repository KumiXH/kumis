import hashlib
import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree
from pypdf import PdfReader


REPORT = Path(r"D:\Repository\ReadPaper\daily\PortraitSR\report\人像超分与人脸细节恢复_阶段性洞察_20260806.docx")
ROOT = REPORT.parents[1]
MARKDOWN = REPORT.with_suffix(".md")


def main() -> None:
    with ZipFile(REPORT) as archive:
        bad_entry = archive.testzip()
        names = archive.namelist()
        media = [name for name in names if name.startswith("word/media/")]
        relationships = etree.fromstring(archive.read("word/_rels/document.xml.rels"))
        image_targets = [
            relationship.get("Target")
            for relationship in relationships
            if "image" in (relationship.get("Type") or "")
        ]
        missing = []
        for target in image_targets:
            candidates = {
                "word/" + target.lstrip("/"),
                "word/" + target.replace("../", ""),
            }
            if not any(candidate in names for candidate in candidates):
                missing.append(target)

    document = Document(REPORT)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required = [
        "8. Mask 条件驱动的人像增强与可控编辑",
        "8.4 区域归一化 LOSS",
        "8.6 训练数据如何制作",
        "8.8 最小消融与评测矩阵",
        "L_region = sum_r",
        "SoftShadow",
        "MatAnyone",
        "mask_conditioning_evidence_matrix.json",
    ]
    errors = []
    if bad_entry is not None:
        errors.append(f"DOCX CRC failure: {bad_entry}")
    if missing:
        errors.append(f"missing DOCX image relationships: {missing}")
    if len(document.inline_shapes) != 33:
        errors.append(f"expected 33 DOCX images, got {len(document.inline_shapes)}")
    for item in required:
        if item not in text:
            errors.append(f"missing DOCX text: {item}")

    manifest = json.loads((ROOT / "metadata" / "mask_conditioning_download_manifest.json").read_text(encoding="utf-8"))
    if len(manifest) != 11:
        errors.append(f"expected 11 PDF records, got {len(manifest)}")
    pdf_checks = []
    for record in manifest:
        path = Path(record["local_path"])
        check = {
            "key": record["key"],
            "exists": path.exists(),
            "size_match": False,
            "sha256_match": False,
            "page_count_match": False,
        }
        if path.exists():
            check["size_match"] = path.stat().st_size == record["size"]
            check["sha256_match"] = hashlib.sha256(path.read_bytes()).hexdigest() == record["sha256"]
            check["page_count_match"] = len(PdfReader(str(path)).pages) == record["page_count"]
        if not all(value for key, value in check.items() if key != "key"):
            errors.append(f"PDF validation failed: {record['key']}")
        pdf_checks.append(check)

    evidence = json.loads((ROOT / "metadata" / "mask_conditioning_evidence_matrix.json").read_text(encoding="utf-8"))
    figures = list((ROOT / "figures" / "mask_conditioning").glob("*.png"))
    if len(evidence) != 14:
        errors.append(f"expected 14 evidence records, got {len(evidence)}")
    if len(figures) != 15:
        errors.append(f"expected 15 rendered source pages, got {len(figures)}")

    markdown_text = MARKDOWN.read_text(encoding="utf-8")
    markdown_images = re.findall(r"!\[[^\]]*\]\(([^)]+)\)", markdown_text)
    missing_markdown_images = [
        reference
        for reference in markdown_images
        if not (MARKDOWN.parent / reference).resolve().exists()
    ]
    if missing_markdown_images:
        errors.append(f"missing Markdown images: {missing_markdown_images}")

    print(json.dumps({
        "path": str(REPORT),
        "size": REPORT.stat().st_size,
        "sha256": hashlib.sha256(REPORT.read_bytes()).hexdigest(),
        "errors": errors,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "media_files": len(media),
        "image_relationships": len(image_targets),
        "pdf_checks": pdf_checks,
        "evidence_records": len(evidence),
        "rendered_source_pages": len(figures),
        "markdown_image_references": len(markdown_images),
        "missing_markdown_images": missing_markdown_images,
    }, ensure_ascii=False, indent=2))
    raise SystemExit(1 if errors else 0)


if __name__ == "__main__":
    main()
